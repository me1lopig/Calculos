import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import matplotlib.patches as patches
import matplotlib.path as mpath
import re
import io
import zipfile
from docx import Document
from docx.shared import Inches

# --- CONFIGURACIÓN INICIAL ---
st.set_page_config(page_title="Visor CPTU Profesional", layout="wide")

st.title("📊 Análisis Geotécnico Avanzado de Ensayo CPTU")
st.markdown("Generador de informes técnicos. **Clasificación 9 Zonas, Estado, Dinámica y Propiedades Físicas (Robertson 2010)**.")
st.divider()

# --- DICCIONARIOS Y GEOMETRÍA DE 9 ZONAS ---
SBT_COLORS = {
    1: '#e63926', 2: '#a85c32', 3: '#4f7296', 4: '#5ba48a', 
    5: '#83c393', 6: '#d6b86e', 7: '#c9893d', 8: '#9c9c9c', 
    9: '#c4c4c4', 0: '#ffffff'
}
SBT_NAMES = {
    1: '1. Fino sensitivo', 2: '2. Suelo orgánico', 3: '3. Arcilla',
    4: '4. Limo arcilloso', 5: '5. Arena limosa', 6: '6. Arena limpia',
    7: '7. Grava a arena', 8: '8. Arena muy rígida', 9: '9. Fino muy rígido',
    0: 'Desconocido'
}

EXTENDED_POLYGONS = {
    7: [(-2.0, 4.0), (0.1, 4.0), (-0.3, 2.15), (-2.0, 2.3)],
    6: [(-2.0, 2.3), (-0.3, 2.15), (-0.15, 1.95), (-2.0, 2.0)],
    5: [(-2.0, 2.0), (-0.15, 1.95), (0.05, 1.75), (-2.0, 1.7)],
    4: [(-2.0, 1.7), (0.05, 1.75), (0.3, 1.55), (0.45, 1.45), (-2.0, 1.3)],
    3: [(-2.0, 1.3), (0.45, 1.45), (2.0, 1.2), (2.0, 0.9), (0.0, 0.6), (-2.0, 0.8)],
    2: [(0.0, 0.6), (2.0, 0.9), (2.0, -2.0), (0.2, -2.0)],
    1: [(-2.0, 0.8), (0.0, 0.6), (0.2, -2.0), (-2.0, -2.0)],
    8: [(-0.3, 2.15), (0.1, 4.0), (0.6, 4.0), (0.3, 1.55), (0.05, 1.75), (-0.15, 1.95)],
    9: [(0.3, 1.55), (0.6, 4.0), (2.0, 4.0), (2.0, 1.2), (0.45, 1.45)]
}

# --- FUNCIONES GEOTÉCNICAS AVANZADAS ---
@st.cache_data 
def calcular_geotecnia(df, gwl, a_cone):
    pa = 0.1; gamma_w = 9.81; Gs = 2.65 # Gravedad específica típica
    df_c = df[(df['Qc'] > 0.01) & (df['Rf'] > 0.01)].copy()
    
    # 1. Tensiones In-situ y Parámetros Básicos
    df_c['Qc_pa'] = df_c['Qc'] / pa
    df_c['Gamma_kN3'] = (0.27 * np.log10(df_c['Rf']) + 0.36 * np.log10(df_c['Qc_pa']) + 1.236) * gamma_w
    df_c['Gamma_kN3'] = np.clip(df_c['Gamma_kN3'], 12, 22)
    df_c['dz'] = df_c['Depth_m'].diff().fillna(0)
    df_c['sigma_v0_kPa'] = (df_c['Gamma_kN3'] * df_c['dz']).cumsum()
    
    df_c['u0_kPa'] = np.where(df_c['Depth_m'] > gwl, (df_c['Depth_m'] - gwl) * gamma_w, 0)
    df_c['sigma_v0_eff_kPa'] = df_c['sigma_v0_kPa'] - df_c['u0_kPa']
    df_c['sigma_v0_eff_kPa'] = np.clip(df_c['sigma_v0_eff_kPa'], 1, None)
    
    # 2. Normalización y Bq
    df_c['u2_MPa'] = df_c['U2'] / 1000.0
    df_c['qt_MPa'] = df_c['Qc'] + df_c['u2_MPa'] * (1 - a_cone)
    esfuerzo_neto_kPa = (df_c['qt_MPa'] * 1000) - df_c['sigma_v0_kPa']
    esfuerzo_neto_kPa = np.clip(esfuerzo_neto_kPa, 1.0, None)
    
    df_c['Qt'] = esfuerzo_neto_kPa / df_c['sigma_v0_eff_kPa']
    df_c['Qt'] = np.clip(df_c['Qt'], 1, 1000)
    df_c['Fr_percent'] = ((df_c['Fs'] / 1000.0) / (esfuerzo_neto_kPa / 1000.0)) * 100.0
    df_c['Fr_percent'] = np.clip(df_c['Fr_percent'], 0.01, 10)
    
    df_c['Bq'] = (df_c['U2'] - df_c['u0_kPa']) / esfuerzo_neto_kPa
    df_c['Bq'] = np.clip(df_c['Bq'], -1.0, 2.0)
    
    # 3. Índice Ic y Zonas SBT
    df_c['Ic'] = np.sqrt((3.47 - np.log10(df_c['Qt']))**2 + (np.log10(df_c['Fr_percent']) + 1.22)**2)
    points = np.column_stack((np.log10(df_c['Fr_percent']), np.log10(df_c['Qt'])))
    zones = np.zeros(len(df_c), dtype=int)
    for zone, coords in EXTENDED_POLYGONS.items():
        path = mpath.Path(coords)
        mask = path.contains_points(points, radius=1e-5)
        zones[mask] = zone
    df_c['SBT_Zone'] = np.where(zones == 0, 0, zones)
    df_c['SBT_Name'] = df_c['SBT_Zone'].map(SBT_NAMES)
    
    # 4. Diseño y Deformación
    df_c['Su_kPa'] = np.where(df_c['Ic'] > 2.6, esfuerzo_neto_kPa / 14.0, np.nan)
    df_c['Phi_deg'] = np.where(df_c['Ic'] <= 2.6, 17.6 + 11.0 * np.log10(df_c['Qt']), np.nan)
    df_c['N60'] = df_c['Qc_pa'] / (10 ** (1.1268 - 0.2817 * df_c['Ic']))
    df_c['Dr_percent'] = np.where(df_c['Ic'] <= 2.6, np.sqrt(df_c['Qt'] / 350.0) * 100.0, np.nan)
    df_c['OCR'] = np.where(df_c['Ic'] > 2.6, 0.33 * df_c['Qt'], np.nan)
    
    alpha = 0.0188 * (10 ** (0.55 * df_c['Ic'] + 1.68))
    df_c['M_MPa'] = alpha * (esfuerzo_neto_kPa / 1000.0)
    alpha_E = 0.015 * (10 ** (0.55 * df_c['Ic'] + 1.68))
    df_c['Es_MPa'] = alpha_E * (esfuerzo_neto_kPa / 1000.0)

    # 5. Estado, Dinámica e Hidráulica
    alpha_vs = 10 ** (0.55 * df_c['Ic'] + 1.68)
    df_c['Vs_ms'] = (alpha_vs * (esfuerzo_neto_kPa / (pa * 1000))) ** 0.5
    df_c['G0_MPa'] = (df_c['Gamma_kN3'] / gamma_w) * (df_c['Vs_ms'] ** 2) / 1000.0
    df_c['St'] = np.clip(7.0 / df_c['Fr_percent'], 1, 100)
    df_c['Psi'] = np.where(df_c['Ic'] <= 2.6, 0.56 - 0.33 * np.log10(df_c['Qt']), np.nan)
    df_c['K0'] = np.clip(0.1 * df_c['Qt'], 0.1, 4.0)
    df_c['k_ms'] = np.where(df_c['Ic'] < 3.27, 10 ** (0.952 - 3.04 * df_c['Ic']), 10 ** (-4.52 - 1.37 * df_c['Ic']))

    # 6. Propiedades Físicas / Índice (Estimaciones Empíricas)
    # Asumimos suelo saturado para derivar el índice de poros (e) desde el peso unitario
    df_c['e_void'] = np.clip((Gs * gamma_w - df_c['Gamma_kN3']) / (df_c['Gamma_kN3'] - gamma_w), 0.2, 3.0)
    df_c['Gamma_dry_kN3'] = (Gs * gamma_w) / (1 + df_c['e_void'])
    df_c['w_per'] = np.clip((df_c['e_void'] / Gs) * 100.0, 5, 100)
    df_c['Ip_per'] = np.where(df_c['Ic'] > 2.2, np.clip((df_c['Ic'] - 2.2) * 25.0, 0, 100), np.nan)

    return df_c

LEYENDA_SUELOS = [patches.Rectangle((0,0),1,1, color=SBT_COLORS[k], label=SBT_NAMES[k]) for k in range(1, 10)]

def plot_stratigraphy_col(ax, df, preforo):
    if preforo > 0: ax.axhspan(0, preforo, color='gray', alpha=0.3, hatch='//')
    for zone in df['SBT_Zone'].unique():
        if zone == 0: continue
        ax.fill_betweenx(df['Depth_m'], 0, 1, where=df['SBT_Zone']==zone, color=SBT_COLORS.get(zone, '#fff'), step='mid')
    ax.set_xlim(0, 1); ax.set_xticks([]); ax.set_ylabel('Profundidad (m)', fontsize=11, fontweight='bold')
    ax.invert_yaxis(); ax.set_title('Estratigrafía', fontsize=10, fontweight='bold')

# --- INTERFAZ PRINCIPAL ---
uploaded_file = st.file_uploader("📂 Sube el archivo CPTU (.CSV)", type=["csv", "CSV"])

if uploaded_file is not None:
    content = uploaded_file.read().decode('utf-8').splitlines()
    header_data = {}
    comentario_preforo = 0.0 
    for line in content[:20]:
        if ';' in line:
            key, val = line.split(';', 1)
            clean_key = key.strip().rstrip(':'); clean_val = val.strip().strip(';')
            if clean_key and clean_val:
                header_data[clean_key] = clean_val
                if clean_key.lower() == 'comments':
                    match = re.search(r'(\d+(?:[.,]\d+)?)', clean_val)
                    if match: comentario_preforo = float(match.group(1).replace(',', '.'))
    
    with st.expander("📋 Ver Datos de la Campaña", expanded=False):
        items = list(header_data.items()); mitad = len(items)//2 + len(items)%2
        c1, c2 = st.columns(2)
        c1.table(pd.DataFrame(items[:mitad], columns=["Parámetro", "Valor"]).set_index("Parámetro"))
        c2.table(pd.DataFrame(items[mitad:], columns=["Parámetro", "Valor"]).set_index("Parámetro"))
                
    uploaded_file.seek(0)
    df = pd.read_csv(uploaded_file, sep=';', decimal=',', skiprows=23)
    df['Depth_m'] = df['Depth'] / 100.0 
    
    cota_analitica = df[df['Qc'] > 0.05]['Depth_m'].min() if not df[df['Qc'] > 0.05].empty else 0.0
    st.sidebar.header("⚙️ Configuración")
    cota_preforo = st.sidebar.number_input("Preforo (m)", 0.0, float(df['Depth_m'].max()), float(max(cota_analitica, comentario_preforo)))
    gwl = st.sidebar.number_input("Nivel Freático Estimado (m)", 0.0, float(df['Depth_m'].max()), 2.0, step=0.1)
    a_cone = st.sidebar.number_input("Relación de Área Neta Cono (a)", 0.50, 1.00, 0.80, step=0.01)
    
    df_calc = calcular_geotecnia(df, gwl, a_cone)
    
    # --- PESTAÑAS AMPLIADAS (9 Pestañas) ---
    tab_b, tab_r, tab_d, tab_est, tab_fis, tab_c, tab_cap, tab_det, tab_f = st.tabs([
        "📉 Básicos e Hidro", "🏗️ Resistencia", "📉 Deformación", "🌊 Estado y Dinámica", "🧪 Prop. Físicas", "🕵️ Calidad", "📑 Capas", "📋 Detalles", "📚 Formulación"
    ])
    leyenda_preforo = [patches.Rectangle((0,0),1,1, color='gray', alpha=0.3, hatch='//', label='PREFORO')] + LEYENDA_SUELOS if cota_preforo > 0 else LEYENDA_SUELOS

    def generar_figura_perfil(tipo):
        # Aumentamos a 7 columnas para que quepan más curvas por gráfico
        fig, axs = plt.subplots(1, 7, figsize=(18, 8), sharey=True, gridspec_kw={'width_ratios': [1, 2, 2, 2, 2, 2, 1.5]})
        plot_stratigraphy_col(axs[0], df_calc, cota_preforo)
        
        if tipo == 'basicos':
            axs[1].plot(df_calc['qt_MPa'], df_calc['Depth_m'], '#1f77b4', lw=1); axs[1].set_title('qt corregido (MPa)')
            axs[2].plot(df_calc['Fr_percent'], df_calc['Depth_m'], '#ff7f0e', lw=1); axs[2].set_xlim(0, 10); axs[2].set_title('Fricción Fr (%)')
            axs[3].plot(df_calc['U2'], df_calc['Depth_m'], '#d62728', lw=1); axs[3].set_title('U2 (kPa)')
            axs[4].plot(df_calc['Bq'], df_calc['Depth_m'], 'teal', lw=1); axs[4].set_xlim(-0.2, 1.5); axs[4].set_title('Pore Pres. Bq')
            axs[5].plot(df_calc['Ic'], df_calc['Depth_m'], '#2ca02c', lw=1); axs[5].axvline(2.6, color='red', ls='--'); axs[5].set_title('SBTn Index (Ic)')
        
        elif tipo == 'resistencia':
            axs[1].plot(df_calc['Su_kPa'], df_calc['Depth_m'], 'brown', lw=1.5); axs[1].set_title('Su (kPa)')
            axs[2].plot(df_calc['Phi_deg'], df_calc['Depth_m'], 'orange', lw=1.5); axs[2].set_title('Phi (°)')
            axs[3].plot(df_calc['Dr_percent'], df_calc['Depth_m'], 'olive', lw=1.5); axs[3].set_xlim(0, 100); axs[3].set_title('Dr (%)')
            axs[4].plot(df_calc['N60'], df_calc['Depth_m'], 'black', lw=1.5); axs[4].set_title('SPT N60 Eq.')
            axs[5].axis('off')
        
        elif tipo == 'deformacion':
            axs[1].plot(df_calc['M_MPa'], df_calc['Depth_m'], 'navy', lw=1.5); axs[1].set_title('Módulo M (MPa)')
            axs[2].plot(df_calc['Es_MPa'], df_calc['Depth_m'], 'blue', lw=1.5); axs[2].set_title('Módulo Es (MPa)')
            axs[3].plot(df_calc['OCR'], df_calc['Depth_m'], 'magenta', lw=1.5); axs[3].axvline(1.0, color='red', ls='--'); axs[3].set_title('OCR')
            axs[4].axis('off'); axs[5].axis('off')
        
        elif tipo == 'estado':
            axs[1].plot(df_calc['Vs_ms'], df_calc['Depth_m'], 'darkcyan', lw=1.5); axs[1].set_title('Vs (m/s)')
            axs[2].plot(df_calc['G0_MPa'], df_calc['Depth_m'], 'darkblue', lw=1.5); axs[2].set_title('Mód. Corte G0 (MPa)')
            axs[3].plot(df_calc['K0'], df_calc['Depth_m'], 'indigo', lw=1.5); axs[3].set_title('K0 In-situ')
            axs[4].plot(df_calc['Psi'], df_calc['Depth_m'], 'darkgoldenrod', lw=1.5); axs[4].axvline(0.0, color='black', ls='--'); axs[4].set_title('Estado Ψ')
            axs[5].plot(df_calc['k_ms'], df_calc['Depth_m'], 'dodgerblue', lw=1.5); axs[5].set_xscale('log'); axs[5].set_xlim(1e-10, 1e-1); axs[5].set_title('Permeab. k (m/s)')

        elif tipo == 'fisicas':
            axs[1].plot(df_calc['Gamma_kN3'], df_calc['Depth_m'], 'purple', lw=1.5, label='Húmedo'); axs[1].plot(df_calc['Gamma_dry_kN3'], df_calc['Depth_m'], 'violet', lw=1.5, ls='--', label='Seco'); axs[1].legend(loc='lower left', fontsize=8); axs[1].set_title('Pesos \u03B3 (kN/m³)')
            axs[2].plot(df_calc['w_per'], df_calc['Depth_m'], 'deepskyblue', lw=1.5); axs[2].set_xlim(0, 80); axs[2].set_title('Humedad w (%)')
            axs[3].plot(df_calc['e_void'], df_calc['Depth_m'], 'saddlebrown', lw=1.5); axs[3].set_xlim(0, 2); axs[3].set_title('Índice Poros (e)')
            axs[4].plot(df_calc['Ip_per'], df_calc['Depth_m'], 'green', lw=1.5); axs[4].set_xlim(0, 100); axs[4].set_title('Ind. Plasticidad Ip (%)')
            axs[5].axis('off')

        elif tipo == 'calidad':
            axs[1].plot(df_calc['Tilt'], df_calc['Depth_m'], 'red', lw=1.5); axs[1].axvline(15, color='k', ls='--'); axs[1].set_title('Tilt (°)')
            axs[2].plot(df_calc['Speed'], df_calc['Depth_m'], 'teal', lw=1.5); axs[2].axvline(2.0, color='g', ls='-', lw=2); axs[2].set_title('Speed (cm/s)')
            axs[3].axis('off'); axs[4].axis('off'); axs[5].axis('off')

        for i in range(1, 6):
            if axs[i].axison:
                axs[i].grid(True, ls='--', alpha=0.5)
                if cota_preforo > 0: axs[i].axhspan(0, cota_preforo, color='gray', alpha=0.3, hatch='//')

        axs[6].axis('off'); axs[6].legend(handles=leyenda_preforo, loc='center', title="Robertson 1990", fontsize=8)
        plt.tight_layout()
        return fig

    with tab_b: st.pyplot(generar_figura_perfil('basicos'))
    with tab_r: st.pyplot(generar_figura_perfil('resistencia'))
    with tab_d: st.pyplot(generar_figura_perfil('deformacion'))
    with tab_est: st.pyplot(generar_figura_perfil('estado'))
    with tab_fis: st.pyplot(generar_figura_perfil('fisicas')) # NUEVA PESTAÑA FÍSICA
    with tab_c: st.pyplot(generar_figura_perfil('calidad'))
    
    # --- TABLAS Y EXPORTACIÓN ---
    df_v = df_calc[df_calc['Depth_m'] >= cota_preforo].copy()
    df_v['Depth_Interval'] = np.floor(df_v['Depth_m'])
    
    with tab_cap: 
        res_1m = df_v.groupby('Depth_Interval').agg(SBT_Predominante=('SBT_Name', lambda x: x.mode()[0]), Qc_Medio=('qt_MPa', 'mean'), Z_min=('Depth_m', 'min'), Z_max=('Depth_m', 'max')).reset_index()
        res_1m = res_1m[res_1m['Qc_Medio'] > 0.05].copy()
        res_1m['Grupo'] = (res_1m['SBT_Predominante'] != res_1m['SBT_Predominante'].shift()).cumsum()
        capas = res_1m.groupby(['Grupo', 'SBT_Predominante']).agg(Desde_m=('Z_min', 'min'), Hasta_m=('Z_max', 'max')).reset_index()
        capas['Espesor_m'] = capas['Hasta_m'] - capas['Desde_m']
        st.dataframe(capas[['Desde_m', 'Hasta_m', 'Espesor_m', 'SBT_Predominante']], hide_index=True)
        
    with tab_det: 
        st.dataframe(df_v[['Depth_m', 'SBT_Name', 'Ic', 'qt_MPa', 'Fr_percent', 'Bq', 'Su_kPa', 'Phi_deg', 'Dr_percent', 'M_MPa', 'Es_MPa', 'Vs_ms', 'G0_MPa', 'k_ms', 'w_per', 'e_void', 'Ip_per']], hide_index=True)
    
    with tab_f:
        st.subheader("📚 Metodología y Formulación Geotécnica Completa")
        st.write("Ecuaciones de normalización, parámetros dinámicos y propiedades físicas empíricas (Suelo Saturado, $G_s = 2.65$).")
        st.divider()

        col_A, col_B = st.columns(2)
        with col_A:
            st.markdown("### 1. Parámetros Normalizados y Bq")
            st.latex(r"Q_t = \frac{q_t - \sigma_{v0}}{\sigma'_{v0}} \quad ; \quad F_r (\%) = \left[ \frac{f_s}{q_t - \sigma_{v0}} \right] \cdot 100")
            st.latex(r"B_q = \frac{u_2 - u_0}{q_t - \sigma_{v0}}")
            st.markdown("### 2. Módulos y Dinámica ($E_s$, $V_s$, $G_0$)")
            st.latex(r"E_s = \alpha_E \cdot (q_t - \sigma_{v0}) \quad (\text{Robertson 2009})")
            st.latex(r"V_s = \left[ \left(10^{0.55 I_c + 1.68}\right) \frac{q_t - \sigma_{v0}}{p_a} \right]^{0.5}")
            st.latex(r"G_0 = \rho \cdot V_s^2")
            st.markdown("### 3. Permeabilidad ($k$)")
            st.latex(r"k = 10^{(0.952 - 3.04 \cdot I_c)} \quad \text{si } I_c < 3.27")

        with col_B:
            st.markdown("### 4. Parámetros de Estado")
            st.latex(r"S_t = \frac{7.0}{F_r} \quad ; \quad K_0 = 0.1 \cdot Q_t")
            st.latex(r"\psi = 0.56 - 0.33 \log_{10}(Q_t)")
            st.markdown("### 5. Propiedades Físicas / Índice")
            st.write("Deducidas asumiendo suelo 100% saturado por debajo del N.F.")
            st.latex(r"e = \frac{G_s \cdot \gamma_w - \gamma_{sat}}{\gamma_{sat} - \gamma_w} \quad ; \quad w (\%) = \frac{e}{G_s} \cdot 100")
            st.latex(r"\gamma_{dry} = \frac{G_s \cdot \gamma_w}{1 + e}")
            st.latex(r"I_p (\%) = 25 \cdot (I_c - 2.2) \quad \text{(Estimación empírica)}")

    # --- EXPORTACIÓN ---
    st.sidebar.divider()
    if st.sidebar.button("🚀 GENERAR INFORME", use_container_width=True):
        progress = st.sidebar.progress(0)
        status = st.sidebar.empty()
        
        status.text("Creando Informe Word...")
        doc = Document()
        doc.add_heading(f'Informe Geotécnico Avanzado: {header_data.get("Location", "CPTU")}', 0)
        
        # Añadimos todas las pestañas al Word
        for i, t in enumerate(['basicos', 'resistencia', 'deformacion', 'estado', 'fisicas', 'calidad']):
            fig = generar_figura_perfil(t)
            buf = io.BytesIO(); fig.savefig(buf, format='png', dpi=150)
            doc.add_heading(f'Perfil de {t.capitalize()}', level=1)
            doc.add_picture(buf, width=Inches(6.5))
            progress.progress(15 + i*15)

        status.text("Creando Excel Multihoja...")
        excel_buf = io.BytesIO()
        with pd.ExcelWriter(excel_buf, engine='xlsxwriter') as writer:
            pd.DataFrame(list(header_data.items())).to_excel(writer, sheet_name='Metadatos')
            df_v[['Depth_m', 'SBT_Name', 'Ic', 'qt_MPa', 'Fr_percent', 'Bq', 'Su_kPa', 'Phi_deg', 'Dr_percent', 'M_MPa', 'Es_MPa', 'Vs_ms', 'G0_MPa', 'St', 'K0', 'Psi', 'k_ms', 'e_void', 'w_per', 'Gamma_dry_kN3', 'Ip_per']].to_excel(writer, sheet_name='Datos_Completos', index=False)

        zip_buf = io.BytesIO()
        with zipfile.ZipFile(zip_buf, 'w') as zf:
            w_buf = io.BytesIO(); doc.save(w_buf)
            zf.writestr("Informe_Word_CPTU.docx", w_buf.getvalue())
            zf.writestr("Libro_Calculo_CPTU.xlsx", excel_buf.getvalue())
        
        progress.progress(100)
        status.success("¡Informe Creado!")
        st.sidebar.download_button("📥 DESCARGAR INFORME (.ZIP)", zip_buf.getvalue(), "Resultados_CPTU.zip", "application/zip", use_container_width=True)