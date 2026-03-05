import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import matplotlib.patches as patches
import matplotlib.path as mpath
import re
import io
import zipfile
from docx import Document
from docx.shared import Inches

# --- CONFIGURACIÓN INICIAL ---
st.set_page_config(page_title="Visor CPTU Profesional", layout="wide")

st.title("📊 Análisis Geotécnico Avanzado de Ensayo CPTU")
st.markdown("Generador de informes técnicos y parámetros de diseño. **Clasificación exacta 9 Zonas (Robertson 1990)**.")
st.divider()

# --- DICCIONARIOS Y GEOMETRÍA DE 9 ZONAS (ROBERTSON 1990) ---
SBT_COLORS = {
    1: '#e63926', 2: '#a85c32', 3: '#4f7296', 4: '#5ba48a', 
    5: '#83c393', 6: '#d6b86e', 7: '#c9893d', 8: '#9c9c9c', 
    9: '#c4c4c4', 0: '#ffffff'
}

SBT_NAMES = {
    1: '1. Fino sensitivo', 2: '2. Suelo orgánico', 3: '3. Arcilla',
    4: '4. Limo arcilloso', 5: '5. Arena limosa', 6: '6. Arena limpia',
    7: '7. Grava a arena', 8: '8. Arena muy rígida', 9: '9. Fino muy rígido',
    0: 'Desconocido'
}

# Polígonos matemáticos exactos (log Fr, log Qt)
EXTENDED_POLYGONS = {
    7: [(-2.0, 4.0), (0.1, 4.0), (-0.3, 2.15), (-2.0, 2.3)],
    6: [(-2.0, 2.3), (-0.3, 2.15), (-0.15, 1.95), (-2.0, 2.0)],
    5: [(-2.0, 2.0), (-0.15, 1.95), (0.05, 1.75), (-2.0, 1.7)],
    4: [(-2.0, 1.7), (0.05, 1.75), (0.3, 1.55), (0.45, 1.45), (-2.0, 1.3)],
    3: [(-2.0, 1.3), (0.45, 1.45), (2.0, 1.2), (2.0, 0.9), (0.0, 0.6), (-2.0, 0.8)],
    2: [(0.0, 0.6), (2.0, 0.9), (2.0, -2.0), (0.2, -2.0)],
    1: [(-2.0, 0.8), (0.0, 0.6), (0.2, -2.0), (-2.0, -2.0)],
    8: [(-0.3, 2.15), (0.1, 4.0), (0.6, 4.0), (0.3, 1.55), (0.05, 1.75), (-0.15, 1.95)],
    9: [(0.3, 1.55), (0.6, 4.0), (2.0, 4.0), (2.0, 1.2), (0.45, 1.45)]
}

# --- FUNCIONES GEOTÉCNICAS ---
@st.cache_data 
def calcular_geotecnia(df, gwl, a_cone):
    pa = 0.1; gamma_w = 9.81
    df_c = df[(df['Qc'] > 0.01) & (df['Rf'] > 0.01)].copy()
    
    # 1. Tensiones In-situ
    df_c['Qc_pa'] = df_c['Qc'] / pa
    df_c['Gamma_kN3'] = (0.27 * np.log10(df_c['Rf']) + 0.36 * np.log10(df_c['Qc_pa']) + 1.236) * gamma_w
    df_c['Gamma_kN3'] = np.clip(df_c['Gamma_kN3'], 12, 22)
    df_c['dz'] = df_c['Depth_m'].diff().fillna(0)
    df_c['sigma_v0_kPa'] = (df_c['Gamma_kN3'] * df_c['dz']).cumsum()
    
    df_c['u0_kPa'] = np.where(df_c['Depth_m'] > gwl, (df_c['Depth_m'] - gwl) * gamma_w, 0)
    df_c['sigma_v0_eff_kPa'] = df_c['sigma_v0_kPa'] - df_c['u0_kPa']
    df_c['sigma_v0_eff_kPa'] = np.clip(df_c['sigma_v0_eff_kPa'], 1, None)
    
    # 2. Normalización
    df_c['u2_MPa'] = df_c['U2'] / 1000.0
    df_c['qt_MPa'] = df_c['Qc'] + df_c['u2_MPa'] * (1 - a_cone)
    df_c['Qt'] = ((df_c['qt_MPa'] * 1000) - df_c['sigma_v0_kPa']) / df_c['sigma_v0_eff_kPa']
    df_c['Qt'] = np.clip(df_c['Qt'], 1, 1000)
    
    esfuerzo_neto = np.abs(df_c['qt_MPa'] - (df_c['sigma_v0_kPa'] / 1000.0))
    esfuerzo_neto = np.clip(esfuerzo_neto, 0.001, None)
    df_c['Fr_percent'] = ((df_c['Fs'] / 1000.0) / esfuerzo_neto) * 100.0
    df_c['Fr_percent'] = np.clip(df_c['Fr_percent'], 0.01, 10)
    
    # 3. Mapeo por Polígonos Exactos (Calculadora masiva)
    points = np.column_stack((np.log10(df_c['Fr_percent']), np.log10(df_c['Qt'])))
    zones = np.zeros(len(df_c), dtype=int)
    for zone, coords in EXTENDED_POLYGONS.items():
        path = mpath.Path(coords)
        mask = path.contains_points(points, radius=1e-5)
        zones[mask] = zone
        
    df_c['SBT_Zone'] = np.where(zones == 0, 0, zones)
    df_c['SBT_Name'] = df_c['SBT_Zone'].map(SBT_NAMES)
    
    # Índice Ic continuo para cálculos empíricos
    df_c['Ic'] = np.sqrt((3.47 - np.log10(df_c['Qt']))**2 + (np.log10(df_c['Fr_percent']) + 1.22)**2)
    
    # 4. Parámetros de Diseño y Estado
    Nkt = 14 
    df_c['Su_kPa'] = np.where(df_c['Ic'] > 2.6, ((df_c['qt_MPa'] * 1000) - df_c['sigma_v0_kPa']) / Nkt, np.nan)
    df_c['Phi_deg'] = np.where(df_c['Ic'] <= 2.6, 17.6 + 11.0 * np.log10(df_c['Qt']), np.nan)
    df_c['N60'] = df_c['Qc_pa'] / (10 ** (1.1268 - 0.2817 * df_c['Ic']))
    df_c['Dr_percent'] = np.where(df_c['Ic'] <= 2.6, np.sqrt(df_c['Qt'] / 350.0) * 100.0, np.nan)
    df_c['Dr_percent'] = np.clip(df_c['Dr_percent'], 0, 100)
    df_c['OCR'] = np.where(df_c['Ic'] > 2.6, 0.33 * df_c['Qt'], np.nan)
    
    alpha = 0.0188 * (10 ** (0.55 * df_c['Ic'] + 1.68))
    df_c['M_MPa'] = alpha * (df_c['qt_MPa'] - (df_c['sigma_v0_kPa'] / 1000.0))
    alpha_E = 0.015 * (10 ** (0.55 * df_c['Ic'] + 1.68))
    df_c['Es_MPa'] = alpha_E * (df_c['qt_MPa'] - (df_c['sigma_v0_kPa'] / 1000.0))
    
    return df_c

LEYENDA_SUELOS = [patches.Rectangle((0,0),1,1, color=SBT_COLORS[k], label=SBT_NAMES[k]) for k in range(1, 10)]

def plot_stratigraphy_col(ax, df, preforo):
    if preforo > 0: ax.axhspan(0, preforo, color='gray', alpha=0.3, hatch='//')
    for zone in df['SBT_Zone'].unique():
        if zone == 0: continue
        ax.fill_betweenx(df['Depth_m'], 0, 1, where=df['SBT_Zone']==zone, color=SBT_COLORS.get(zone, '#fff'), step='mid')
    ax.set_xlim(0, 1); ax.set_xticks([]); ax.set_ylabel('Profundidad (m)', fontsize=11, fontweight='bold')
    ax.invert_yaxis(); ax.set_title('Estratigrafía', fontsize=10, fontweight='bold')

def generar_grafico_distribucion(df, cota_preforo):
    df_v = df[df['Depth_m'] >= cota_preforo].copy()
    if df_v.empty: return plt.subplots(figsize=(8, 4))[0]
    
    # Agrupar por zona y sumar los espesores (dz)
    dist = df_v.groupby('SBT_Zone')['dz'].sum().reset_index()
    dist.columns = ['Zona', 'Metros']
    dist = dist[dist['Zona'] != 0] # Filtrar desconocidos
    dist['Porcentaje'] = (dist['Metros'] / dist['Metros'].sum()) * 100
    dist = dist.sort_values(by='Zona', ascending=False)
    
    fig, ax = plt.subplots(figsize=(10, 5))
    bars = ax.barh(dist['Zona'].map(SBT_NAMES), dist['Porcentaje'], color=dist['Zona'].map(SBT_COLORS), edgecolor='black')
    
    for bar in bars:
        width = bar.get_width()
        ax.text(width + 0.5, bar.get_y() + bar.get_height()/2, f'{width:.1f}%', va='center', fontsize=10, fontweight='bold')
        
    ax.set_xlabel('Porcentaje del Sondeo (%)', fontsize=11, fontweight='bold')
    ax.set_title('Distribución de Tipos de Suelo (SBT 9 Zonas)', fontsize=13, fontweight='bold')
    ax.grid(True, axis='x', ls='--', alpha=0.5)
    ax.set_xlim(0, max(dist['Porcentaje']) * 1.15) # Dar espacio al texto
    plt.tight_layout()
    return fig

# --- INTERFAZ PRINCIPAL ---
uploaded_file = st.file_uploader("📂 Sube el archivo CPTU (.CSV)", type=["csv", "CSV"])

if uploaded_file is not None:
    content = uploaded_file.read().decode('utf-8').splitlines()
    header_data = {}
    comentario_preforo = 0.0 
    for line in content[:20]:
        if ';' in line:
            key, val = line.split(';', 1)
            clean_key = key.strip().rstrip(':'); clean_val = val.strip().strip(';')
            if clean_key and clean_val:
                header_data[clean_key] = clean_val
                if clean_key.lower() == 'comments':
                    match = re.search(r'(\d+(?:[.,]\d+)?)', clean_val)
                    if match: comentario_preforo = float(match.group(1).replace(',', '.'))
    
    with st.expander("📋 Ver Datos de la Campaña", expanded=False):
        items = list(header_data.items()); mitad = len(items)//2 + len(items)%2
        c1, c2 = st.columns(2)
        c1.table(pd.DataFrame(items[:mitad], columns=["Parámetro", "Valor"]).set_index("Parámetro"))
        c2.table(pd.DataFrame(items[mitad:], columns=["Parámetro", "Valor"]).set_index("Parámetro"))
                
    uploaded_file.seek(0)
    df = pd.read_csv(uploaded_file, sep=';', decimal=',', skiprows=23)
    df['Depth_m'] = df['Depth'] / 100.0 
    
    cota_analitica = df[df['Qc'] > 0.05]['Depth_m'].min() if not df[df['Qc'] > 0.05].empty else 0.0
    st.sidebar.header("⚙️ Configuración")
    cota_preforo = st.sidebar.number_input("Preforo (m)", 0.0, float(df['Depth_m'].max()), float(max(cota_analitica, comentario_preforo)))
    gwl = st.sidebar.number_input("Nivel Freático Estimado (m)", 0.0, float(df['Depth_m'].max()), 2.0, step=0.1)
    a_cone = st.sidebar.number_input("Relación de Área Neta Cono (a)", 0.50, 1.00, 0.80, step=0.01)
    
    df_calc = calcular_geotecnia(df, gwl, a_cone)
    
    # --- PESTAÑAS (Añadimos Distribución) ---
    tab_b, tab_r, tab_d, tab_dist, tab_c, tab_cap, tab_det, tab_f = st.tabs([
        "📉 Básicos", "🏗️ Resistencia", "📉 Deformación", "📊 Distribución SBT", "🕵️ Calidad", "📑 Capas", "📋 Detalles", "📚 Formulación"
    ])
    leyenda_preforo = [patches.Rectangle((0,0),1,1, color='gray', alpha=0.3, hatch='//', label='PREFORO')] + LEYENDA_SUELOS if cota_preforo > 0 else LEYENDA_SUELOS

    def generar_figura_perfil(tipo):
        fig, axs = plt.subplots(1, 6, figsize=(16, 8), sharey=True, gridspec_kw={'width_ratios': [1, 2, 2, 2, 2, 1.5]})
        plot_stratigraphy_col(axs[0], df_calc, cota_preforo)
        
        if tipo == 'basicos':
            axs[1].plot(df_calc['qt_MPa'], df_calc['Depth_m'], '#1f77b4', lw=1); axs[1].set_title('qt corregido (MPa)')
            axs[2].plot(df_calc['Fs'], df_calc['Depth_m'], '#ff7f0e', lw=1); axs[2].set_title('Fs (kPa)')
            axs[3].plot(df_calc['U2'], df_calc['Depth_m'], '#d62728', lw=1); axs[3].set_title('U2 (kPa)')
            axs[4].plot(df_calc['Ic'], df_calc['Depth_m'], '#2ca02c', lw=1); axs[4].axvline(2.6, color='red', ls='--'); axs[4].set_title('SBTn Index (Ic)')
        elif tipo == 'resistencia':
            axs[1].plot(df_calc['Su_kPa'], df_calc['Depth_m'], 'brown', lw=1.5); axs[1].set_title('Su (kPa)')
            axs[2].plot(df_calc['Phi_deg'], df_calc['Depth_m'], 'orange', lw=1.5); axs[2].set_title('Phi (°)')
            axs[3].plot(df_calc['Dr_percent'], df_calc['Depth_m'], 'olive', lw=1.5); axs[3].set_xlim(0, 100); axs[3].set_title('Dr (%)')
            axs[4].plot(df_calc['N60'], df_calc['Depth_m'], 'black', lw=1.5); axs[4].set_title('SPT N60 Eq.')
        elif tipo == 'deformacion':
            axs[1].plot(df_calc['M_MPa'], df_calc['Depth_m'], 'navy', lw=1.5); axs[1].set_title('M (MPa)')
            axs[2].plot(df_calc['Es_MPa'], df_calc['Depth_m'], 'blue', lw=1.5); axs[2].set_title('Es (MPa)')
            axs[3].plot(df_calc['Gamma_kN3'], df_calc['Depth_m'], 'purple', lw=1.5); axs[3].set_title('\u03B3 (kN/m³)')
            axs[4].plot(df_calc['OCR'], df_calc['Depth_m'], 'magenta', lw=1.5); axs[4].axvline(1.0, color='red', ls='--'); axs[4].set_title('OCR')
        elif tipo == 'calidad':
            axs[1].plot(df_calc['Tilt'], df_calc['Depth_m'], 'red', lw=1.5); axs[1].axvline(15, color='k', ls='--'); axs[1].set_title('Tilt (°)')
            axs[2].plot(df_calc['Speed'], df_calc['Depth_m'], 'teal', lw=1.5); axs[2].axvline(2.0, color='g', ls='-', lw=2); axs[2].set_title('Speed (cm/s)')
            axs[3].axis('off'); axs[4].axis('off')

        for i in range(1, 5):
            if axs[i].axison:
                axs[i].grid(True, ls='--', alpha=0.5)
                if cota_preforo > 0: axs[i].axhspan(0, cota_preforo, color='gray', alpha=0.3, hatch='//')

        axs[5].axis('off'); axs[5].legend(handles=leyenda_preforo, loc='center', title="Robertson 1990", fontsize=8)
        plt.tight_layout()
        return fig

    with tab_b: st.pyplot(generar_figura_perfil('basicos'))
    with tab_r: st.pyplot(generar_figura_perfil('resistencia'))
    with tab_d: st.pyplot(generar_figura_perfil('deformacion'))
    with tab_c: st.pyplot(generar_figura_perfil('calidad'))
    
    # --- PESTAÑA: DISTRIBUCIÓN ESTADÍSTICA ---
    with tab_dist:
        st.subheader("Estadística de Perfil de Suelo (Omitiendo Preforo)")
        st.markdown("Porcentaje de cada tipo de material detectado en el sondeo en base a la geometría de polígonos de 9 zonas.")
        fig_dist = generar_grafico_distribucion(df_calc, cota_preforo)
        st.pyplot(fig_dist)

    # --- TABLAS DE RESULTADOS ---
    df_v = df_calc[df_calc['Depth_m'] >= cota_preforo].copy()
    df_v['Depth_Interval'] = np.floor(df_v['Depth_m'])
    res_1m = df_v.groupby('Depth_Interval').agg(SBT_Predominante=('SBT_Name', lambda x: x.mode()[0]), Qc_Medio=('qt_MPa', 'mean'), U2_Medio=('U2', 'mean'), Su_Medio=('Su_kPa', 'mean'), Phi_Medio=('Phi_deg', 'mean'), Dr_Medio=('Dr_percent', 'mean'), Es_Medio=('Es_MPa', 'mean'), M_Medio=('M_MPa', 'mean'), OCR_Medio=('OCR', 'mean'), Z_min=('Depth_m', 'min'), Z_max=('Depth_m', 'max')).reset_index()
    res_1m = res_1m[res_1m['Qc_Medio'] > 0.05].copy()
    res_1m['Grupo'] = (res_1m['SBT_Predominante'] != res_1m['SBT_Predominante'].shift()).cumsum()
    capas = res_1m.groupby(['Grupo', 'SBT_Predominante']).agg(Desde_m=('Z_min', 'min'), Hasta_m=('Z_max', 'max'), qt_MPa=('Qc_Medio', 'mean'), Su_kPa=('Su_Medio', 'mean'), Phi_deg=('Phi_Medio', 'mean'), Dr_per=('Dr_Medio', 'mean'), Es_MPa=('Es_Medio', 'mean'), M_MPa=('M_Medio', 'mean')).reset_index()
    capas['Espesor_m'] = capas['Hasta_m'] - capas['Desde_m']
    capas_final = capas[['Desde_m', 'Hasta_m', 'Espesor_m', 'SBT_Predominante', 'qt_MPa', 'Su_kPa', 'Phi_deg', 'Dr_per', 'Es_MPa', 'M_MPa']]
    
    with tab_cap: st.dataframe(capas_final, hide_index=True)
    with tab_det: st.dataframe(df_v[['Depth_m', 'SBT_Name', 'Ic', 'qt_MPa', 'Fr_percent', 'Qt', 'Su_kPa', 'Phi_deg', 'Dr_percent', 'M_MPa', 'Es_MPa', 'OCR', 'N60']], hide_index=True)
    
    with tab_f:
        st.subheader("📚 Metodología y Formulación Geotécnica")
        st.write("Interpretación avanzada basada en corrección por área neta, tensiones efectivas y clasificación poligonal (Robertson 1990, 2009 y 2010).")
        st.divider()

        col_A, col_B = st.columns(2)
        with col_A:
            st.markdown("### 1. Resistencia Corregida y Tensión Efectiva")
            st.latex(r"q_t = q_c + u_2 \cdot (1 - a)")
            st.latex(r"\sigma'_{v0} = \sum (\gamma \cdot \Delta z) - u_0")
            
            st.markdown("### 2. Parámetros Normalizados ($Q_t$ y $F_r$)")
            st.latex(r"Q_t = \frac{q_t - \sigma_{v0}}{\sigma'_{v0}}")
            st.latex(r"F_r (\%) = \left[ \frac{f_s}{q_t - \sigma_{v0}} \right] \times 100")

            st.markdown("### 3. Índice de Comportamiento del Suelo ($I_c$)")
            st.latex(r"I_c = \sqrt{(3.47 - \log_{10}Q_t)^2 + (\log_{10}F_r + 1.22)^2}")

            st.markdown("### 4. Densidad Relativa ($D_r$)")
            st.latex(r"D_r (\%) = \sqrt{\frac{Q_t}{350}} \times 100")

        with col_B:
            st.markdown("### 5. Resistencia al Corte sin Drenaje ($S_u$)")
            st.latex(r"S_u = \frac{q_t - \sigma_{v0}}{N_{kt}}")

            st.markdown("### 6. Módulos de Deformación ($M$ y $E_s$)")
            st.latex(r"M = \alpha_M \cdot (q_t - \sigma_{v0}) \quad \text{donde } \alpha_M = 0.0188 \cdot 10^{(0.55 I_c + 1.68)}")
            st.latex(r"E_s = \alpha_E \cdot (q_t - \sigma_{v0}) \quad \text{donde } \alpha_E = 0.015 \cdot 10^{(0.55 I_c + 1.68)}")

            st.markdown("### 7. Ratio de Sobreconsolidación (OCR)")
            st.latex(r"OCR = 0.33 \cdot Q_t")

            st.markdown("### 8. Equivalencia SPT ($N_{60}$)")
            st.latex(r"N_{60} = \frac{q_c/p_a}{10^{(1.1268 - 0.2817 \cdot I_c)}}")

    # --- EXPORTACIÓN ---
    st.sidebar.divider()
    if st.sidebar.button("🚀 GENERAR INFORME", use_container_width=True):
        progress = st.sidebar.progress(0)
        status = st.sidebar.empty()
        
        status.text("Creando Informe Word...")
        doc = Document()
        doc.add_heading(f'Informe Geotécnico Avanzado: {header_data.get("Location", "CPTU")}', 0)
        
        for i, t in enumerate(['basicos', 'resistencia', 'deformacion', 'calidad']):
            fig = generar_figura_perfil(t)
            buf = io.BytesIO(); fig.savefig(buf, format='png', dpi=150)
            doc.add_heading(f'Perfil de {t.capitalize()}', level=1)
            doc.add_picture(buf, width=Inches(6.5))
            progress.progress(15 + i*15)
            
        status.text("Renderizando Gráfico de Distribución...")
        doc.add_heading('Distribución de Tipos de Suelo (SBT 9 Zonas)', level=1)
        buf_dist = io.BytesIO(); fig_dist.savefig(buf_dist, format='png', dpi=150)
        doc.add_picture(buf_dist, width=Inches(6.0))
        progress.progress(80)

        status.text("Creando Excel Multihoja...")
        excel_buf = io.BytesIO()
        with pd.ExcelWriter(excel_buf, engine='xlsxwriter') as writer:
            pd.DataFrame(list(header_data.items())).to_excel(writer, sheet_name='Metadatos')
            capas_final.to_excel(writer, sheet_name='Resumen_Capas', index=False)
            df_v[['Depth_m', 'SBT_Name', 'Ic', 'qt_MPa', 'Fr_percent', 'Qt', 'Su_kPa', 'Phi_deg', 'Dr_percent', 'M_MPa', 'Es_MPa', 'OCR', 'N60']].to_excel(writer, sheet_name='Datos_Completos', index=False)

        zip_buf = io.BytesIO()
        with zipfile.ZipFile(zip_buf, 'w') as zf:
            w_buf = io.BytesIO(); doc.save(w_buf)
            zf.writestr("Informe_Word_CPTU.docx", w_buf.getvalue())
            zf.writestr("Libro_Calculo_CPTU.xlsx", excel_buf.getvalue())
        
        progress.progress(100)
        status.success("¡Informe Creado!")
        st.sidebar.download_button("📥 DESCARGAR INFORME (.ZIP)", zip_buf.getvalue(), "Resultados_CPTU.zip", "application/zip", use_container_width=True)