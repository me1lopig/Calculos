import streamlit as st
import pandas as pd
import numpy as np
import io
from docx import Document
from docx.shared import Inches
from apsg import fol, lin, folset, StereoNet
from sklearn.cluster import KMeans

# --- CONFIGURACIÓN DE PÁGINA ---
st.set_page_config(page_title="Geomecánica PRO", layout="wide")
st.title("Análisis Estereográfico y Cinemático de Rocas")

# =====================================================================
# 🧠 MOTOR MATEMÁTICO PURO (Independiente de APSG)
# =====================================================================
def calcular_plano_medio(dd_array, dip_array):
    if len(dd_array) == 0:
        return fol(0, 0), 0.0, 0.0
        
    trend_polo = np.radians((dd_array + 180) % 360)
    plunge_polo = np.radians(90 - dip_array)
    
    x = np.sin(trend_polo) * np.cos(plunge_polo)
    y = np.cos(trend_polo) * np.cos(plunge_polo)
    z = np.sin(plunge_polo)
    
    V = np.column_stack((x, y, z))
    T = np.dot(V.T, V)
    
    eigenvalues, eigenvectors = np.linalg.eigh(T)
    polo_medio = eigenvectors[:, np.argmax(eigenvalues)]
    
    if polo_medio[2] < 0:
        polo_medio = -polo_medio
        
    plunge_pm = np.arcsin(polo_medio[2])
    trend_pm = np.arctan2(polo_medio[0], polo_medio[1])
    
    trend_pm_deg = np.degrees(trend_pm) % 360
    plunge_pm_deg = np.degrees(plunge_pm)
    
    dip_m = 90 - plunge_pm_deg
    dd_m = (trend_pm_deg - 180) % 360
    
    return fol(dd_m, dip_m), dd_m, dip_m

def calcular_interseccion_cuna(dd1, dip1, dd2, dip2):
    t1 = np.radians((dd1 + 180) % 360)
    p1 = np.radians(90 - dip1)
    v1 = np.array([np.sin(t1)*np.cos(p1), np.cos(t1)*np.cos(p1), np.sin(p1)])
    
    t2 = np.radians((dd2 + 180) % 360)
    p2 = np.radians(90 - dip2)
    v2 = np.array([np.sin(t2)*np.cos(p2), np.cos(t2)*np.cos(p2), np.sin(p2)])
    
    inter = np.cross(v1, v2)
    norma = np.linalg.norm(inter)
    if norma == 0:
        return lin(0,0), 0.0, 0.0
        
    inter = inter / norma
    if inter[2] < 0:
        inter = -inter
        
    plunge_c = np.degrees(np.arcsin(inter[2]))
    trend_c = np.degrees(np.arctan2(inter[0], inter[1])) % 360
    
    return lin(trend_c, plunge_c), trend_c, plunge_c
# =====================================================================

# --- BARRA LATERAL: PARÁMETROS GLOBALES ---
st.sidebar.header("1. Parámetros del Talud")
talud_dd = st.sidebar.number_input("Dirección de Buzamiento (Dip Dir)", 0, 360, 180)
talud_dip = st.sidebar.number_input("Buzamiento (Dip)", 0, 90, 60)
friccion = st.sidebar.slider("Ángulo de Fricción Interna (°)", 10, 80, 30)

talud = fol(talud_dd, talud_dip)
eje_vertical = lin(0, 90)

st.sidebar.header("2. Entrada de Datos")
archivo = st.sidebar.file_uploader("Sube tu archivo CSV o Excel ('Dip_Dir', 'Dip')", type=['csv', 'xlsx'])

if archivo is None:
    st.sidebar.info("Usando datos de prueba (2 Familias).")
    datos_prueba = {'Dip_Dir': [175, 182, 178, 160, 180, 172, 260, 265, 255, 270, 262, 258], 
                    'Dip': [45, 48, 42, 50, 46, 44, 80, 85, 82, 78, 81, 79]}
    df = pd.DataFrame(datos_prueba)
else:
    if archivo.name.endswith('csv'):
        df = pd.read_csv(archivo)
    else:
        df = pd.read_excel(archivo)

st.sidebar.header("3. Clustering (K-Means)")
num_familias = st.sidebar.number_input("Número de Familias a detectar", min_value=1, max_value=5, value=2)

# --- PROCESAMIENTO PRINCIPAL ---
if not df.empty and 'Dip_Dir' in df.columns and 'Dip' in df.columns:
    
    lista_foliaciones = [fol(d, dp) for d, dp in zip(df['Dip_Dir'], df['Dip'])]
    macizo = folset(lista_foliaciones, name="Macizo")
    
    trend_k = np.radians((df['Dip_Dir'] + 180) % 360)
    plunge_k = np.radians(90 - df['Dip'])
    x_k = np.sin(trend_k) * np.cos(plunge_k)
    y_k = np.cos(trend_k) * np.cos(plunge_k)
    z_k = np.sin(plunge_k)
    
    vectores = np.column_stack((x_k, y_k, z_k))
    kmeans = KMeans(n_clusters=num_familias, random_state=42, n_init=10).fit(vectores)
    df['Cluster'] = kmeans.labels_
    
    familias = []
    planos_medios_data = [] 
    colores = ['blue', 'orange', 'purple', 'cyan', 'magenta']
    
    for i in range(num_familias):
        datos_cluster = df[df['Cluster'] == i]
        grupo_fam = folset([fol(d, dp) for d, dp in zip(datos_cluster['Dip_Dir'], datos_cluster['Dip'])])
        familias.append(grupo_fam)
        
        f_obj, dd_m, dip_m = calcular_plano_medio(datos_cluster['Dip_Dir'].values, datos_cluster['Dip'].values)
        planos_medios_data.append((f_obj, dd_m, dip_m))

    # --- PESTAÑAS ---
    tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs([
        "📊 Datos", "🎯 Familias (Clustering)", "📉 Rotura Plana", "🔻 Rotura Cuña", "🔄 Vuelco", "📄 Informe Word"
    ])
    
    # --- TAB 1: VISOR DE DATOS ---
    with tab1:
        st.subheader("Datos Estructurales de Campo")
        col1, col2 = st.columns([1, 2])
        
        with col1:
            st.markdown(f"**Total de mediciones analizadas:** `{len(df)}` datos")
            st.dataframe(df.drop(columns=['Cluster'], errors='ignore'), use_container_width=True, height=450)
            
        with col2:
            st.markdown("**Visualización Estereográfica**")
            tipo_viz = st.radio("Configuración de la vista:", ["📍 Solo Polos", "🌐 Solo Planos", "✨ Ambos"], horizontal=True)
            
            s_datos = StereoNet()
            
            if tipo_viz in ["📍 Solo Polos", "✨ Ambos"]:
                s_datos.pole(macizo, color='black', markersize=4, label="Polos medidos")
                
            if tipo_viz in ["🌐 Solo Planos", "✨ Ambos"]:
                for f_obj in macizo:
                    s_datos.great_circle(f_obj, color='gray', linewidth=0.5, alpha=0.3)
            
            # --- SOLUCIÓN MODO DIOS: Exportación directa a imagen ---
            buf_datos = io.BytesIO()
            s_datos.savefig(buf_datos, format="png", dpi=150, bbox_inches="tight")
            st.image(buf_datos, use_container_width=True)
            
            st.download_button(label="📥 Descargar Falsilla (PNG Alta Resolución)", data=buf_datos.getvalue(), file_name="Falsilla_Datos.png", mime="image/png", use_container_width=True)

    with tab2:
        st.subheader(f"Identificación Automática: {num_familias} Familias")
        col1, col2 = st.columns(2)
        with col1:
            s_cluster = StereoNet()
            s_cluster.contour(macizo, cmap='YlOrRd', sigma=2)
            for i, fam in enumerate(familias):
                s_cluster.pole(fam, color=colores[i], markersize=4, label=f"Set {i+1}")
                
            buf_cluster = io.BytesIO()
            s_cluster.savefig(buf_cluster, format="png", dpi=150, bbox_inches="tight")
            st.image(buf_cluster, use_container_width=True)
            
        with col2:
            s_mean = StereoNet()
            for i, data in enumerate(planos_medios_data):
                f_obj, dd_m, dip_m = data
                st.metric(f"Familia {i+1} (Dip Dir / Dip)", f"{dd_m:.1f}° / {dip_m:.1f}°")
                s_mean.great_circle(f_obj, color=colores[i], linewidth=2, label=f"Fam {i+1}")
                s_mean.pole(f_obj, color=colores[i], marker='s', markersize=8)
                
            buf_mean = io.BytesIO()
            s_mean.savefig(buf_mean, format="png", dpi=150, bbox_inches="tight")
            st.image(buf_mean, use_container_width=True)

    with tab3:
        st.subheader("Análisis de Rotura Plana")
        polos_criticos_dd = []
        polos_criticos_dip = []
        
        for index, row in df.iterrows():
            p_dd = row['Dip_Dir']
            p_dip = row['Dip']
            dif_azimut = abs(p_dd - talud_dd)
            if dif_azimut > 180: dif_azimut = 360 - dif_azimut
            if dif_azimut <= 20 and friccion < p_dip < talud_dip:
                polos_criticos_dd.append(p_dd)
                polos_criticos_dip.append(p_dip)
                
        prob_plana = (len(polos_criticos_dd) / len(df)) * 100
        
        col1, col2 = st.columns([2, 1])
        with col1:
            s_plana = StereoNet()
            s_plana.great_circle(talud, color='red', linewidth=2, label="Talud")
            s_plana.cone(eje_vertical, 90 - friccion, color='green', fill=False)
            s_plana.pole(macizo, color='black', markersize=3)
            if polos_criticos_dd:
                g_criticos = folset([fol(d, dp) for d, dp in zip(polos_criticos_dd, polos_criticos_dip)])
                s_plana.pole(g_criticos, color='red', markersize=6)
                
            buf_plana = io.BytesIO()
            s_plana.savefig(buf_plana, format="png", dpi=150, bbox_inches="tight")
            st.image(buf_plana, use_container_width=True)
            
        with col2:
            st.error(f"**Probabilidad:** {prob_plana:.1f}%")
            st.write(f"Polos críticos: {len(polos_criticos_dd)}")

    with tab4:
        st.subheader("Análisis de Rotura en Cuña")
        if num_familias >= 2:
            f1_obj, f1_dd, f1_dip = planos_medios_data[0]
            f2_obj, f2_dd, f2_dip = planos_medios_data[1]
            
            linea_cuna_obj, cuna_trend, cuna_plunge = calcular_interseccion_cuna(f1_dd, f1_dip, f2_dd, f2_dip)
            riesgo_cuna = friccion < cuna_plunge < talud_dip
            
            col1, col2 = st.columns([2, 1])
            with col1:
                s_cuna = StereoNet()
                s_cuna.great_circle(talud, color='red', linewidth=2, label="Talud")
                s_cuna.great_circle(f1_obj, color=colores[0], label="Fam 1")
                s_cuna.great_circle(f2_obj, color=colores[1], label="Fam 2")
                s_cuna.cone(eje_vertical, 90 - friccion, color='green', fill=False)
                s_cuna.line(linea_cuna_obj, color='red' if riesgo_cuna else 'black', marker='^', markersize=10)
                
                buf_cuna = io.BytesIO()
                s_cuna.savefig(buf_cuna, format="png", dpi=150, bbox_inches="tight")
                st.image(buf_cuna, use_container_width=True)
                
            with col2:
                st.write(f"**Inmersión de la Cuña:** {cuna_plunge:.1f}°")
                st.write(f"**Dirección de la Cuña:** {cuna_trend:.1f}°")
                if riesgo_cuna:
                    st.error("⚠️ **Peligro:** La cuña aflora y supera la fricción.")
                else:
                    st.success("✅ Cuña estable.")
        else:
            st.warning("⚠️ Necesitas al menos 2 familias para analizar cuñas.")

    with tab5:
        st.subheader("Análisis de Vuelco Flexural")
        polos_vuelco_dd = []
        polos_vuelco_dip = []
        
        for index, row in df.iterrows():
            p_dd = row['Dip_Dir']
            p_dip = row['Dip']
            dif_azimut = abs(p_dd - talud_dd)
            if dif_azimut > 180: dif_azimut = 360 - dif_azimut
            if dif_azimut >= 160 and (p_dip + friccion) >= talud_dip:
                polos_vuelco_dd.append(p_dd)
                polos_vuelco_dip.append(p_dip)
                
        prob_vuelco = (len(polos_vuelco_dd) / len(df)) * 100
        
        col1, col2 = st.columns([2, 1])
        with col1:
            s_vuelco = StereoNet()
            s_vuelco.great_circle(talud, color='red', linewidth=2, label="Talud")
            s_vuelco.pole(macizo, color='black', markersize=3)
            if polos_vuelco_dd:
                g_vuelco = folset([fol(d, dp) for d, dp in zip(polos_vuelco_dd, polos_vuelco_dip)])
                s_vuelco.pole(g_vuelco, color='purple', markersize=6)
                
            buf_vuelco = io.BytesIO()
            s_vuelco.savefig(buf_vuelco, format="png", dpi=150, bbox_inches="tight")
            st.image(buf_vuelco, use_container_width=True)
            
        with col2:
            st.error(f"**Probabilidad:** {prob_vuelco:.1f}%")

    with tab6:
        st.subheader("Generar Entregable Profesional")
        st.write("Exporta todo el análisis y los gráficos a un documento de Word.")
        
        def generar_word():
            doc = Document()
            doc.add_heading('Estudio Geomecánico y Cinemático de Taludes', 0)
            
            doc.add_heading('1. Parámetros de Diseño', level=1)
            doc.add_paragraph(f"• Talud (Dip Dir / Dip): {talud_dd}° / {talud_dip}°\n"
                              f"• Fricción Interna: {friccion}°\n"
                              f"• Datos analizados: {len(df)}")
            
            doc.add_heading('2. Familias Identificadas (Clustering)', level=1)
            for i, data in enumerate(planos_medios_data):
                _, dd_m, dip_m = data
                doc.add_paragraph(f"• Familia {i+1}: {dd_m:.1f}° / {dip_m:.1f}°")
            
            doc.add_heading('3. Resultados Cinemáticos', level=1)
            doc.add_paragraph(f"• Rotura Plana (Probabilidad): {prob_plana:.1f}%\n"
                              f"• Vuelco Flexural (Probabilidad): {prob_vuelco:.1f}%")
            if num_familias >= 2:
                doc.add_paragraph(f"• Rotura en Cuña (Fam 1 vs Fam 2): Inmersión a {cuna_plunge:.1f}°")
            
            doc.add_heading('4. Anexos Estereográficos', level=1)
            
            def agregar_grafico(stereo_obj, titulo):
                doc.add_paragraph(titulo, style='Heading 2')
                buf = io.BytesIO()
                stereo_obj.savefig(buf, format='png', bbox_inches='tight', dpi=150)
                buf.seek(0)
                doc.add_picture(buf, width=Inches(5.5))
            
            agregar_grafico(s_cluster, "Densidad y Clustering de Familias")
            agregar_grafico(s_plana, "Análisis de Rotura Plana")
            if num_familias >= 2:
                agregar_grafico(s_cuna, "Análisis de Rotura en Cuña")
            agregar_grafico(s_vuelco, "Análisis de Vuelco")
            
            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            return doc_buffer

        st.download_button(
            label="📄 Descargar Informe Completo (.docx)",
            data=generar_word(),
            file_name="Informe_Geomecanico.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

else:
    st.error("Error cargando datos. Asegúrate de que las columnas se llamen 'Dip_Dir' y 'Dip'.")