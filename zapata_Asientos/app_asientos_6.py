import streamlit as st
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import io
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# ══════════════════════════════════════════════════════════════════════════
# CONSTANTES
# ══════════════════════════════════════════════════════════════════════════
GAMMA_AGUA = 9.81  # kN/m³

# ══════════════════════════════════════════════════════════════════════════
# TENSIONES DE HOLL — BAJO EL CENTRO (compartido por ambos métodos)
# ══════════════════════════════════════════════════════════════════════════
def holl_esquina(p, B, L, z):
    """Tensiones bajo la ESQUINA de una carga rectangular BxL (Solución de Holl)."""
    if z <= 1e-6:
        return p, p / 2.0, p / 2.0
    R1 = np.sqrt(L**2 + z**2)
    R2 = np.sqrt(B**2 + z**2)
    R3 = np.sqrt(L**2 + B**2 + z**2)
    arc = np.arctan((B * L) / (z * R3))
    sz = (p / (2*np.pi)) * (arc + B*L*(1/R1**2 + 1/R2**2)*(z/R3))
    sx = (p / (2*np.pi)) * (arc - (B*L*z)/(R1**2*R3))
    sy = (p / (2*np.pi)) * (arc - (B*L*z)/(R2**2*R3))
    return sz, sx, sy

def holl_centro(p, B, L, z):
    """Tensiones bajo el CENTRO: superposición ×4 de cuadrantes B/2 × L/2."""
    sz, sx, sy = holl_esquina(p, B/2.0, L/2.0, z)
    return 4*sz, 4*sx, 4*sy

# ══════════════════════════════════════════════════════════════════════════
# MÉTODO 1 — STEINBRENNER (φ1, φ2, s(z) analítico)
# ══════════════════════════════════════════════════════════════════════════
def phi1(m, n):
    if m == 0:
        t1 = np.log(np.sqrt(1+n**2)+n)
        t2 = n*np.log((np.sqrt(1+n**2)+1)/n)
    else:
        t1 = np.log((np.sqrt(1+m**2+n**2)+n)/np.sqrt(1+m**2))
        t2 = n*np.log((np.sqrt(1+m**2+n**2)+1)/np.sqrt(n**2+m**2))
    return (1/np.pi)*(t1+t2)

def phi2(m, n):
    if m == 0: return 0.0
    return (m/np.pi)*np.arctan(n/(m*np.sqrt(1+m**2+n**2)))

def s_z(p, B, E, nu, z, L):
    """Asiento teórico acumulado desde superficie hasta z (Steinbrenner)."""
    n = L/B
    m = (2*z)/B
    corchete = (1-nu**2)*phi1(m,n) - (1-nu-2*nu**2)*phi2(m,n)
    return (p*B/E)*corchete

def calcular_steinbrenner(p, B, L, df, z_max):
    """
    Método 1 — Steinbrenner BAJO EL CENTRO de la cimentación.
    s_z() da el asiento bajo la ESQUINA de B\xd7L. Para el CENTRO:
        s_centro(z) = 4 × s_z(p, B/2, E, nu, z, L/2)
    Igual que la superposición usada por holl_centro().
    """
    total = 0.0
    resultados = []
    z_actual = 0.0
    n_factor = L / B   # = (L/2)/(B/2): la esbeltez del cuadrante es idéntica

    for _, row in df.iterrows():
        if z_actual >= z_max:
            break
        h_i   = float(row["Espesor (m)"])
        E_i   = float(row["E (kPa)"])
        nu_i  = float(row["nu"])
        nombre= str(row["Descripción"])

        z_techo = z_actual
        z_base  = min(z_actual + h_i, z_max)

        # m relativo al cuadrante B/2 (m_cuad = 2z/(B/2) = 4z/B)
        m_t = (2*z_techo) / (B/2)
        m_b = (2*z_base)  / (B/2)
        # Asiento CENTRO = 4 × esquina del cuadrante B/2 × L/2
        s_t = 4 * s_z(p, B/2, E_i, nu_i, z_techo, L/2)
        s_b = 4 * s_z(p, B/2, E_i, nu_i, z_base,  L/2)
        ds  = s_t - s_b
        total += ds

        resultados.append({
            "Capa":               nombre,
            "z Techo [m]":        round(z_techo, 3),
            "z Base [m]":         round(z_base,  3),
            "m_techo":            round(m_t, 4),
            "φ1_techo":           round(phi1(m_t, n_factor), 4),
            "φ2_techo":           round(phi2(m_t, n_factor), 4),
            "s_techo [mm]":       round(s_t*1000, 3),
            "m_base":             round(m_b, 4),
            "φ1_base":            round(phi1(m_b, n_factor), 4),
            "φ2_base":            round(phi2(m_b, n_factor), 4),
            "s_base [mm]":        round(s_b*1000, 3),
            "Δs [mm]":            round(ds*1000, 3),
        })
        z_actual = z_base

    return total, pd.DataFrame(resultados)

# ══════════════════════════════════════════════════════════════════════════
# MÉTODO 2 — (integración directa de deformaciones unitarias)
# ══════════════════════════════════════════════════════════════════════════
def calcular_ec68(p, B, L, df, z_max, dz_sub=0.25):
    """
    Método 2 — Ec. 68:  s = Σ [ h/E · (Δσz - ν(Δσx+Δσy)) ]_i
    Cada estrato se subdivide en subcapas de dz_sub metros para mejorar
    la precisión de la integración numérica (regla del punto medio).
    La tabla de resultados muestra valores AGREGADOS por capa original.
    """
    total = 0.0
    resultados = []
    z_actual = 0.0

    for _, row in df.iterrows():
        if z_actual >= z_max:
            break
        h_i    = float(row["Espesor (m)"])
        E_i    = float(row["E (kPa)"])
        nu_i   = float(row["nu"])
        nombre = str(row["Descripción"])

        z_techo = z_actual
        z_base  = min(z_actual + h_i, z_max)
        h_ef    = z_base - z_techo

        # Número de subcapas (al menos 1)
        n_sub  = max(1, int(np.ceil(h_ef / dz_sub)))
        dz     = h_ef / n_sub

        # Integración por subcapas (punto medio de cada una)
        ds_capa  = 0.0
        sz_medio = 0.0
        sx_medio = 0.0
        sy_medio = 0.0
        ez_medio = 0.0

        for k in range(n_sub):
            z_sub_t = z_techo + k * dz
            z_mid   = z_sub_t + dz / 2.0
            dsz, dsx, dsy = holl_centro(p, B, L, z_mid)
            dep_z  = (dsz - nu_i*(dsx+dsy)) / E_i
            ds_sub = dep_z * dz
            ds_capa  += ds_sub
            sz_medio += dsz
            sx_medio += dsx
            sy_medio += dsy
            ez_medio += dep_z

        # Promedios representativos de la capa (solo para mostrar)
        sz_medio /= n_sub
        sx_medio /= n_sub
        sy_medio /= n_sub
        ez_medio /= n_sub

        total += ds_capa

        resultados.append({
            "Capa":          nombre,
            "z Techo [m]":   round(z_techo,  3),
            "z Base [m]":    round(z_base,   3),
            "h_ef [m]":      round(h_ef,     3),
            "Sub-capas":      n_sub,
            "Δσz med [kPa]": round(sz_medio, 3),
            "Δσx med [kPa]": round(sx_medio, 3),
            "Δσy med [kPa]": round(sy_medio, 3),
            "Δεz med [-]":   round(ez_medio, 6),
            "Δs [mm]":        round(ds_capa*1000, 3),
        })
        z_actual = z_base

    return total, pd.DataFrame(resultados)

# ══════════════════════════════════════════════════════════════════════════
# TENSIÓN EFECTIVA Y ZONA DE INFLUENCIA
# ══════════════════════════════════════════════════════════════════════════
def sigma_v0(z, df, NF):
    sv = 0.0; z_act = 0.0
    for _, row in df.iterrows():
        h  = float(row["Espesor (m)"])
        g  = float(row["Peso Esp. (kN/m³)"])
        gs = float(row["Peso Esp. Sat (kN/m³)"])
        zt = z_act; zb = z_act + h
        if z <= zt: break
        ze = min(z, zb)
        z_sec_b = min(ze, NF)
        if z_sec_b > zt: sv += g*(z_sec_b-zt)
        z_sat_t = max(zt, NF)
        if ze > z_sat_t: sv += (gs-GAMMA_AGUA)*(ze-z_sat_t)
        z_act = zb
    return sv

def z_influencia_ec7(p, B, L, df, NF):
    et = float(pd.to_numeric(df["Espesor (m)"]).sum())
    z = 0.05
    while z <= et:
        dsz, _, _ = holl_centro(p, B, L, z)
        sv = sigma_v0(z, df, NF)
        if sv > 0 and dsz <= 0.20*sv:
            return z
        z += 0.05
    return et

# ══════════════════════════════════════════════════════════════════════════
# INFORME WORD ESTÉTICO
# ══════════════════════════════════════════════════════════════════════════
def _fig_bytes(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=200, bbox_inches='tight')
    buf.seek(0)
    return buf

def _add_styled_table(doc, df, title):
    doc.add_heading(title, level=2)
    # Convertir dataframe limpiando formatos
    df = df.astype(str)
    table = doc.add_table(rows=1+len(df), cols=len(df.columns))
    table.style = 'Medium List 1 Accent 1' # Estilo nativo de Word muy limpio
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Cabeceras
    hdr_cells = table.rows[0].cells
    for i, column in enumerate(df.columns):
        hdr_cells[i].text = column
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255)
    
    # Filas de datos
    for i, row in enumerate(df.itertuples(index=False)):
        row_cells = table.rows[i+1].cells
        for j, value in enumerate(row):
            row_cells[j].text = str(value)
            for paragraph in row_cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
                    
    # Añadir espacio post-tabla
    doc.add_paragraph()

def generar_word(B, L, p, NF, z_max, zi, df_terreno,
                 df_st, tot_st, df_ec, tot_ec,
                 fig_bulbo_bytes):
    fecha = datetime.now().strftime("%d de %B de %Y — %H:%M")
    doc = Document()
    
    # Configurar márgenes elegantes
    for sec in doc.sections:
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(2.5)

    # Estilos globales
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Portada principal
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run('CÁLCULO DE ASIENTOS')
    r_title.bold = True
    r_title.font.size = Pt(18)
    r_title.font.color.rgb = RGBColor(31, 73, 125) # Azul oscuro profesional
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run('Comparativa de Formulaciones: Steinbrenner vs Guía ecuación elástica \n')
    r_sub.italic = True
    r_sub.font.size = Pt(12)
    r_sub.font.color.rgb = RGBColor(128, 128, 128)
    
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_date = p_date.add_run(f'Fecha de emisión: {fecha}')
    r_date.font.size = Pt(10)

    doc.add_paragraph() # Espacio
    
    # ── 1. Datos de Entrada ──────────────────────────────────────────────
    doc.add_heading('1. Parámetros de la Cimentación', level=1)
    
    # Tabla sin bordes feos para datos (Grid = feo, Light = limpio)
    table_params = doc.add_table(rows=7, cols=2)
    table_params.style = 'Light List Accent 1'
    
    data = [
        ('Dimensiones en planta (B × L)', f'{B:.2f} m  ×  {L:.2f} m'),
        ('Esbeltez geométrica (n = L/B)', f'{L/B:.3f}'),
        ('Presión de trabajo (p)', f'{p:.1f} kPa'),
        ('Nivel freático (NF)', f'{NF:.1f} m desde cimentación'),
        ('Profundidad máxima de corte (z_max)', f'{z_max:.1f} m'),
        ('Profundidad bulbo criterio EC7 (z_i)', f'{zi:.2f} m'),
        ('Tensiones', 'Calculadas bajo el CENTRO (Superposición 4x)')
    ]
    
    for i, (key, val) in enumerate(data):
        row = table_params.rows[i].cells
        row[0].text = key
        row[1].text = val
        row[0].paragraphs[0].runs[0].font.bold = True
        row[0].paragraphs[0].runs[0].font.size = Pt(10)
        row[1].paragraphs[0].runs[0].font.size = Pt(10)
        
    doc.add_paragraph()
    
    _add_styled_table(doc, df_terreno, '1.2 Estratigrafía del Perfil')
    doc.add_page_break()

    doc.add_heading('2. Comparativa de Resultados', level=1)
    p2 = doc.add_paragraph()
    r1 = p2.add_run(f'Steinbrenner: {tot_st*1000:.3f} mm    ')
    r1.font.bold=True; r1.font.size=Pt(11); r1.font.color.rgb=RGBColor(26,58,92)
    r2 = p2.add_run(f'Ec. 68: {tot_ec*1000:.3f} mm')
    r2.font.bold=True; r2.font.size=Pt(11); r2.font.color.rgb=RGBColor(0,100,0)
    doc.add_paragraph()

    # 3. Detalles
    _add_styled_table(doc, df_st, '3. Detalle Steinbrenner (por capa)')
    doc.add_page_break()
    _add_styled_table(doc, df_ec, '4. Detalle Ecuación Elástica (por capa)')
    doc.add_page_break()

    # 5. Bulbo
    h5 = doc.add_heading('5. Bulbo de Presiones', level=1)
    doc.add_picture(fig_bulbo_bytes, width=Cm(12))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    nota = doc.add_paragraph(f'B={B} m · L={L} m · p={p} kPa · NF={NF} m — {fecha}')
    nota.runs[0].font.size=Pt(8); nota.runs[0].font.color.rgb=RGBColor(128,128,128)
    nota.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf

# ══════════════════════════════════════════════════════════════════════════
# SESSION STATE
# ══════════════════════════════════════════════════════════════════════════
def reset_calculo():
    st.session_state.calculo_realizado = False

if 'calculo_realizado' not in st.session_state:
    st.session_state.calculo_realizado = False

if 'df_terreno' not in st.session_state:
    st.session_state.df_terreno = pd.DataFrame({
        "Descripción":           ["Relleno",  "Arcilla",  "Grava"],
        "Espesor (m)":           [1.5,         3.0,        5.0],
        "E (kPa)":               [10000.0,     15000.0,     40000.0],
        "nu":                    [0.30,         0.45,       0.25],
        "Peso Esp. (kN/m³)":     [18.0,         19.0,       21.0],
        "Peso Esp. Sat (kN/m³)": [20.0,         20.0,       22.0],
    })

# ══════════════════════════════════════════════════════════════════════════
# CONFIGURACIÓN
# ══════════════════════════════════════════════════════════════════════════
st.set_page_config(
    page_title="Asientos Comparativa de Métodos",
    layout="wide", page_icon="🏗️"
)

st.sidebar.title("Navegación")
modo = st.sidebar.radio("Vista:", [
    "🧮 Panel de Cálculo",
    "📋 Modelo Steinbrenner",
    "📋 Modelo Elástico",
    "📉 Bulbo de Presiones",
    "📖 Fundamento Teórico",
])

st.sidebar.markdown("---")
st.sidebar.header("📥 Datos de Entrada")

B  = st.sidebar.number_input("Ancho (B) [m]",             min_value=0.1, value=2.0,   step=0.1,  on_change=reset_calculo)
L  = st.sidebar.number_input("Longitud (L) [m]",          min_value=0.1, value=3.0,   step=0.1,  on_change=reset_calculo)
p  = st.sidebar.number_input("Presión neta (p) [kPa]",    min_value=1.0, value=150.0, step=10.0, on_change=reset_calculo)
NF = st.sidebar.number_input("Nivel Freático [m]",         min_value=0.0, value=100.0, step=0.5,  on_change=reset_calculo)

if L < B:
    B, L = L, B
    st.sidebar.warning("⚠️ L<B: valores intercambiados.")
st.sidebar.info(f"**n = L/B:** {L/B:.2f}")

# Profundidad de influencia
espesor_total = max(float(pd.to_numeric(st.session_state.df_terreno["Espesor (m)"]).sum()), 0.1)
zi = z_influencia_ec7(p, B, L, st.session_state.df_terreno, NF)

st.sidebar.markdown("---")
st.sidebar.subheader("📐 Profundidad de Cálculo")
if zi >= espesor_total - 0.05:
    st.sidebar.warning(f"⚠️ Bulbo supera estratigrafía (>{espesor_total:.1f} m).")
else:
    st.sidebar.success(f"💡 **z_i EC7 = {zi:.2f} m**")

z_max_user = st.sidebar.number_input(
    "Profundidad de corte (z_max) [m]",
    min_value=0.1, max_value=espesor_total,
    value=float(min(round(zi, 1), espesor_total)),
    step=0.1, on_change=reset_calculo
)

# Precisión de ecuación elástica
st.sidebar.markdown("---")
st.sidebar.subheader("🔧 Precisión Ecuación Elástica")
dz_sub = st.sidebar.select_slider(
    "Tamaño de subcapa (dz) [m]",
    options=[2.0, 1.0, 0.5, 0.25, 0.10, 0.05],
    value=0.10,
    on_change=reset_calculo,
    help="Subdivisión de cada estrato para modelo elástico. "
         "Menos dz = más precisión (converge a Steinbrenner)."
)
st.sidebar.caption(f"Subcapas estimadas: ~{int(np.ceil(espesor_total / dz_sub))} en total")

# ══════════════════════════════════════════════════════════════════════════
# BOTÓN CALCULAR
# ══════════════════════════════════════════════════════════════════════════
st.sidebar.markdown("---")
if st.sidebar.button("🚀 Calcular", type="primary", use_container_width=True):
    tot_st, df_st = calcular_steinbrenner(p, B, L, st.session_state.df_terreno, z_max_user)
    tot_ec, df_ec = calcular_ec68(        p, B, L, st.session_state.df_terreno, z_max_user, dz_sub)

    st.session_state.tot_st  = tot_st
    st.session_state.df_st   = df_st
    st.session_state.tot_ec  = tot_ec
    st.session_state.df_ec   = df_ec
    st.session_state.dz_used = dz_sub
    st.session_state.calculo_realizado = True

# ══════════════════════════════════════════════════════════════════════════
# BOTÓN INFORME WORD
# ══════════════════════════════════════════════════════════════════════════
st.sidebar.markdown("---")
if st.session_state.calculo_realizado:
    df_st = st.session_state.df_st
    df_ec = st.session_state.df_ec
    tot_st= st.session_state.tot_st
    tot_ec= st.session_state.tot_ec

    # ── Figura bulbo ──────────────────────────────────────────────────────
    z_vals = np.linspace(0.05, espesor_total, 200)
    sz_v,sx_v,sy_v,sv0_v = [],[],[],[]
    for z in z_vals:
        sz,sx,sy = holl_centro(p, B, L, z)
        sz_v.append(sz); sx_v.append(sx); sy_v.append(sy)
        sv0_v.append(sigma_v0(z, st.session_state.df_terreno, NF)*0.20)
    fig_b, ax_b = plt.subplots(figsize=(5, 7))
    ax_b.plot(sz_v, z_vals, label=r"$\Delta\sigma_z$", color='red', lw=2)
    ax_b.plot(sx_v, z_vals, label=r"$\Delta\sigma_x$", color='blue', ls='--')
    ax_b.plot(sy_v, z_vals, label=r"$\Delta\sigma_y$", color='purple', ls='-.')
    ax_b.plot(sv0_v,z_vals, label=r"$0.20\sigma'_{v0}$", color='green', lw=2)
    if zi <= espesor_total:
        ax_b.axhline(y=zi, color='orange', ls=':', lw=1.5, label=f'z_i={zi:.2f} m')
    if NF < espesor_total:
        ax_b.axhline(y=NF, color='deepskyblue', ls='-.', lw=1.2, label=f'NF={NF:.1f} m')
    ax_b.set_ylim(espesor_total,0); ax_b.set_xlim(left=0)
    ax_b.set_xlabel("Tensión (kPa)"); ax_b.set_ylabel("Profundidad z (m)")
    ax_b.set_title("Bulbo de presiones"); ax_b.legend(fontsize=8)
    ax_b.grid(True, linestyle=':', alpha=0.4)
    ax_b.spines[['top','right']].set_visible(False)
    plt.tight_layout()
    fig_bulbo_bytes = _fig_bytes(fig_b)
    plt.close(fig_b)

    word_buf = generar_word(
        B, L, p, NF, z_max_user, zi,
        st.session_state.df_terreno,
        df_st, tot_st, df_ec, tot_ec,
        fig_bulbo_bytes
    )
    st.sidebar.download_button(
        "📝 Descargar Informe Word", data=word_buf,
        file_name=f"informe_comparativo_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )
else:
    st.sidebar.button("📝 Descargar Informe Word", disabled=True,
                      use_container_width=True,
                      help="Primero calcula.")

# ══════════════════════════════════════════════════════════════════════════
# ÁREA PRINCIPAL
# ══════════════════════════════════════════════════════════════════════════
st.title("🏗️ Cálculo de Asientos de cimentaciones rectangulares")
st.markdown("**Método 1 — Modelo Steinbrenner**")
st.markdown("**Método 2 — Modelo Elástico**")
st.markdown("---")

# ══════════════════════════════════════════════════
# VISTA 1: PANEL DE CÁLCULO
# ══════════════════════════════════════════════════
if modo == "🧮 Panel de Cálculo":
    st.header("1. Estratigrafía del Terreno")
    st.caption("Edita las capas libremente. Los pesos específicos se usan para el criterio EC7 (z_i).")

    df_edit = st.data_editor(
        st.session_state.df_terreno,
        num_rows="dynamic",
        use_container_width=True,
        column_config={
            "E (kPa)":               st.column_config.NumberColumn("E (kPa)",      min_value=100.0, step=500.0),
            "nu":                    st.column_config.NumberColumn("ν",             min_value=0.0, max_value=0.5, step=0.01, format="%.2f"),
            "Peso Esp. (kN/m³)":     st.column_config.NumberColumn("γ (kN/m³)",    min_value=10.0, max_value=25.0, step=0.5),
            "Peso Esp. Sat (kN/m³)": st.column_config.NumberColumn("γsat (kN/m³)", min_value=10.0, max_value=25.0, step=0.5),
        }
    )
    if not df_edit.equals(st.session_state.df_terreno):
        st.session_state.df_terreno = df_edit
        st.session_state.calculo_realizado = False
        st.rerun()

    st.markdown("---")
    st.header("2. Resultados y Comparativa")

    if not st.session_state.calculo_realizado:
        st.info("👈 Pulsa **Calcular** en el panel izquierdo.")
    else:
        df_st = st.session_state.df_st
        df_ec = st.session_state.df_ec
        tot_st= st.session_state.tot_st
        tot_ec= st.session_state.tot_ec

        # ── Métricas principales ──────────────────────────────────────────
        dz_used = st.session_state.get('dz_used', 0.25)
        st.success(f"✅ Cálculo completado. **Ec. elástica** integrada con subcapas de **{dz_used} m**.")
        c1, c2, c3 = st.columns(3)
        c1.metric("🔵 Steinbrenner",f"{tot_st*1000:.3f} mm")
        c2.metric("🟢 Ec. Elástica", f"{tot_ec*1000:.3f} mm",
                  help=f"Integrada con dz = {dz_used} m")
        dif = abs(tot_st - tot_ec)*1000
        pct = abs(tot_st - tot_ec)/max(abs(tot_st), 1e-9)*100
        c3.metric("📊 Diferencia",f"{dif:.3f} mm", f"{pct:.1f}%")

        # ── Tabla comparativa por capa ────────────────────────────────────
        st.subheader("Tabla comparativa por capa")
        df_comp = pd.DataFrame({
            "Capa":                  df_st["Capa"],
            "z Techo [m]":           df_st["z Techo [m]"],
            "z Base [m]":            df_st["z Base [m]"],
            "Δs Steinbrenner [mm]":  df_st["Δs [mm]"],
            "Δs Ec.68 [mm]":         df_ec["Δs [mm]"].values,
            "Diferencia [mm]":       (df_st["Δs [mm]"].values - df_ec["Δs [mm]"].values).round(3),
        })
        st.dataframe(df_comp, use_container_width=True, hide_index=True,
            column_config={
                "Capa":                 st.column_config.TextColumn("Capa", width="medium"),
                "z Techo [m]":          st.column_config.NumberColumn("z Techo [m]",         width="small", format="%.2f"),
                "z Base [m]":           st.column_config.NumberColumn("z Base [m]",           width="small", format="%.2f"),
                "Δs Steinbrenner [mm]": st.column_config.NumberColumn("Δs Steinbrenner [mm]", width="medium",format="%.3f"),
                "Δs Ec.elástica [mm]":        st.column_config.NumberColumn("Δs Ec. 68 [mm]",       width="medium",format="%.3f"),
                "Diferencia [mm]":      st.column_config.NumberColumn("Diferencia [mm]",       width="medium",format="%.3f"),
            }
        )

# ══════════════════════════════════════════════════
# VISTA 2: DETALLE STEINBRENNER
# ══════════════════════════════════════════════════
elif modo == "📋 Detalle Steinbrenner":
    st.header("📋 Detalle Método Steinbrenner")
    st.latex(r"s = \frac{p \cdot B}{E}\left[(1-\nu^2)\phi_1 - (1-\nu-2\nu^2)\phi_2\right]"
             r"\quad \Delta s_i = s(z_{techo}) - s(z_{base})")

    if not st.session_state.calculo_realizado:
        st.warning("⚠️ Calcula primero.")
    else:
        df_st = st.session_state.df_st
# ══════════════════════════════════════════════════
# VISTA 2: DETALLE STEINBRENNER
# ══════════════════════════════════════════════════
elif modo == "📋 Modelo Steinbrenner":
    st.header("📋 Detalle Método Steinbrenner")
    st.markdown("Cálculo capa a capa integrando las funciones de influencia $\\phi_1$ y $\\phi_2$.")
    
    if not st.session_state.calculo_realizado:
        st.warning("⚠️ Calcula primero en el panel izquierdo.")
    else:
        df_st = st.session_state.df_st
        st.markdown("##### 🔼 Valores en el Techo")
        st.dataframe(df_st[["Capa","z Techo [m]","m_techo","φ1_techo","φ2_techo","s_techo [mm]"]],
                     use_container_width=True, hide_index=True)
        st.markdown("##### 🔽 Valores en la Base")
        st.dataframe(df_st[["Capa","z Base [m]","m_base","φ1_base","φ2_base","s_base [mm]"]],
                     use_container_width=True, hide_index=True)
        st.markdown("##### 📊 Asiento por estrato")
        st.dataframe(df_st[["Capa","Δs [mm]"]], use_container_width=True, hide_index=True)
        st.metric("🔵 Asiento Total Steinbrenner",
                  f"{st.session_state.tot_st*1000:.3f} mm")

# ══════════════════════════════════════════════════
# VISTA 3: DETALLE EC ELÁSTICO
# ══════════════════════════════════════════════════
elif modo == "📋 Modelo Elástico":
    st.header("📋 Detalle Método Ecuación Elástica")
    st.latex(r"s = \sum_{i=1}^{n}\left[\frac{h}{E}\left(\Delta\sigma_z - \nu(\Delta\sigma_x+\Delta\sigma_y)\right)\right]_i")

    if not st.session_state.calculo_realizado:
        st.warning("⚠️ Calcula primero.")
    else:
        df_ec = st.session_state.df_ec
        dz_used = st.session_state.get('dz_used', 0.25)
        st.caption(f"Integración con subcapas de **{dz_used} m** por estrato. Los valores de Δσ y Δεz son promedios de las subcapas.")
        st.markdown("##### ⚡ Tensiones de Holl — promedio por capa")
        st.dataframe(df_ec[["Capa","Sub-capas","Δσz med [kPa]","Δσx med [kPa]","Δσy med [kPa]"]],
                     use_container_width=True, hide_index=True)
        st.markdown("##### 📐 Deformación unitaria media y asiento")
        st.dataframe(df_ec[["Capa","h_ef [m]","Sub-capas","Δεz med [-]","Δs [mm]"]],
                     use_container_width=True, hide_index=True)
        st.metric("🟢 Asiento Total Ec. 68",
                  f"{st.session_state.tot_ec*1000:.3f} mm")

# ══════════════════════════════════════════════════
# VISTA 4: BULBO DE PRESIONES
# ══════════════════════════════════════════════════
elif modo == "📉 Bulbo de Presiones":
    st.header("Bulbo de Presiones y Zona de Influencia")
    st.markdown(
        r"Tensiones bajo el centro ($\times 4$ superposición $B/2 \times L/2$). "
        r"El criterio EC7 es: $\Delta\sigma_z \leq 0.20\,\sigma'_{v0}$."
    )
    col1, col2 = st.columns([1, 3])
    with col1:
        z_gr = st.slider("Profundidad máxima [m]:", 1.0, espesor_total,
                         min(espesor_total, 15.0), 0.5)
        st.markdown(f"**p:** {p} kPa · **B:** {B} m · **L:** {L} m")
        if NF < 100.0: st.markdown(f"**NF:** {NF:.1f} m")
        st.metric("📐 z_i (EC7)", f"{zi:.2f} m")
        st.markdown("---")
        st.info("Las tensiones de Holl son **idénticas** para ambos métodos. La diferencia entre métodos está en cómo se integra el asiento.")
    with col2:
        z_vals = np.linspace(0.05, z_gr, 200)
        sz_v,sx_v,sy_v,sv0_v,umb20_v = [],[],[],[],[]
        for z in z_vals:
            sz,sx,sy = holl_centro(p, B, L, z)
            sv = sigma_v0(z, st.session_state.df_terreno, NF)
            sz_v.append(sz); sx_v.append(sx); sy_v.append(sy)
            sv0_v.append(sv); umb20_v.append(0.20*sv)

        fig, ax = plt.subplots(figsize=(9, 7))
        ax.plot(sz_v,   z_vals, label=r"Vertical $\Delta\sigma_z$",           color='red',         lw=2)
        ax.plot(sx_v,   z_vals, label=r"Horiz. Trans. $\Delta\sigma_x$",      color='blue',        ls='--')
        ax.plot(sy_v,   z_vals, label=r"Horiz. Long. $\Delta\sigma_y$",       color='purple',      ls='-.')
        ax.plot(sv0_v,  z_vals, label=r"$\sigma'_{v0}$ (tensión efect.)",     color='saddlebrown', ls=':', lw=1.5)
        ax.plot(umb20_v,z_vals, label=r"$0.20\,\sigma'_{v0}$ (criterio EC7)", color='green',       lw=2)
        if zi <= z_gr:
            ax.axhline(y=zi, color='orange', ls='--', lw=1.5,
                       label=f'z_i EC7 = {zi:.2f} m')
            ax.annotate(f' z_i = {zi:.2f} m', xy=(0, zi),
                        xytext=(p*0.04, zi - z_gr*0.03),
                        color='darkorange', fontsize=10, fontweight='bold')
        if NF < z_gr and NF < 100.0:
            ax.axhline(y=NF, color='deepskyblue', ls='-.', lw=1.2, label=f'NF = {NF:.1f} m')
        ax.set_ylim(z_gr, 0); ax.set_xlim(left=0)
        ax.set_xlabel("Tensión (kPa)", fontsize=11)
        ax.set_ylabel("Profundidad z (m)", fontsize=11)
        ax.set_title("Bulbo de presiones — Centro de la cimentación", fontsize=13)
        ax.legend(loc='lower right', fontsize=9)
        ax.grid(True, linestyle=':', alpha=0.5)
        ax.spines[['top','right']].set_visible(False)
        st.pyplot(fig); plt.close(fig)

# ══════════════════════════════════════════════════
# VISTA 5: FUNDAMENTO TEÓRICO
# ══════════════════════════════════════════════════
elif modo == "📖 Fundamento Teórico":
    st.header("Fundamento Teórico — Comparativa de Formulaciones")

    col_a, col_b = st.columns(2)

    with col_a:
        st.subheader("🔵 Método 1 — Steinbrenner")
        st.markdown("Integración analítica del campo de asientos usando los factores geométricos φ₁ y φ₂:")
        st.latex(r"s(z) = \frac{p \cdot B}{E}\left[(1-\nu^2)\phi_1 - (1-\nu-2\nu^2)\phi_2\right]")
        st.latex(r"\phi_1 = \frac{1}{\pi}\left[\ln\frac{\sqrt{1+m^2+n^2}+n}{\sqrt{1+m^2}} + n\ln\frac{\sqrt{1+m^2+n^2}+1}{\sqrt{n^2+m^2}}\right]")
        st.latex(r"\phi_2 = \frac{m}{\pi}\arctan\frac{n}{m\sqrt{1+m^2+n^2}}")
        st.markdown(r"Con $n = L/B$ y $m = 2z/B$. El asiento de cada estrato:")
        st.latex(r"\Delta s_i = s(z_{techo}) - s(z_{base})")
        st.info("El asiento integra implícitamente la distribución de tensiones en profundidad.")

    with col_b:
        st.subheader("🟢 Método 2 — Ecuación Elástica")
        st.markdown("Integración explícita de la deformación unitaria vertical en cada estrato:")
        st.latex(r"s = \sum_{i=1}^{n}\left[\frac{h}{E}\left(\Delta\sigma_z - \nu(\Delta\sigma_x+\Delta\sigma_y)\right)\right]_i")
        st.markdown(r"Las tensiones se evalúan en el **punto medio** de cada estrato ($z_{mid}$):")
        st.latex(r"\Delta\varepsilon_z = \frac{\Delta\sigma_z - \nu(\Delta\sigma_x+\Delta\sigma_y)}{E}")
        st.latex(r"\Delta s_i = \Delta\varepsilon_z \cdot h_i")
        st.info("Permite apreciar la contribución de las tensiones horizontales al asiento vertical.")

    st.markdown("---")
    st.subheader("🔁 Tensiones de Holl — compartidas por ambos métodos")
    st.markdown(
        r"Ambas formulaciones usan las tensiones de Holl bajo la **esquina** de una carga rectangular, "
        r"aplicando superposición ×4 con $B/2 \times L/2$ para obtener el **centro** de la zapata "
        r"(UNE-EN 1997-1, 6.6.2(15)):"
    )
    st.latex(r"\sigma_z = \frac{p}{2\pi}\left[\arctan\frac{BL}{zR_3} + BL\left(\frac{1}{R_1^2}+\frac{1}{R_2^2}\right)\frac{z}{R_3}\right]")
    st.latex(r"\sigma_x = \frac{p}{2\pi}\left[\arctan\frac{BL}{zR_3} - \frac{BLz}{R_1^2 R_3}\right]")
    st.latex(r"\sigma_y = \frac{p}{2\pi}\left[\arctan\frac{BL}{zR_3} - \frac{BLz}{R_2^2 R_3}\right]")
    st.latex(r"R_1=\sqrt{L^2+z^2}\quad R_2=\sqrt{B^2+z^2}\quad R_3=\sqrt{L^2+B^2+z^2}")

    st.markdown("---")
    st.subheader("📐 Criterio de Profundidad de Influencia ")
    st.latex(r"\Delta\sigma_z(z_i) \leq 0.20\,\sigma'_{v0}(z_i)")
    st.markdown(
        r"Con $\sigma'_{v0}$ = tensión efectiva geoestática, considerando el nivel freático: "
        r"$\gamma_i$ sobre NF y $\gamma_{sat,i} - \gamma_w$ bajo NF."
    )
