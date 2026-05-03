import streamlit as st
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import io
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from scipy.optimize import brentq

# ══════════════════════════════════════════════════════════════════════════
# CONSTANTES
# ══════════════════════════════════════════════════════════════════════════
GAMMA_AGUA = 9.81  # kN/m³

# ══════════════════════════════════════════════════════════════════════════
# TENSIONES DE HOLL — BAJO EL CENTRO (compartido por ambos métodos)
# ══════════════════════════════════════════════════════════════════════════
def holl_esquina(p, B, L, z):
    if z <= 1e-6:
        return p, p / 2.0, p / 2.0
    R1 = np.sqrt(L**2 + z**2)
    R2 = np.sqrt(B**2 + z**2)
    R3 = np.sqrt(L**2 + B**2 + z**2)
    arc = np.arctan((B * L) / (z * R3))
    sz = (p / (2*np.pi)) * (arc + B*L*(1/R1**2 + 1/R2**2)*(z/R3))
    sx = (p / (2*np.pi)) * (arc - (B*L*z)/(R1**2*R3))
    sy = (p / (2*np.pi)) * (arc - (B*L*z)/(R2**2*R3))
    return sz, sx, sy

def holl_centro(p, B, L, z):
    sz, sx, sy = holl_esquina(p, B/2.0, L/2.0, z)
    return 4*sz, 4*sx, 4*sy

# ══════════════════════════════════════════════════════════════════════════
# MÉTODO 1 — STEINBRENNER (φ1, φ2, s(z) analítico)
# ══════════════════════════════════════════════════════════════════════════
def phi1(m, n):
    if m == 0:
        t1 = np.log(np.sqrt(1+n**2)+n)
        t2 = n*np.log((np.sqrt(1+n**2)+1)/n)
    else:
        t1 = np.log((np.sqrt(1+m**2+n**2)+n)/np.sqrt(1+m**2))
        t2 = n*np.log((np.sqrt(1+m**2+n**2)+1)/np.sqrt(n**2+m**2))
    return (1/np.pi)*(t1+t2)

def phi2(m, n):
    if m == 0: return 0.0
    return (m/np.pi)*np.arctan(n/(m*np.sqrt(1+m**2+n**2)))

def s_z(p, B, E, nu, z, L):
    n = L/B
    m = z/B
    corchete = (1-nu**2)*phi1(m,n) - (1-nu-2*nu**2)*phi2(m,n)
    return (p*B/E)*corchete

def calcular_steinbrenner(p, B, L, df, z_max):
    total = 0.0
    resultados = []
    z_actual = 0.0
    n_factor = L / B

    for _, row in df.iterrows():
        if z_actual >= z_max:
            break
        h_i   = float(row["Espesor (m)"])
        E_i   = float(row["E (kPa)"])
        nu_i  = float(row["nu"])
        nombre= str(row["Descripción"])

        z_techo = z_actual
        z_base  = min(z_actual + h_i, z_max)

        m_t = z_techo / (B/2)
        m_b = z_base  / (B/2)
        
        s_t = 4 * s_z(p, B/2, E_i, nu_i, z_techo, L/2)
        s_b = 4 * s_z(p, B/2, E_i, nu_i, z_base,  L/2)
        ds  = s_t - s_b
        total += ds

        resultados.append({
            "Capa":               nombre,
            "z Techo [m]":        round(z_techo, 3),
            "z Base [m]":         round(z_base,  3),
            "m_techo":            round(m_t, 4),
            "φ1_techo":           round(phi1(m_t, n_factor), 4),
            "φ2_techo":           round(phi2(m_t, n_factor), 4),
            "s_techo [mm]":       round(s_t*1000, 3),
            "m_base":             round(m_b, 4),
            "φ1_base":            round(phi1(m_b, n_factor), 4),
            "φ2_base":            round(phi2(m_b, n_factor), 4),
            "s_base [mm]":        round(s_b*1000, 3),
            "Δs [mm]":            round(ds*1000, 3),
        })
        z_actual = z_base

    return total, pd.DataFrame(resultados)

# ══════════════════════════════════════════════════════════════════════════
# MÉTODO 2 — (integración directa de deformaciones unitarias)
# ══════════════════════════════════════════════════════════════════════════
def calcular_ec68(p, B, L, df, z_max, dz_sub=0.25):
    total = 0.0
    resultados = []
    z_actual = 0.0

    for _, row in df.iterrows():
        if z_actual >= z_max:
            break
        h_i    = float(row["Espesor (m)"])
        E_i    = float(row["E (kPa)"])
        nu_i   = float(row["nu"])
        nombre = str(row["Descripción"])

        z_techo = z_actual
        z_base  = min(z_actual + h_i, z_max)
        h_ef    = z_base - z_techo

        n_sub  = max(1, int(np.ceil(h_ef / dz_sub)))
        dz     = h_ef / n_sub

        ds_capa, sz_medio, sx_medio, sy_medio, ez_medio = 0.0, 0.0, 0.0, 0.0, 0.0

        for k in range(n_sub):
            z_sub_t = z_techo + k * dz
            z_mid   = z_sub_t + dz / 2.0
            dsz, dsx, dsy = holl_centro(p, B, L, z_mid)
            dep_z  = (dsz - nu_i*(dsx+dsy)) / E_i
            ds_sub = dep_z * dz
            ds_capa  += ds_sub
            sz_medio += dsz
            sx_medio += dsx
            sy_medio += dsy
            ez_medio += dep_z

        sz_medio /= n_sub
        sx_medio /= n_sub
        sy_medio /= n_sub
        ez_medio /= n_sub

        total += ds_capa

        resultados.append({
            "Capa":          nombre,
            "z Techo [m]":   round(z_techo,  3),
            "z Base [m]":    round(z_base,   3),
            "h_ef [m]":      round(h_ef,     3),
            "Sub-capas":      n_sub,
            "Δσz med [kPa]": round(sz_medio, 3),
            "Δσx med [kPa]": round(sx_medio, 3),
            "Δσy med [kPa]": round(sy_medio, 3),
            "Δεz med [-]":   round(ez_medio, 6),
            "Δs [mm]":        round(ds_capa*1000, 3),
        })
        z_actual = z_base

    return total, pd.DataFrame(resultados)

# ══════════════════════════════════════════════════════════════════════════
# TENSIÓN EFECTIVA Y ZONA DE INFLUENCIA
# ══════════════════════════════════════════════════════════════════════════
def sigma_v0(z, df, NF):
    sv = 0.0; z_act = 0.0
    for _, row in df.iterrows():
        h  = float(row["Espesor (m)"])
        g  = float(row["Peso Esp. (kN/m³)"])
        gs = float(row["Peso Esp. Sat (kN/m³)"])
        zt = z_act; zb = z_act + h
        if z <= zt: break
        ze = min(z, zb)
        z_sec_b = min(ze, NF)
        if z_sec_b > zt: sv += g*(z_sec_b-zt)
        z_sat_t = max(zt, NF)
        if ze > z_sat_t: sv += (gs-GAMMA_AGUA)*(ze-z_sat_t)
        z_act = zb
    return sv

def z_influencia_seleccionada(p, B, L, df, NF, z_max_usuario, criterio_elegido):
    et = float(pd.to_numeric(df["Espesor (m)"]).sum())
    
    if criterio_elegido == "Límite Geométrico (2B)":
        z_calc = 2.0 * B
    elif criterio_elegido == "Límite Geométrico (3B)":
        z_calc = 3.0 * B
    else: 
        z = 0.05
        z_calc = et
        while z <= et:
            dsz, _, _ = holl_centro(p, B, L, z)
            sv = sigma_v0(z, df, NF)
            if sv > 0 and dsz <= 0.20 * sv:
                z_calc = z
                break
            z += 0.05
            
    return min(z_calc, et, z_max_usuario)

# ══════════════════════════════════════════════════════════════════════════
# SOLVER INVERSO: ENCONTRAR PRESION ADMISIBLE
# ══════════════════════════════════════════════════════════════════════════
def encontrar_presion_admisible(metodo_func, s_max_objetivo, B, L, df_terreno, NF, z_max_usuario, criterio_elegido, dz_sub=None):
    def diferencia_asiento(p_prueba):
        zi_actual = z_influencia_seleccionada(p_prueba, B, L, df_terreno, NF, z_max_usuario, criterio_elegido)
        if dz_sub is None:
            asiento_calc, _ = metodo_func(p_prueba, B, L, df_terreno, zi_actual)
        else:
            asiento_calc, _ = metodo_func(p_prueba, B, L, df_terreno, zi_actual, dz_sub)
        return asiento_calc - s_max_objetivo

    try:
        p_optima = brentq(diferencia_asiento, 1.0, 5000.0, xtol=0.1)
        zi_final = z_influencia_seleccionada(p_optima, B, L, df_terreno, NF, z_max_usuario, criterio_elegido)
        if dz_sub is None:
            tot, df_res = metodo_func(p_optima, B, L, df_terreno, zi_final)
        else:
            tot, df_res = metodo_func(p_optima, B, L, df_terreno, zi_final, dz_sub)
        return p_optima, tot, df_res, zi_final
    except ValueError:
        return None, 0.0, pd.DataFrame(), 0.0

# ══════════════════════════════════════════════════════════════════════════
# INFORME WORD ESTÉTICO
# ══════════════════════════════════════════════════════════════════════════
def _fig_bytes(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=250, bbox_inches='tight')
    buf.seek(0)
    return buf

def _add_styled_table(doc, df, title):
    h = doc.add_heading(title, level=2)
    h.runs[0].font.color.rgb = RGBColor(31, 73, 125)
    df = df.astype(str)
    table = doc.add_table(rows=1+len(df), cols=len(df.columns))
    table.style = 'Light Shading Accent 1' 
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    hdr_cells = table.rows[0].cells
    for i, column in enumerate(df.columns):
        hdr_cells[i].text = column
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for paragraph in hdr_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(8) 
                run.font.color.rgb = RGBColor(23, 54, 93) 
    
    for i, row in enumerate(df.itertuples(index=False)):
        row_cells = table.rows[i+1].cells
        for j, value in enumerate(row):
            row_cells[j].text = str(value)
            row_cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in row_cells[j].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(8)
    doc.add_paragraph()

# ORDEN CORREGIDO: criterio_bulbo antes de df_terreno
def generar_word(B, L, s_max_mm, p_st, p_ec, NF, z_max, zi_final, criterio_bulbo, df_terreno,
                 df_st, df_ec, fig_bulbo_bytes):
    fecha = datetime.now().strftime("%d de %B de %Y — %H:%M")
    doc = Document()
    
    for sec in doc.sections:
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(2.0) 
        sec.right_margin = Cm(2.0)
        footer = sec.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"Memoria de Cálculo Geotécnico — Generado automáticamente el {fecha}"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer_para.runs[0].font.size = Pt(8)
        footer_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)

    h1_style = doc.styles['Heading 1']
    h1_font = h1_style.font
    h1_font.name = 'Calibri Light'
    h1_font.size = Pt(14)
    h1_font.color.rgb = RGBColor(23, 54, 93)
    h1_font.bold = True

    doc.add_paragraph()
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run('CÁLCULO DE PRESIÓN ADMISIBLE (ELS)')
    r_title.bold = True
    r_title.font.size = Pt(22)
    r_title.font.color.rgb = RGBColor(23, 54, 93) 
    
    doc.add_paragraph()
    
    doc.add_heading('1. Datos de Entrada', level=1)
    
    table_params = doc.add_table(rows=6, cols=2)
    table_params.style = 'Light Shading Accent 1'
    
    data = [
        ('Dimensiones en planta (B × L)', f'{B:.2f} m  ×  {L:.2f} m'),
        ('Asiento máximo admisible', f'{s_max_mm:.1f} mm'),
        ('Profundidad del Nivel Freático (NF)', f'{NF:.1f} m'),
        ('Criterio de Bulbo Seleccionado', criterio_bulbo),
        ('Profundidad de influencia activa (zi)', f'{zi_final:.2f} m'),
        ('Profundidad de bulbo manual', f'{z_max:.1f} m')
    ]
    
    for i, (key, val) in enumerate(data):
        row = table_params.rows[i].cells
        row[0].text = key
        row[1].text = val
        row[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row[0].paragraphs[0].runs[0].font.bold = True
        
    doc.add_paragraph()
    _add_styled_table(doc, df_terreno, '1.1 Estratigrafía del Perfil Geotécnico')
    doc.add_page_break()

    doc.add_heading('2. Resultados: Presiones Máximas Admisibles', level=1)
    
    table_res = doc.add_table(rows=1, cols=2)
    table_res.style = 'Light Shading Accent 1'
    table_res.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    c1, c2 = table_res.rows[0].cells
    c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    c2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    p1 = c1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1a = p1.add_run('Presión máx. Steinbrenner\n')
    r1a.font.size = Pt(11); r1a.font.color.rgb = RGBColor(89, 89, 89); r1a.bold = True
    r1b = p1.add_run(f'{p_st:.1f} kPa')
    r1b.font.size = Pt(16); r1b.bold = True; r1b.font.color.rgb = RGBColor(23, 54, 93)
    
    p2 = c2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2a = p2.add_run('Presión máx. Ec. Elástica\n')
    r2a.font.size = Pt(11); r2a.font.color.rgb = RGBColor(89, 89, 89); r2a.bold = True
    r2b = p2.add_run(f'{p_ec:.1f} kPa')
    r2b.font.size = Pt(16); r2b.bold = True; r2b.font.color.rgb = RGBColor(33, 115, 70)

    doc.add_paragraph()
    p_info = doc.add_paragraph('Desglose tensional y deformacional para la presión límite evaluada por cada método.')
    p_info.runs[0].font.color.rgb = RGBColor(89, 89, 89)
    
    _add_styled_table(doc, df_st, f'2.1 Método de Steinbrenner (Evaluado a p = {p_st:.1f} kPa)')
    doc.add_page_break()
    _add_styled_table(doc, df_ec, f'2.2 Método Elástico (Evaluado a p = {p_ec:.1f} kPa)')

    doc.add_page_break()
    doc.add_heading('3. Gráfica del Bulbo de Tensiones', level=1)
    p_fig = doc.add_paragraph()
    p_fig.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_fig = p_fig.add_run()
    r_fig.add_picture(fig_bulbo_bytes, width=Cm(14))
    
    nota = doc.add_paragraph(f'Evolución en profundidad de las tensiones bajo el centro geométrico de la zapata. Gráfica trazada para la presión admisible del Método Elástico (p = {p_ec:.1f} kPa).')
    nota.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nota.runs[0].font.size = Pt(9)
    nota.runs[0].font.italic = True
    nota.runs[0].font.color.rgb = RGBColor(128,128,128)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ══════════════════════════════════════════════════════════════════════════
# SESSION STATE
# ══════════════════════════════════════════════════════════════════════════
def reset_calculo():
    st.session_state.calculo_realizado = False

if 'calculo_realizado' not in st.session_state:
    st.session_state.calculo_realizado = False

if 'df_terreno' not in st.session_state:
    st.session_state.df_terreno = pd.DataFrame({
        "Descripción":           ["UG-01" , "UG-02",  "UG-03"],
        "Espesor (m)":           [1.5,         3.0,        5.0],
        "E (kPa)":               [10000.0,     15000.0,     40000.0],
        "nu":                    [0.30,         0.45,       0.25],
        "Peso Esp. (kN/m³)":     [18.0,         19.0,       21.0],
        "Peso Esp. Sat (kN/m³)": [20.0,         20.0,       22.0],
    })

# ══════════════════════════════════════════════════════════════════════════
# CONFIGURACIÓN UI
# ══════════════════════════════════════════════════════════════════════════
st.set_page_config(
    page_title="Cálculo de Presión Admisible (ELS)",
    layout="wide", page_icon="🏗️"
)

st.sidebar.title("Navegación")
modo = st.sidebar.radio("Vista:", [
    "🧮 Panel de Resultados",
    "📋 Detalle Steinbrenner",
    "📋 Detalle Elástico",
    "📉 Bulbo de Presiones",
    "📖 Fundamento Teórico",
])

st.sidebar.markdown("---")
st.sidebar.header("📥 Datos de Entrada")

B  = st.sidebar.number_input("Ancho (B) [m]",             min_value=0.1, value=2.0,   step=0.1,  on_change=reset_calculo)
L  = st.sidebar.number_input("Longitud (L) [m]",          min_value=0.1, value=3.0,   step=0.1,  on_change=reset_calculo)
s_max_mm = st.sidebar.number_input("Asiento máx admisible [mm]", min_value=1.0, value=25.0, step=1.0, on_change=reset_calculo)
NF = st.sidebar.number_input("Nivel Freático [m]",         min_value=0.0, value=100.0, step=0.5,  on_change=reset_calculo)

s_max_m = s_max_mm / 1000.0

if L < B:
    B, L = L, B
    st.sidebar.warning("⚠️ L<B: valores intercambiados.")

espesor_total = max(float(pd.to_numeric(st.session_state.df_terreno["Espesor (m)"]).sum()), 0.1)

st.sidebar.markdown("---")
st.sidebar.subheader("📐 Criterio de Bulbo de Presiones")

criterio_bulbo = st.sidebar.selectbox(
    "Selecciona el criterio limitante:",
    options=["Criterio EC7 (Δσz ≤ 0.20 σ'v0)", "Límite Geométrico (2B)", "Límite Geométrico (3B)"],
    index=0,
    on_change=reset_calculo
)

z_max_user = st.sidebar.number_input(
    "Profundidad de bulbo manual [m]",
    min_value=0.1, max_value=espesor_total,
    value=float(espesor_total),
    step=0.1, on_change=reset_calculo,
    help="Límite físico forzado adicional. El cálculo se detendrá aquí si el criterio seleccionado arriba llega más profundo."
)

st.sidebar.markdown("---")
st.sidebar.subheader("🔧 Precisión Ecuación Elástica")
dz_sub = st.sidebar.select_slider(
    "Tamaño de subcapa (dz) [m]",
    options=[2.0, 1.0, 0.5, 0.25, 0.10, 0.05],
    value=0.10,
    on_change=reset_calculo
)

# ══════════════════════════════════════════════════════════════════════════
# BOTÓN CALCULAR (SOLVER INVERSO)
# ══════════════════════════════════════════════════════════════════════════
st.sidebar.markdown("---")
if st.sidebar.button("🚀 Calcular Presión Admisible", type="primary", use_container_width=True):
    with st.spinner("Iterando presiones admisibles..."):
        p_st, tot_st, df_st, zi_st = encontrar_presion_admisible(
            calcular_steinbrenner, s_max_m, B, L, st.session_state.df_terreno, NF, z_max_user, criterio_bulbo
        )
        
        p_ec, tot_ec, df_ec, zi_ec = encontrar_presion_admisible(
            calcular_ec68, s_max_m, B, L, st.session_state.df_terreno, NF, z_max_user, criterio_bulbo, dz_sub
        )

        st.session_state.p_st = p_st
        st.session_state.p_ec = p_ec
        st.session_state.df_st = df_st
        st.session_state.df_ec = df_ec
        st.session_state.zi_final = max(zi_st, zi_ec) 
        st.session_state.dz_used = dz_sub
        st.session_state.criterio_usado = criterio_bulbo
        st.session_state.calculo_realizado = True

# ══════════════════════════════════════════════════════════════════════════
# BOTÓN INFORME WORD
# ══════════════════════════════════════════════════════════════════════════
st.sidebar.markdown("---")
if st.session_state.calculo_realizado and st.session_state.p_ec is not None:
    p_ref = st.session_state.p_ec 
    z_vals = np.linspace(0.05, espesor_total, 200)
    sz_v, sx_v, sy_v, sv0_v = [], [], [], []
    for z in z_vals:
        sz, sx, sy = holl_centro(p_ref, B, L, z)
        sz_v.append(sz); sx_v.append(sx); sy_v.append(sy)
        sv0_v.append(sigma_v0(z, st.session_state.df_terreno, NF)*0.20)
        
    fig_b, ax_b = plt.subplots(figsize=(5, 7))
    ax_b.plot(sz_v, z_vals, label=r"$\Delta\sigma_z$", color='red', lw=2)
    ax_b.plot(sx_v, z_vals, label=r"$\Delta\sigma_x$", color='blue', ls='--')
    ax_b.plot(sy_v, z_vals, label=r"$\Delta\sigma_y$", color='purple', ls='-.')
    ax_b.plot(sv0_v,z_vals, label=r"$0.20\sigma'_{v0}$", color='green', lw=2)
    zi_plot = st.session_state.zi_final
    if zi_plot <= espesor_total:
        ax_b.axhline(y=zi_plot, color='orange', ls=':', lw=1.5, label=f'z_i={zi_plot:.2f} m')
    if NF < espesor_total:
        ax_b.axhline(y=NF, color='deepskyblue', ls='-.', lw=1.2, label=f'NF={NF:.1f} m')
    ax_b.set_ylim(espesor_total, 0); ax_b.set_xlim(left=0)
    ax_b.set_xlabel("Tensión (kPa)"); ax_b.set_ylabel("Profundidad z (m)")
    ax_b.set_title(f"Bulbo para p_adm = {p_ref:.1f} kPa"); ax_b.legend(fontsize=8)
    ax_b.grid(True, linestyle=':', alpha=0.4)
    ax_b.spines[['top','right']].set_visible(False)
    plt.tight_layout()
    fig_bulbo_bytes = _fig_bytes(fig_b)
    plt.close(fig_b)

    word_buf = generar_word(
        B, L, s_max_mm, st.session_state.p_st, st.session_state.p_ec, NF, z_max_user, 
        st.session_state.zi_final, st.session_state.criterio_usado, st.session_state.df_terreno,
        st.session_state.df_st, st.session_state.df_ec, fig_bulbo_bytes
    )
    st.sidebar.download_button(
        "📝 Descargar Informe Word", data=word_buf,
        file_name=f"presion_admisible_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )

# ══════════════════════════════════════════════════════════════════════════
# ÁREA PRINCIPAL
# ══════════════════════════════════════════════════════════════════════════
st.title("🏗️ Cálculo Inverso: Presión Admisible (ELS)")
st.markdown(f"**Asiento máximo objetivo del proyecto:** {s_max_mm} mm")
st.markdown("---")

if modo == "🧮 Panel de Resultados":
    st.header("1. Estratigrafía del Terreno")
    df_edit = st.data_editor(
        st.session_state.df_terreno,
        num_rows="dynamic",
        use_container_width=True,
        column_config={
            "E (kPa)":               st.column_config.NumberColumn("E (kPa)",      min_value=100.0, step=500.0),
            "nu":                    st.column_config.NumberColumn("ν",             min_value=0.0, max_value=0.5, step=0.01, format="%.2f"),
            "Peso Esp. (kN/m³)":     st.column_config.NumberColumn("γ (kN/m³)",    min_value=10.0, max_value=25.0, step=0.5),
            "Peso Esp. Sat (kN/m³)": st.column_config.NumberColumn("γsat (kN/m³)", min_value=10.0, max_value=25.0, step=0.5),
        }
    )
    if not df_edit.equals(st.session_state.df_terreno):
        st.session_state.df_terreno = df_edit
        st.session_state.calculo_realizado = False
        st.rerun()

    st.markdown("---")
    st.header("2. Presión Máxima Calculada")

    if not st.session_state.calculo_realizado:
        st.info("👈 Pulsa **Calcular Presión Admisible** en el panel izquierdo.")
    else:
        p_st = st.session_state.p_st
        p_ec = st.session_state.p_ec
        zi_final = st.session_state.zi_final

        if p_st is None or p_ec is None:
            st.error("🚨 **Error de convergencia:** El terreno es demasiado rígido (o la zapata enorme). Ni siquiera aplicando 5000 kPa se alcanza el asiento de diseño.")
        else:
            st.success(f"✅ El solver convergió. Profundidad de influencia activa ajustada a **{zi_final:.2f} m**.")
            c1, c2, c3 = st.columns(3)
            c1.metric("🔵 Presión adm. Steinbrenner", f"{p_st:.1f} kPa")
            c2.metric("🟢 Presión adm. Elástica", f"{p_ec:.1f} kPa", help=f"Integrada con dz = {st.session_state.dz_used} m")
            
            dif_p = abs(p_st - p_ec)
            c3.metric("📊 Diferencia de carga", f"{dif_p:.1f} kPa")

            st.markdown(f"> **Interpretación:** Para que esta cimentación de {B}x{L} m no se asiente más de {s_max_mm} mm, no deberías superar estas presiones en tu diseño estructural.")

elif modo == "📋 Detalle Steinbrenner":
    st.header("📋 Detalle Método Steinbrenner")
    if not st.session_state.calculo_realizado:
        st.warning("⚠️ Calcula primero.")
    elif st.session_state.p_st is None:
        st.error("No hay datos para mostrar por error de convergencia.")
    else:
        st.markdown(f"Desglose tensional evaluado a **p = {st.session_state.p_st:.1f} kPa**.")
        df_st = st.session_state.df_st
        st.markdown("##### 📊 Asiento por estrato")
        st.dataframe(df_st[["Capa","z Techo [m]","z Base [m]","Δs [mm]"]], use_container_width=True, hide_index=True)

elif modo == "📋 Detalle Elástico":
    st.header("📋 Detalle Método Ecuación Elástica")
    if not st.session_state.calculo_realizado:
        st.warning("⚠️ Calcula primero.")
    elif st.session_state.p_ec is None:
        st.error("No hay datos para mostrar por error de convergencia.")
    else:
        st.markdown(f"Desglose tensional evaluado a **p = {st.session_state.p_ec:.1f} kPa**.")
        df_ec = st.session_state.df_ec
        st.markdown("##### ⚡ Deformación unitaria media y asiento")
        st.dataframe(df_ec[["Capa","h_ef [m]","Sub-capas","Δεz med [-]","Δs [mm]"]], use_container_width=True, hide_index=True)

elif modo == "📉 Bulbo de Presiones":
    st.header("Bulbo de Presiones")
    if not st.session_state.calculo_realizado:
        st.warning("⚠️ Calcula primero para fijar la presión.")
    elif st.session_state.p_ec is None:
        st.error("No hay datos para graficar por error de convergencia.")
    else:
        p_ref = st.session_state.p_ec
        st.markdown(f"Bulbo trazado para la presión límite del método numérico: **p = {p_ref:.1f} kPa**")
        
        z_gr = st.slider("Zoom profundidad [m]:", 1.0, espesor_total, min(espesor_total, 15.0), 0.5)
        
        z_vals = np.linspace(0.05, z_gr, 200)
        sz_v, sx_v, sy_v, sv0_v, umb20_v = [], [], [], [], []
        for z in z_vals:
            sz, sx, sy = holl_centro(p_ref, B, L, z)
            sv = sigma_v0(z, st.session_state.df_terreno, NF)
            sz_v.append(sz); sx_v.append(sx); sy_v.append(sy)
            sv0_v.append(sv); umb20_v.append(0.20*sv)

        fig, ax = plt.subplots(figsize=(9, 7))
        ax.plot(sz_v,   z_vals, label=r"Vertical $\Delta\sigma_z$",           color='red',         lw=2)
        ax.plot(sx_v,   z_vals, label=r"Horiz. Trans. $\Delta\sigma_x$",      color='blue',        ls='--')
        ax.plot(sy_v,   z_vals, label=r"Horiz. Long. $\Delta\sigma_y$",       color='purple',      ls='-.')
        ax.plot(sv0_v,  z_vals, label=r"$\sigma'_{v0}$ (tensión efect.)",     color='saddlebrown', ls=':', lw=1.5)
        ax.plot(umb20_v,z_vals, label=r"$0.20\,\sigma'_{v0}$ (criterio EC7)", color='green',       lw=2)
        
        zi = st.session_state.zi_final
        if zi <= z_gr:
            ax.axhline(y=zi, color='orange', ls='--', lw=1.5, label=f'z_i Limitante = {zi:.2f} m')
        
        ax.set_ylim(z_gr, 0); ax.set_xlim(left=0)
        ax.set_xlabel("Tensión (kPa)", fontsize=11)
        ax.set_ylabel("Profundidad z (m)", fontsize=11)
        ax.legend(loc='lower right', fontsize=9)
        ax.grid(True, linestyle=':', alpha=0.5)
        ax.spines[['top','right']].set_visible(False)
        st.pyplot(fig); plt.close(fig)

elif modo == "📖 Fundamento Teórico":
    st.header("Fundamento Teórico")
    
    st.subheader("1. Búsqueda de la Presión Admisible (Problema Inverso)")
    st.markdown(
        "El programa no calcula el asiento a partir de una carga, sino al revés. "
        "Define una función matemática basada en la diferencia entre el asiento calculado y tu objetivo: "
        r"$f(p) = s_{calculado}(p) - s_{max}$"
    )
    st.markdown(
        "Utiliza un algoritmo de búsqueda de raíces (`scipy.optimize.brentq`) que prueba iterativamente "
        "distintas presiones $p$ hasta encontrar el valor exacto donde $f(p) = 0$. En cada intento, "
        "el tamaño del bulbo de presiones se recalcula dinámicamente según el criterio seleccionado."
    )
    
    st.markdown("---")
    st.subheader("2. Modelos de Deformación (Problema Directo)")
    col_a, col_b = st.columns(2)

    with col_a:
        st.markdown("**🔵 Método 1 — Steinbrenner**")
        st.markdown("Integración analítica usando los factores de influencia $\phi_1$ y $\phi_2$:")
        st.latex(r"s(z) = \frac{p \cdot B}{E}\left[(1-\nu^2)\phi_1 - (1-\nu-2\nu^2)\phi_2\right]")
        st.latex(r"\phi_1 = \frac{1}{\pi}\left[\ln\frac{\sqrt{1+m^2+n^2}+n}{\sqrt{1+m^2}} + n\ln\frac{\sqrt{1+m^2+n^2}+1}{\sqrt{n^2+m^2}}\right]")
        st.latex(r"\phi_2 = \frac{m}{\pi}\arctan\frac{n}{m\sqrt{1+m^2+n^2}}")
        st.markdown("El asiento de cada estrato se obtiene restando el valor en su base al de su techo.")

    with col_b:
        st.markdown("**🟢 Método 2 — Ecuación Elástica**")
        st.markdown("Integración numérica de la deformación unitaria vertical en cada subcapa:")
        st.latex(r"s = \sum_{i=1}^{n}\left[\frac{h}{E}\left(\Delta\sigma_z - \nu(\Delta\sigma_x+\Delta\sigma_y)\right)\right]_i")
        st.markdown(r"Las tensiones $\Delta\sigma$ se calculan en el centro geométrico de la subcapa con la solución de Holl, superponiendo 4 cuadrantes.")
    
    st.markdown("---")
    st.subheader("3. Límite Activo de Profundidad (Bulbo)")
    st.markdown(
        "Para la integración del asiento, el programa evalúa distintos criterios y adopta el más restrictivo (el menor valor):\n"
        "1. **Criterio Seleccionado:** Puede ser el **EC7 Tensional** ($\Delta\sigma_z \leq 0.20\,\sigma'_{v0}$) o un límite geométrico explícito como **2B** o **3B**.\n"
        "2. **Límite Físico:** El espesor total de la estratigrafía introducida.\n"
        "3. **Límite Manual:** El valor forzado por el usuario como 'Profundidad de bulbo manual'."
    )