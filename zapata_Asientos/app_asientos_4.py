import streamlit as st
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import io
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# ==========================================
# CONSTANTES Y FUNCIONES BASE
# ==========================================
GAMMA_AGUA = 9.81  # kN/m³

def calcular_phi1(m, n):
    if m == 0:
        term1 = np.log(np.sqrt(1 + n**2) + n)
        term2 = n * np.log((np.sqrt(1 + n**2) + 1) / n)
    else:
        term1 = np.log((np.sqrt(1 + m**2 + n**2) + n) / np.sqrt(1 + m**2))
        term2 = n * np.log((np.sqrt(1 + m**2 + n**2) + 1) / np.sqrt(n**2 + m**2))
    return (1 / np.pi) * (term1 + term2)

def calcular_phi2(m, n):
    if m == 0: return 0.0
    return (m / np.pi) * np.arctan(n / (m * np.sqrt(1 + m**2 + n**2)))

def calcular_s_z(p, B, E, nu, z, L):
    n = L / B
    m = (2 * z) / B
    phi1 = calcular_phi1(m, n)
    phi2 = calcular_phi2(m, n)
    corchete = ((1 - nu**2) * phi1) - ((1 - nu - 2 * nu**2) * phi2)
    return (p * B / E) * corchete

def tensiones_holl_centro(p, B, L, z):
    if z <= 0.01: 
        return p, p/2, p/2 
        
    B_cuad = B / 2.0
    L_cuad = L / 2.0
    
    R1 = np.sqrt(L_cuad**2 + z**2)
    R2 = np.sqrt(B_cuad**2 + z**2)
    R3 = np.sqrt(L_cuad**2 + B_cuad**2 + z**2)
    
    term_arctan = np.arctan((B_cuad * L_cuad) / (z * R3))
    
    sigma_z_esq = (p / (2 * np.pi)) * (term_arctan + B_cuad * L_cuad * (1/R1**2 + 1/R2**2) * (z / R3))
    sigma_x_esq = (p / (2 * np.pi)) * (term_arctan - (B_cuad * L_cuad * z) / (R1**2 * R3))
    sigma_y_esq = (p / (2 * np.pi)) * (term_arctan - (B_cuad * L_cuad * z) / (R2**2 * R3))
    
    return 4 * sigma_z_esq, 4 * sigma_x_esq, 4 * sigma_y_esq

def calcular_sigma_v0(z, df_terreno, NF):
    """Calcula la tensión efectiva vertical en reposo a una profundidad z"""
    sigma_v0 = 0.0
    z_actual = 0.0
    for _, row in df_terreno.iterrows():
        espesor = float(row["Espesor (m)"])
        gamma = float(row["Peso Esp. (kN/m³)"])
        gamma_sat = float(row["Peso Esp. Sat (kN/m³)"])
        
        z_techo = z_actual
        z_base = z_actual + espesor
        
        if z <= z_techo: break
            
        z_calc = min(z, z_base)
        dz = z_calc - z_techo
        
        if NF <= z_techo:
            sigma_v0 += dz * (gamma_sat - GAMMA_AGUA)
        elif NF >= z_base:
            sigma_v0 += dz * gamma
        else:
            if z_calc <= NF:
                sigma_v0 += dz * gamma
            else:
                dz_dry = NF - z_techo
                dz_sub = z_calc - NF
                sigma_v0 += dz_dry * gamma + dz_sub * (gamma_sat - GAMMA_AGUA)
                
        z_actual = z_base
        if z <= z_base: break
            
    return sigma_v0

def calcular_z_influencia_ec7(p, B, L, df_terreno, NF):
    """Calcula dinámicamente el punto donde Δσ_z <= 0.20 * σ'_v0"""
    z = 0.1
    espesor_total = float(pd.to_numeric(df_terreno["Espesor (m)"]).sum())
    
    while z <= espesor_total:
        sz_inducida, _, _ = tensiones_holl_centro(p, B, L, z)
        sigma_v0 = calcular_sigma_v0(z, df_terreno, NF)
        
        if sigma_v0 > 0 and sz_inducida <= 0.20 * sigma_v0:
            return z
        z += 0.1
        
    return espesor_total 

# ==========================================
# FUNCIONES AUXILIARES PARA EL INFORME WORD
# ==========================================
def _fig_to_bytes(fig):
    """Convierte una figura matplotlib a bytes PNG en memoria."""
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    buf.seek(0)
    return buf

def _estilo_cabecera(celda, texto, color_fondo=(26, 58, 92)):
    """Aplica estilo de cabecera a una celda de tabla Word."""
    celda.text = texto
    run = celda.paragraphs[0].runs[0]
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    run.font.size = Pt(9)
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc = celda._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), '{:02X}{:02X}{:02X}'.format(*color_fondo))
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)

def _agregar_tabla_word(doc, df, titulo):
    """Inserta un párrafo de título y una tabla Word a partir de un DataFrame."""
    p = doc.add_paragraph(titulo)
    p.style = 'Heading 2'
    cols = list(df.columns)
    tabla = doc.add_table(rows=1 + len(df), cols=len(cols))
    tabla.style = 'Table Grid'
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Cabecera
    for j, col in enumerate(cols):
        _estilo_cabecera(tabla.cell(0, j), col)
    # Datos
    for i, (_, row) in enumerate(df.iterrows(), start=1):
        for j, val in enumerate(row):
            celda = tabla.cell(i, j)
            celda.text = str(val)
            run = celda.paragraphs[0].runs[0]
            run.font.size = Pt(8)
            if i % 2 == 0:
                from docx.oxml.ns import qn
                from docx.oxml import OxmlElement
                tc = celda._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), 'E8F0FE')
                shd.set(qn('w:val'), 'clear')
                tcPr.append(shd)
    doc.add_paragraph()

# ==========================================
# FUNCIÓN: GENERACIÓN DE INFORME WORD
# ==========================================
def generar_informe_word(B, L, p, NF, z_max_user, z_influencia, df_terreno,
                         df_basico, df_detallado, asiento_total):
    """Genera un informe Word y lo devuelve como bytes."""
    fecha = datetime.now().strftime("%d/%m/%Y %H:%M")
    espesor_total = float(pd.to_numeric(df_terreno["Espesor (m)"]).sum())
    doc = Document()

    # --- Márgenes ---
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── PORTADA ──────────────────────────────────────────────────────────
    titulo = doc.add_heading('INFORME DE CÁLCULO DE CIMENTACIONES', level=0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitulo = doc.add_paragraph('Guía de Cimentaciones con Eurocódigo 7 (EC7)')
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitulo.runs[0].font.italic = True
    fecha_p = doc.add_paragraph(f'Fecha: {fecha}')
    fecha_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # ── SECCIÓN 1: DATOS DE ENTRADA ───────────────────────────────────────
    doc.add_heading('1. Datos de Entrada', level=1)

    # Tabla parámetros geométricos
    doc.add_heading('1.1 Parámetros Geométricos y de Carga', level=2)
    params = [
        ('Ancho de zapata (B)',          f'{B:.2f} m'),
        ('Longitud de zapata (L)',        f'{L:.2f} m'),
        ('Esbeltez (n = L/B)',           f'{L/B:.3f}'),
        ('Presión neta aplicada (p)',     f'{p:.1f} kPa'),
        ('Nivel freático (NF)',           f'{NF:.1f} m'),
        ('Profundidad de corte (z_max)', f'{z_max_user:.1f} m'),
        ('Prof. influencia EC7 (z_i)',   f'{z_influencia:.2f} m'),
    ]
    tabla_geo = doc.add_table(rows=len(params), cols=2)
    tabla_geo.style = 'Table Grid'
    for i, (param, valor) in enumerate(params):
        tabla_geo.cell(i, 0).text = param
        tabla_geo.cell(i, 1).text = valor
        tabla_geo.cell(i, 0).paragraphs[0].runs[0].font.bold = True
        for j in range(2):
            tabla_geo.cell(i, j).paragraphs[0].runs[0].font.size = Pt(9)
    doc.add_paragraph()

    # Tabla estratigrafía
    _agregar_tabla_word(doc, df_terreno, '1.2 Estratigrafía del Terreno')

    doc.add_page_break()

    # ── SECCIÓN 2: RESULTADOS ─────────────────────────────────────────────
    doc.add_heading('2. Resultados del Cálculo', level=1)

    # Tabla básica resultados
    _agregar_tabla_word(doc, df_basico, '2.1 Asiento por Estrato')

    # Métrica total
    p_total = doc.add_paragraph()
    run = p_total.add_run(f'   ASIENTO TOTAL ESTIMADO: {asiento_total * 1000:.2f} mm')
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(26, 58, 92)
    doc.add_paragraph()

    # Gráfico de barras de asientos
    doc.add_heading('2.2 Gráfico de Asientos por Estrato', level=2)
    fig_bar, ax_bar = plt.subplots(figsize=(7, 3))
    capas  = df_basico['Capa'].tolist()
    deltas = df_basico['Asiento Aportado [mm]'].tolist()
    bars = ax_bar.barh(capas, deltas, color='#1a3a5c', edgecolor='white')
    ax_bar.bar_label(bars, fmt='%.2f mm', padding=3, fontsize=8)
    ax_bar.set_xlabel('Asiento Aportado (mm)')
    ax_bar.set_title('Asiento por estrato')
    ax_bar.invert_yaxis()
    ax_bar.grid(axis='x', linestyle=':', alpha=0.5)
    ax_bar.spines[['top', 'right']].set_visible(False)
    plt.tight_layout()
    doc.add_picture(_fig_to_bytes(fig_bar), width=Cm(14))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    plt.close(fig_bar)
    doc.add_page_break()

    # ── SECCIÓN 3: CÁLCULOS INTERMEDIOS ──────────────────────────────────
    doc.add_heading('3. Cálculos Intermedios', level=1)
    _agregar_tabla_word(doc, df_detallado, '3.1 Factores Geométricos φ₁ y φ₂')
    doc.add_page_break()

    # ── SECCIÓN 4: BULBO DE PRESIONES ────────────────────────────────────
    doc.add_heading('4. Bulbo de Presiones y Zona de Influencia EC7', level=1)
    z_vals = np.linspace(0.1, espesor_total, 200)
    sz_v, sx_v, sy_v, sv0_v = [], [], [], []
    for z in z_vals:
        sz, sx, sy = tensiones_holl_centro(p, B, L, z)
        sz_v.append(sz); sx_v.append(sx); sy_v.append(sy)
        sv0_v.append(calcular_sigma_v0(z, df_terreno, NF) * 0.20)

    fig_bulbo, ax_b = plt.subplots(figsize=(7, 8))
    ax_b.plot(sz_v,  z_vals, label=r'Vertical $\Delta\sigma_z$',           color='red',    lw=2)
    ax_b.plot(sx_v,  z_vals, label=r'Horiz. Transversal $\Delta\sigma_x$', color='blue',   ls='--')
    ax_b.plot(sy_v,  z_vals, label=r'Horiz. Longitudinal $\Delta\sigma_y$',color='purple', ls='-.')
    ax_b.plot(sv0_v, z_vals, label=r"$0.20\,\sigma'_{v0}$ (criterio EC7)", color='green',  lw=2)
    if z_influencia < espesor_total:
        ax_b.axhline(y=z_influencia, color='orange', ls=':', lw=1.5,
                     label=f'z_i EC7 = {z_influencia:.2f} m')
    if NF < espesor_total:
        ax_b.axhline(y=NF, color='deepskyblue', ls='-.', lw=1.2,
                     label=f'NF = {NF:.1f} m')
    ax_b.set_ylim(espesor_total, 0)
    ax_b.set_xlim(left=0)
    ax_b.set_xlabel('Tensión (kPa)', fontsize=10)
    ax_b.set_ylabel('Profundidad z (m)', fontsize=10)
    ax_b.set_title('Bulbo de presiones — Centro de la cimentación', fontsize=11, fontweight='bold')
    ax_b.legend(fontsize=8)
    ax_b.grid(True, linestyle=':', alpha=0.5)
    ax_b.spines[['top', 'right']].set_visible(False)
    plt.tight_layout()
    doc.add_picture(_fig_to_bytes(fig_bulbo), width=Cm(13))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    plt.close(fig_bulbo)

    # Footer con parámetros
    doc.add_paragraph()
    nota = doc.add_paragraph(
        f'B={B} m · L={L} m · p={p} kPa · NF={NF} m · Generado: {fecha}')
    nota.runs[0].font.size = Pt(8)
    nota.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    nota.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Guardar en buffer
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
    buf = io.BytesIO()
    fecha = datetime.now().strftime("%d/%m/%Y %H:%M")
    espesor_total = float(pd.to_numeric(df_terreno["Espesor (m)"]).sum())

    with PdfPages(buf) as pdf:

        # ── PÁGINA 1: Portada + Datos de entrada ──────────────────────────
        fig = plt.figure(figsize=(11.69, 8.27))  # A4 apaisado
        fig.patch.set_facecolor('#f0f4f8')
        gs = gridspec.GridSpec(2, 2, figure=fig, hspace=0.5, wspace=0.4,
                               left=0.07, right=0.95, top=0.88, bottom=0.08)

        # Título
        fig.text(0.5, 0.94, 'INFORME DE CÁLCULO DE CIMENTACIONES', ha='center',
                 fontsize=16, fontweight='bold', color='#1a3a5c')
        fig.text(0.5, 0.91, f'Guía de Cimentaciones EC7 — Generado: {fecha}',
                 ha='center', fontsize=9, color='gray')

        # Cuadro parámetros geométricos
        ax_geo = fig.add_subplot(gs[0, 0])
        ax_geo.axis('off')
        ax_geo.set_title('Parámetros Geométricos', fontweight='bold', loc='left', fontsize=10)
        datos_geo = [
            ['Ancho (B)', f'{B:.2f} m'],
            ['Longitud (L)', f'{L:.2f} m'],
            ['Esbeltez (L/B)', f'{L/B:.3f}'],
            ['Presión neta (p)', f'{p:.1f} kPa'],
            ['Nivel freático (NF)', f'{NF:.1f} m'],
            ['Prof. de corte (z_max)', f'{z_max_user:.1f} m'],
            ['Prof. influencia EC7', f'{z_influencia:.2f} m'],
        ]
        tbl = ax_geo.table(cellText=datos_geo, colLabels=['Parámetro', 'Valor'],
                           loc='center', cellLoc='left')
        tbl.auto_set_font_size(False)
        tbl.set_fontsize(9)
        tbl.scale(1, 1.4)
        for (r, c), cell in tbl.get_celld().items():
            if r == 0:
                cell.set_facecolor('#1a3a5c')
                cell.set_text_props(color='white', fontweight='bold')
            elif r % 2 == 0:
                cell.set_facecolor('#e8f0fe')

        # Cuadro estratigrafía
        ax_est = fig.add_subplot(gs[0, 1])
        ax_est.axis('off')
        ax_est.set_title('Estratigrafía del Terreno', fontweight='bold', loc='left', fontsize=10)
        cols_est = ['Descripción', 'Espesor\n(m)', 'E\n(kPa)', 'nu',
                    'γ\n(kN/m³)', 'γsat\n(kN/m³)']
        rows_est = []
        for _, row in df_terreno.iterrows():
            rows_est.append([
                str(row['Descripción']),
                f"{float(row['Espesor (m)']):.2f}",
                f"{float(row['E (kPa)']):.0f}",
                f"{float(row['nu']):.2f}",
                f"{float(row['Peso Esp. (kN/m³)']):.1f}",
                f"{float(row['Peso Esp. Sat (kN/m³)']):.1f}",
            ])
        tbl2 = ax_est.table(cellText=rows_est, colLabels=cols_est,
                            loc='center', cellLoc='center')
        tbl2.auto_set_font_size(False)
        tbl2.set_fontsize(8)
        tbl2.scale(1, 1.4)
        for (r, c), cell in tbl2.get_celld().items():
            if r == 0:
                cell.set_facecolor('#1a3a5c')
                cell.set_text_props(color='white', fontweight='bold')
            elif r % 2 == 0:
                cell.set_facecolor('#e8f0fe')

        # Tabla resumen resultados
        ax_res = fig.add_subplot(gs[1, 0])
        ax_res.axis('off')
        ax_res.set_title('Resultados por Capa', fontweight='bold', loc='left', fontsize=10)
        col_bas = ['Capa', 'z Techo\n(m)', 'z Base\n(m)', 'Δs\n(mm)']
        rows_bas = [[r['Capa'], f"{r['Prof. Techo [m]']:.2f}",
                     f"{r['Prof. Base [m]']:.2f}",
                     f"{r['Asiento Aportado [mm]']:.2f}"] for _, r in df_basico.iterrows()]
        rows_bas.append(['TOTAL', '', '', f"{asiento_total*1000:.2f}"])
        tbl3 = ax_res.table(cellText=rows_bas, colLabels=col_bas,
                            loc='center', cellLoc='center')
        tbl3.auto_set_font_size(False)
        tbl3.set_fontsize(9)
        tbl3.scale(1, 1.5)
        for (r, c), cell in tbl3.get_celld().items():
            if r == 0:
                cell.set_facecolor('#1a3a5c')
                cell.set_text_props(color='white', fontweight='bold')
            elif r == len(rows_bas):
                cell.set_facecolor('#ffd700')
                cell.set_text_props(fontweight='bold')
            elif r % 2 == 0:
                cell.set_facecolor('#e8f0fe')

        # Gráfico de barras de asientos
        ax_bar = fig.add_subplot(gs[1, 1])
        capas = df_basico['Capa'].tolist()
        deltas = df_basico['Asiento Aportado [mm]'].tolist()
        bars = ax_bar.barh(capas, deltas, color='#1a3a5c', edgecolor='white')
        ax_bar.bar_label(bars, fmt='%.2f mm', padding=3, fontsize=8)
        ax_bar.set_xlabel('Asiento Aportado (mm)', fontsize=8)
        ax_bar.set_title('Asiento por estrato', fontsize=10, fontweight='bold')
        ax_bar.invert_yaxis()
        ax_bar.grid(axis='x', linestyle=':', alpha=0.5)
        ax_bar.spines[['top', 'right']].set_visible(False)

        pdf.savefig(fig, bbox_inches='tight')
        plt.close(fig)

        # ── PÁGINA 2: Cálculos intermedios ────────────────────────────────
        fig2, ax2 = plt.subplots(figsize=(11.69, 8.27))
        fig2.patch.set_facecolor('#f0f4f8')
        ax2.axis('off')
        fig2.text(0.5, 0.96, 'CÁLCULOS INTERMEDIOS — Factores Geométricos',
                  ha='center', fontsize=13, fontweight='bold', color='#1a3a5c')

        cols_det = ['Capa', 'z\ntecho', 'm\ntecho', 'φ₁\ntecho', 'φ₂\ntecho',
                    's techo\n(mm)', 'z\nbase', 'm\nbase', 'φ₁\nbase', 'φ₂\nbase',
                    's base\n(mm)', 'Δs\n(mm)']
        rows_det = []
        for _, r in df_detallado.iterrows():
            rows_det.append([
                r['Capa'],
                f"{r['z_techo [m]']:.2f}", f"{r['m_techo']:.4f}",
                f"{r['φ1_techo']:.4f}",    f"{r['φ2_techo']:.4f}",
                f"{r['s_techo_teórico [mm]']:.3f}",
                f"{r['z_base [m]']:.2f}",  f"{r['m_base']:.4f}",
                f"{r['φ1_base']:.4f}",     f"{r['φ2_base']:.4f}",
                f"{r['s_base_teórico [mm]']:.3f}",
                f"{r['Δs Real [mm]']:.2f}",
            ])
        tbl4 = ax2.table(cellText=rows_det, colLabels=cols_det,
                         loc='center', cellLoc='center', bbox=[0, 0.05, 1, 0.88])
        tbl4.auto_set_font_size(False)
        tbl4.set_fontsize(8)
        tbl4.scale(1, 2.0)
        for (r, c), cell in tbl4.get_celld().items():
            if r == 0:
                cell.set_facecolor('#1a3a5c')
                cell.set_text_props(color='white', fontweight='bold')
            elif r % 2 == 0:
                cell.set_facecolor('#e8f0fe')
            if c == len(cols_det) - 1 and r > 0:
                cell.set_facecolor('#d4edda')
                cell.set_text_props(fontweight='bold')

        pdf.savefig(fig2, bbox_inches='tight')
        plt.close(fig2)

        # ── PÁGINA 3: Bulbo de presiones ──────────────────────────────────
        fig3, ax3 = plt.subplots(figsize=(8.27, 11.69))  # A4 vertical
        z_vals = np.linspace(0.1, espesor_total, 200)
        sz_vals, sx_vals, sy_vals, sv0_vals = [], [], [], []
        for z in z_vals:
            sz, sx, sy = tensiones_holl_centro(p, B, L, z)
            sz_vals.append(sz); sx_vals.append(sx); sy_vals.append(sy)
            sv0_vals.append(calcular_sigma_v0(z, df_terreno, NF) * 0.20)

        ax3.plot(sz_vals, z_vals, label=r'Vertical $\Delta\sigma_z$', color='red', lw=2)
        ax3.plot(sx_vals, z_vals, label=r'Horiz. Transversal $\Delta\sigma_x$', color='blue', ls='--')
        ax3.plot(sy_vals, z_vals, label=r'Horiz. Longitudinal $\Delta\sigma_y$', color='purple', ls='-.')
        ax3.plot(sv0_vals, z_vals, label=r"$0.20\,\sigma'_{v0}$ (criterio EC7)", color='green', lw=2)
        if z_influencia < espesor_total:
            ax3.axhline(y=z_influencia, color='orange', ls=':', lw=1.5,
                        label=f'z_i EC7 = {z_influencia:.2f} m')
        if NF < espesor_total:
            ax3.axhline(y=NF, color='deepskyblue', ls='-.', lw=1.2,
                        label=f'Nivel freático NF = {NF:.1f} m')
        ax3.set_ylim(espesor_total, 0)
        ax3.set_xlim(left=0)
        ax3.set_xlabel('Tensión (kPa)', fontsize=11)
        ax3.set_ylabel('Profundidad z (m)', fontsize=11)
        ax3.set_title('Bulbo de Presiones y Zona de Influencia EC7', fontsize=13, fontweight='bold')
        ax3.legend(fontsize=9)
        ax3.grid(True, linestyle=':', alpha=0.5)
        ax3.spines[['top', 'right']].set_visible(False)
        fig3.text(0.5, 0.02, f'B={B} m · L={L} m · p={p} kPa · NF={NF} m — {fecha}',
                  ha='center', fontsize=8, color='gray')

        pdf.savefig(fig3, bbox_inches='tight')
        plt.close(fig3)

    buf.seek(0)
    return buf

# ==========================================
# GESTIÓN DE ESTADO (SESSION STATE)
# ==========================================
def reset_calculo():
    st.session_state.calculo_realizado = False

if 'calculo_realizado' not in st.session_state:
    st.session_state.calculo_realizado = False

if 'df_terreno' not in st.session_state:
    st.session_state.df_terreno = pd.DataFrame({
        "Descripción": ["Relleno", "Arcilla", "Grava"],
        "Espesor (m)": [1.5, 3.0, 5.0],
        "E (kPa)": [10000.0, 5000.0, 40000.0],
        "nu": [0.30, 0.45, 0.25],
        "Peso Esp. (kN/m³)": [18.0, 19.0, 21.0],
        "Peso Esp. Sat (kN/m³)": [20.0, 20.0, 22.0]
    })

# ==========================================
# CONFIGURACIÓN DE PÁGINA Y BARRA LATERAL
# ==========================================
st.set_page_config(page_title="Cálculo de Cimentaciones EC7", layout="wide", page_icon="🏗️")

st.sidebar.title("Navegación")
modo_vista = st.sidebar.radio(
    "Selecciona la vista principal:",
    ("🧮 Panel de Cálculo", "🔍 Desglose de Asientos", "📉 Incremento de Tensiones", "📖 Fundamento Teórico")
)

st.sidebar.markdown("---")
st.sidebar.header("📥 Datos de Entrada")

B = st.sidebar.number_input("Ancho zapata (B) [m]", min_value=0.1, value=2.0, step=0.1, on_change=reset_calculo)
L = st.sidebar.number_input("Longitud zapata (L) [m]", min_value=0.1, value=3.0, step=0.1, on_change=reset_calculo)
p = st.sidebar.number_input("Presión neta (p) [kPa]", min_value=1.0, value=150.0, step=10.0, on_change=reset_calculo)
NF = st.sidebar.number_input("Nivel Freático [m desde cimentación]", min_value=0.0, value=10.0, step=0.5, on_change=reset_calculo)

if L < B:
    B, L = L, B

st.sidebar.markdown("---")

# --- CÁLCULO PREVIO DE Z_INFLUENCIA RECOMENDADA ---
z_influencia_teorica = calcular_z_influencia_ec7(p, B, L, st.session_state.df_terreno, NF)
espesor_total = float(pd.to_numeric(st.session_state.df_terreno["Espesor (m)"]).sum())

st.sidebar.subheader("📐 Profundidad de Cálculo")
if z_influencia_teorica >= espesor_total - 0.1:
    st.sidebar.warning(f"⚠️ **Aviso:** El bulbo profundiza más que los estratos definidos (>{espesor_total} m).")
else:
    st.sidebar.success(f"💡 **Recomendación EC7:** El bulbo se agota a **{z_influencia_teorica:.2f} m**.")

z_max_user = st.sidebar.number_input("Profundidad de corte (z_max) [m]", min_value=0.1, max_value=espesor_total, value=float(round(z_influencia_teorica, 1)), step=0.5, on_change=reset_calculo)

# --- BOTÓN DE CÁLCULO EN LA BARRA LATERAL ---
if st.sidebar.button("🚀 Calcular Asiento Total", type="primary", use_container_width=True):
    asiento_total = 0.0
    z_actual = 0.0
    resultados_basicos = []
    resultados_detallados = []
    n_factor = L / B
    
    for index, row in st.session_state.df_terreno.iterrows():
        if z_actual >= z_max_user:
            break 
            
        espesor = float(row["Espesor (m)"])
        E = float(row["E (kPa)"])
        nu = float(row["nu"])
        nombre = str(row["Descripción"])
        
        z_techo = z_actual
        z_base = min(z_actual + espesor, z_max_user) 
        
        m_techo = (2 * z_techo) / B
        m_base = (2 * z_base) / B
        
        s_techo = calcular_s_z(p, B, E, nu, z_techo, L)
        s_base = calcular_s_z(p, B, E, nu, z_base, L)
        
        delta_s = s_techo - s_base
        asiento_total += delta_s
        
        resultados_basicos.append({
            "Capa": f"{nombre} (hasta {z_base:.1f}m)" if z_base < z_actual + espesor else nombre,
            "Prof. Techo [m]": round(z_techo, 2),
            "Prof. Base [m]": round(z_base, 2),
            "Asiento Aportado [mm]": round(delta_s * 1000, 2)
        })
        
        resultados_detallados.append({
            "Capa": nombre,
            "z_techo [m]": round(z_techo, 2),
            "m_techo": round(m_techo, 4),
            "φ1_techo": round(calcular_phi1(m_techo, n_factor), 4),
            "φ2_techo": round(calcular_phi2(m_techo, n_factor), 4),
            "s_techo_teórico [mm]": round(s_techo * 1000, 3),
            "z_base [m]": round(z_base, 2),
            "m_base": round(m_base, 4),
            "φ1_base": round(calcular_phi1(m_base, n_factor), 4),
            "φ2_base": round(calcular_phi2(m_base, n_factor), 4),
            "s_base_teórico [mm]": round(s_base * 1000, 3),
            "Δs Real [mm]": round(delta_s * 1000, 2)
        })
        
        z_actual = z_base 
        
    st.session_state.df_basico = pd.DataFrame(resultados_basicos)
    st.session_state.df_detallado = pd.DataFrame(resultados_detallados)
    st.session_state.asiento_total = asiento_total
    st.session_state.calculo_realizado = True

# --- BOTÓN DE INFORME WORD ---
st.sidebar.markdown("---")
if st.session_state.calculo_realizado:
    word_bytes = generar_informe_word(
        B, L, p, NF, z_max_user, z_influencia_teorica,
        st.session_state.df_terreno,
        st.session_state.df_basico,
        st.session_state.df_detallado,
        st.session_state.asiento_total
    )
    st.sidebar.download_button(
        label="📝 Descargar Informe Word",
        data=word_bytes,
        file_name=f"informe_cimentacion_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )
else:
    st.sidebar.button("📝 Descargar Informe Word", disabled=True,
                      use_container_width=True,
                      help="Primero calcula el asiento total.")

# ==========================================
# ÁREA PRINCIPAL CENTRAL
# ==========================================
st.title("🏗️ Proyecto de Cimentaciones Superficiales")

if modo_vista == "🧮 Panel de Cálculo":
    
    st.header("1. Estratigrafía del Terreno")
    st.info("Añade los pesos específicos para calcular correctamente la Tensión Efectiva (Norma EC7).")
    
    df_actualizado = st.data_editor(st.session_state.df_terreno, num_rows="dynamic", use_container_width=True)
    
    if not df_actualizado.equals(st.session_state.df_terreno):
        st.session_state.df_terreno = df_actualizado
        st.session_state.calculo_realizado = False
        st.rerun() 
    
    st.markdown("---")
    st.header("2. Resultados del Cálculo")
        
    if st.session_state.calculo_realizado:
        col_res1, col_res2 = st.columns([2, 1])
        with col_res1:
            st.dataframe(st.session_state.df_basico, use_container_width=True)
            st.bar_chart(st.session_state.df_basico.set_index("Capa")["Asiento Aportado [mm]"])
        with col_res2:
            st.success("Cálculo actualizado.")
            st.metric(label=f"Asiento Total Estimado", value=f"{round(st.session_state.asiento_total * 1000, 2)} mm")
    else:
        st.info("👈 Verifica la recomendación de profundidad y haz clic en **Calcular Asiento Total**.")

elif modo_vista == "🔍 Desglose de Asientos":
    st.header("📋 Cálculos Intermedios y Factores Geométricos")
    
    if not st.session_state.calculo_realizado:
        st.warning("⚠️ Los datos han cambiado o no se ha calculado aún. Pulsa el botón de calcular en el panel izquierdo.")
    else:
        df_det = st.session_state.df_detallado

        # --- Sub-tabla TECHO ---

        cfg_num_4 = {"format": "%.4f", "width": "small"}
        cfg_num_3 = {"format": "%.3f", "width": "small"}
        cfg_num_2 = {"format": "%.2f", "width": "small"}

        st.markdown("##### 🔼 Valores en el Techo del estrato")
        df_techo = df_det[["Capa", "z_techo [m]", "m_techo", "φ1_techo", "φ2_techo", "s_techo_teórico [mm]"]].copy()
        st.dataframe(
            df_techo,
            use_container_width=True,
            hide_index=True,
            column_config={
                "Capa":                 st.column_config.TextColumn("Capa",     width="medium"),
                "z_techo [m]":          st.column_config.NumberColumn("z [m]",  **cfg_num_2),
                "m_techo":              st.column_config.NumberColumn("m",       **cfg_num_4),
                "φ1_techo":             st.column_config.NumberColumn("φ₁",      **cfg_num_4),
                "φ2_techo":             st.column_config.NumberColumn("φ₂",      **cfg_num_4),
                "s_techo_teórico [mm]": st.column_config.NumberColumn("s [mm]", **cfg_num_3),
            },
        )

        # --- Sub-tabla BASE ---
        st.markdown("##### 🔽 Valores en la Base del estrato")
        df_base = df_det[["Capa", "z_base [m]", "m_base", "φ1_base", "φ2_base", "s_base_teórico [mm]"]].copy()
        st.dataframe(
            df_base,
            use_container_width=True,
            hide_index=True,
            column_config={
                "Capa":                st.column_config.TextColumn("Capa",     width="medium"),
                "z_base [m]":          st.column_config.NumberColumn("z [m]",  **cfg_num_2),
                "m_base":              st.column_config.NumberColumn("m",       **cfg_num_4),
                "φ1_base":             st.column_config.NumberColumn("φ₁",      **cfg_num_4),
                "φ2_base":             st.column_config.NumberColumn("φ₂",      **cfg_num_4),
                "s_base_teórico [mm]": st.column_config.NumberColumn("s [mm]", **cfg_num_3),
            },
        )

        # --- Resumen Δs ---
        st.markdown("##### 📊 Asiento aportado por estrato")
        df_delta = df_det[["Capa", "Δs Real [mm]"]].copy()
        st.dataframe(
            df_delta,
            use_container_width=True,
            hide_index=True,
            column_config={
                "Capa":        st.column_config.TextColumn("Capa",             width="medium"),
                "Δs Real [mm]": st.column_config.NumberColumn("Δs Real [mm]", width="small", format="%.2f"),
            },
        )

        st.info(f"**Factor de esbeltez:** n = L/B = **{L/B:.4f}**")

elif modo_vista == "📉 Incremento de Tensiones":
    st.header("Profundidad de Influencia y Bulbo de Presiones")
    
    fig, ax = plt.subplots(figsize=(8, 6))
    z_vals = np.linspace(0.1, espesor_total, 100)
    
    # Preparamos las listas para las 4 curvas
    sigma_z_vals, sigma_x_vals, sigma_y_vals, sigma_v0_vals = [], [], [], []
    
    for z in z_vals:
        sz, sx, sy = tensiones_holl_centro(p, B, L, z)
        sigma_z_vals.append(sz)
        sigma_x_vals.append(sx)  # Recuperado
        sigma_y_vals.append(sy)  # Recuperado
        sigma_v0_vals.append(calcular_sigma_v0(z, st.session_state.df_terreno, NF) * 0.20)
        
    # Dibujamos todas las curvas
    ax.plot(sigma_z_vals, z_vals, label=r'Vertical ($\Delta\sigma_z$)', color='red', linewidth=2)
    ax.plot(sigma_x_vals, z_vals, label=r'Horiz. Transversal ($\Delta\sigma_x$)', color='blue', linestyle='--')
    ax.plot(sigma_y_vals, z_vals, label=r'Horiz. Longitudinal ($\Delta\sigma_y$)', color='purple', linestyle='-.')
    ax.plot(sigma_v0_vals, z_vals, label=r"20% Tensión Efectiva ($0.2\sigma'_{v0}$)", color='green', linewidth=2)
    
    # Marcador de profundidad teórica
    if z_influencia_teorica < espesor_total:
        ax.axhline(y=z_influencia_teorica, color='gray', linestyle=':', label=f'Corte Teórico ({z_influencia_teorica:.2f}m)')
    
    ax.set_ylim(espesor_total, 0)
    ax.set_xlim(0, max(p, max(sigma_v0_vals)))
    ax.set_xlabel("Tensión (kPa)")
    ax.set_ylabel("Profundidad z (m)")
    ax.grid(True, alpha=0.5)
    ax.legend()
    st.pyplot(fig)

elif modo_vista == "📖 Fundamento Teórico":
    st.header("Metodología de Cálculo")
    
    st.subheader("1. Cálculo de Asientos (Steinbrenner)")
    st.markdown("Basado en el Apartado 5.2.8.3 de la Guía EC7. El asiento a profundidad $z$ en medio semi-infinito es:")
    st.latex(r"s(z) = \frac{p \cdot B}{E} \left[ (1 - \nu^2) \phi_1 - (1 - \nu - 2\nu^2) \phi_2 \right]")
    st.markdown(r"Donde $\phi_1$ y $\phi_2$ dependen de $n = L/B$ y $m = 2z/B$:")
    st.latex(r"\phi_1 = \frac{1}{\pi} \left[ \ln \left( \frac{\sqrt{1+m^2+n^2}+n}{\sqrt{1+m^2}} \right) + n \ln \left( \frac{\sqrt{1+m^2+n^2}+1}{\sqrt{n^2+m^2}} \right) \right]")
    st.latex(r"\phi_2 = \frac{m}{\pi} \arctan \left( \frac{n}{m \sqrt{1+m^2+n^2}} \right)")
    st.markdown("Para estratos múltiples, se aplica superposición (asiento techo - asiento base con mismos parámetros):")
    st.latex(r"\Delta s_i = s(z_i) - s(z_{i+1}) \quad \rightarrow \quad s_{total} = \sum \Delta s_i")
    
    st.subheader("2. Distribución de Tensiones (Holl)")
    st.markdown(r"Incrementos de tensión bajo la **esquina** de un rectángulo $B \times L$ a profundidad $z$. En esta app, se evalúa para $B/2 \times L/2$ y se multiplica por 4 (centro de la zapata).")
    st.latex(r"\sigma_z = \frac{p}{2\pi} \left[ \arctan\left(\frac{BL}{zR_3}\right) + \frac{BLz}{R_3} \left(\frac{1}{R_1^2} + \frac{1}{R_2^2}\right) \right]")
    st.latex(r"\sigma_x = \frac{p}{2\pi} \left[ \arctan\left(\frac{BL}{zR_3}\right) - \frac{BLz}{R_1^2 R_3} \right]")
    st.latex(r"\sigma_y = \frac{p}{2\pi} \left[ \arctan\left(\frac{BL}{zR_3}\right) - \frac{BLz}{R_2^2 R_3} \right]")
    st.markdown(r"Siendo $R_1 = \sqrt{L^2+z^2}$, $R_2 = \sqrt{B^2+z^2}$ y $R_3 = \sqrt{L^2+B^2+z^2}$.")

    st.subheader("3. Profundidad de Influencia (Criterio EC7)")
    st.markdown("""
    La profundidad de influencia $z_i$ es aquella a partir de la cual el incremento de tensión 
    generado por la cimentación es despreciable frente al estado tensional natural del terreno. 
    El **criterio del EC7** establece:
    """)
    st.latex(r"\Delta\sigma_z(z_i) \leq 0.20 \cdot \sigma'_{v0}(z_i)")
    st.markdown(r"""
    Donde $\sigma'_{v0}$ es la **tensión efectiva vertical geoestática**, calculada en base a los pesos 
    específicos ($\gamma$ y $\gamma_{sat}$) y la posición del Nivel Freático.
    """)