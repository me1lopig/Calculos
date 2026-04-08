import streamlit as st
import numpy as np
import pandas as pd
import math
import io
import plotly.graph_objects as go
import plotly.express as px
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date

st.set_page_config(page_title="Cálculo de Pilotes - CTE", layout="wide", page_icon="🏗️")

st.title("🏗️ Diseño de Pilotes (CTE DB-SE-C)")
st.markdown("Cálculo Analítico y Estimación de Asientos según el Código Técnico de la Edificación.")

# ══════════════════════════════════════════════════════════════════════════
# INICIALIZACIÓN DE LA TABLA BASE 
# ══════════════════════════════════════════════════════════════════════════
if 'df_base' not in st.session_state:
    st.session_state.df_base = pd.DataFrame({
        "Estrato": ["UG-01", "UG-02", "UG-03", "UG-04"],
        "Espesor (m)": [2.0, 5.0, 8.0, 10.0],
        "Gamma Seco (kN/m3)": [18.0, 17.0, 19.0, 20.0],
        "Gamma Sat. (kN/m3)": [20.0, 18.0, 21.0, 21.0],
        "Condición": ["Largo Plazo", "Corto Plazo", "Largo Plazo", "Corto Plazo"],
        "c / cu (kPa)": [0.0, 100.0, 0.0, 150.0],
        "phi (grados)": [28.0, 0.0, 35.0, 0.0]
    })

if 'calculado' not in st.session_state: st.session_state.calculado = False
if 'word_buffer' not in st.session_state: st.session_state.word_buffer = None
if 'excel_buffer' not in st.session_state: st.session_state.excel_buffer = None
if 'fig_final_guardada' not in st.session_state: st.session_state.fig_final_guardada = None

# ══════════════════════════════════════════════════════════════════════════
# INTERFAZ PRINCIPAL DE PESTAÑAS 
# ══════════════════════════════════════════════════════════════════════════
tab_datos, tab_tensiones, tab_matriz_punta, tab_matriz_fuste, tab_matriz_total, tab_matriz_tope, tab_auditoria, tab_asientos, tab_formulacion = st.tabs([
    "📋 1. Estratigrafía", "🌊 2. Tensiones", "🔻 3. Punta", "🟫 4. Fuste", "🌍 5. Total", "🛑 6. Tope Estruct.", "🔍 7. Auditoría", "📉 8. Asientos", "📖 9. Formulación CTE"
])

with tab_datos:
    st.subheader("Definición de Unidades")
    
    # --- TABLA BLINDADA CONTRA NEGATIVOS E INCONGRUENCIAS ---
    df_edit = st.data_editor(
        st.session_state.df_base, 
        key="tabla_estratos",
        num_rows="dynamic", 
        use_container_width=True,
        column_config={
            "Condición": st.column_config.SelectboxColumn(
                options=["Corto Plazo", "Largo Plazo"], 
                required=True
            ),
            "c / cu (kPa)": st.column_config.NumberColumn(
                "c / cu (kPa)", min_value=0.0, step=1.0
            ),
            "phi (grados)": st.column_config.NumberColumn(
                "phi (grados)", min_value=0.0, max_value=60.0, step=1.0
            ),
            "Espesor (m)": st.column_config.NumberColumn(
                "Espesor (m)", min_value=0.1, step=0.5
            ),
            "Gamma Seco (kN/m3)": st.column_config.NumberColumn(
                "Gamma Seco (kN/m3)", min_value=0.0, step=0.5
            ),
            "Gamma Sat. (kN/m3)": st.column_config.NumberColumn(
                "Gamma Sat. (kN/m3)", min_value=0.0, step=0.5
            )
        }
    )
    
    # --- CORRECCIÓN BIDIRECCIONAL C/PHI ---
    mascara_corto = df_edit["Condición"] == "Corto Plazo"
    cambio_corto = (df_edit.loc[mascara_corto, "phi (grados)"] != 0.0).any()
    
    mascara_largo = df_edit["Condición"] == "Largo Plazo"
    cambio_largo = (df_edit.loc[mascara_largo, "c / cu (kPa)"] != 0.0).any()
    
    if cambio_corto or cambio_largo:
        if cambio_corto:
            df_edit.loc[mascara_corto, "phi (grados)"] = 0.0
        if cambio_largo:
            df_edit.loc[mascara_largo, "c / cu (kPa)"] = 0.0
            
        st.session_state.df_base = df_edit.copy() 
        st.toast("💡 Parámetros de resistencia (c/φ) ajustados automáticamente según la condición", icon="⚙️")
        st.rerun()
        
    espesores_invalidos = (df_edit["Espesor (m)"] <= 0).any()
    if espesores_invalidos:
        st.error("⛔ **Error de datos:** Existen estratos con espesor ≤ 0 m.")
        st.session_state.calculado = False

    z_max_total = df_edit["Espesor (m)"].sum()
    st.info(f"Profundidad máxima actual del sondeo: **{z_max_total:.2f} m**.")

# ══════════════════════════════════════════════════════════════════════════
# BARRA LATERAL: CONFIGURACIÓN CTE
# ══════════════════════════════════════════════════════════════════════════
st.sidebar.header("💧 Condiciones Generales")
zw = st.sidebar.number_input("Prof. Nivel Freático, zw (m)", min_value=0.0, value=3.0, step=0.5)

FS = 3.00
sit_str = "CTE (Método Analítico FS=3.0)"
st.sidebar.info(f"Factor de Seguridad Global: **{FS}**\n(S/ Tabla 5.2 DB-SE-C)")

st.sidebar.header("🏗️ Ejecución y Material (CTE)")
tipo_ejecucion = st.sidebar.selectbox("Método de Ejecución", ["Perforado", "Hincado"])

if tipo_ejecucion == "Perforado":
    material_pilote = "Hormigón in situ"
    st.sidebar.markdown(f"**Material:** {material_pilote}")
    metodo_perf = st.sidebar.selectbox("Técnica (Tabla 5.1)", ["En seco", "Lodos", "Entubado", "Barrena sin control", "Barrena con control"])
    apoyo_roca = st.sidebar.checkbox("¿Apoyo en Roca Dura?")
    control_integridad = st.sidebar.checkbox("¿Control de Integridad Adecuado (+25%)?")
    
    if metodo_perf == "Entubado": 
        tope_base = 5.0 if apoyo_roca else 6.0
    elif metodo_perf in ["En seco", "Lodos"]: 
        tope_base = 4.0 if apoyo_roca else 5.0
    elif metodo_perf == "Barrena sin control": 
        tope_base = 3.5
    elif metodo_perf == "Barrena con control": 
        tope_base = 4.0
        
    tope_calc = tope_base * 1.25 if control_integridad else tope_base

else: # HINCADO
    material_pilote = st.sidebar.selectbox("Material del Pilote", ["Hormigón prefabricado", "Acero", "Madera"])
    
    if material_pilote == "Hormigón prefabricado":
        fck = st.sidebar.number_input("Resistencia Hormigón fck (MPa):", min_value=20.0, value=50.0, step=5.0)
        tope_calc = 0.30 * fck
    elif material_pilote == "Acero":
        fyk = st.sidebar.number_input("Límite Elástico Acero fyk (MPa):", min_value=200.0, value=355.0, step=5.0)
        tope_calc = 0.30 * fyk
    else:
        tope_calc = 5.0 # Madera
        
sigma_tope_mpa = st.sidebar.number_input("🛑 Tope Estructural Adoptado (MPa):", min_value=0.1, value=float(tope_calc), step=0.1)

st.sidebar.header("📐 Matriz Geométrica del Pilote")
col_d1, col_d2 = st.sidebar.columns(2)
D_min = col_d1.number_input("Ø min (m)", value=0.6, min_value=0.3, step=0.1)
D_max = col_d2.number_input("Ø max (m)", value=1.5, min_value=0.3, step=0.1)
D_step = st.sidebar.number_input("Paso Ø (m)", value=0.3, min_value=0.1)

col_l1, col_l2 = st.sidebar.columns(2)
L_min_default = min(10.0, float(z_max_total)) 
L_min = col_l1.number_input("L min (m)", value=L_min_default, min_value=1.0, max_value=float(z_max_total), step=1.0)
L_max_default = min(20.0, float(z_max_total))
L_max = col_l2.number_input("L max (m)", value=L_max_default, min_value=float(L_min), max_value=float(z_max_total), step=1.0)
L_step = st.sidebar.number_input("Paso L (m)", value=2.5, min_value=0.5)

st.sidebar.markdown("---")
if st.sidebar.button("⚙️ Calcular Pilotes", type="primary", use_container_width=True, disabled=espesores_invalidos):
    st.session_state.calculado = True
    st.session_state.word_buffer = None
    st.session_state.excel_buffer = None

# ══════════════════════════════════════════════════════════════════════════
# MOTORES DE CÁLCULO CTE DB-SE-C
# ══════════════════════════════════════════════════════════════════════════
def calcular_perfil_tensiones(df, zw, z_max):
    z_max_extendido = z_max + 10.0 
    z_vals = np.arange(0, z_max_extendido + 0.1, 0.1)
    sigma_v, u, sigma_v_eff = [], [], []
    sv_acum = 0.0
    for z in z_vals:
        z_acum, gamma_actual = 0.0, 18.0
        for idx, row in df.iterrows():
            is_last = (idx == len(df) - 1)
            if z <= z_acum + row["Espesor (m)"] + 1e-5 or is_last:
                gamma_actual = row["Gamma Seco (kN/m3)"] if z <= zw + 1e-5 else row["Gamma Sat. (kN/m3)"]
                break
            z_acum += row["Espesor (m)"]
            
        sv_acum = sv_acum + gamma_actual * 0.1 if z > 0 else 0.0
        u_val = max(0, (z - zw) * 9.81)
        sigma_v.append(sv_acum)
        u.append(u_val)
        sigma_v_eff.append(sv_acum - u_val)
        
    return z_vals, np.array(sigma_v), np.array(u), np.array(sigma_v_eff)

def obtener_tension_a_profundidad(z_target, z_vals, sigma_array):
    return np.interp(z_target, z_vals, sigma_array)

def calcular_pilote_cte(D, L, df, zw, fS_val, sigma_tope_mpa, ejecucion, material):
    z_max = df["Espesor (m)"].sum()
    if L > z_max: return None 
    
    z_vals, sig_v, u, sig_v_eff = calcular_perfil_tensiones(df, zw, z_max)
    
    f_p = 3.0 if ejecucion == "Hincado" else 2.5
    K_f = 1.0 if ejecucion == "Hincado" else 0.75
    
    if material == "Hormigón in situ": f_mat = 1.0
    elif material == "Hormigón prefabricado": f_mat = 0.9
    elif material == "Acero": f_mat = 0.8
    else: f_mat = 1.0
    
    # --- PUNTA CTE ---
    sig_v_eff_punta = obtener_tension_a_profundidad(L, z_vals, sig_v_eff)
    
    # --- Bulbo CTE (6D Arriba / 3D Abajo según 5.3.4.1.2) ---
    z_sup_bulbo = max(0.0, L - (6 * D))
    z_inf_bulbo = L + (3 * D)
    espesor_bulbo = z_inf_bulbo - z_sup_bulbo
    
    qp_eq_acumulado = 0.0
    z_acum = 0.0
    detalle_bulbo_grafico = [] 
    
    for idx, row in df.iterrows():
        is_last = (idx == len(df) - 1)
        h_estrato = row["Espesor (m)"]
        z_bot_estrato = max(z_acum + h_estrato, z_inf_bulbo) if is_last else z_acum + h_estrato
            
        overlap_top = max(z_sup_bulbo, z_acum)
        overlap_bot = min(z_inf_bulbo, z_bot_estrato)
        overlap_h = max(0.0, overlap_bot - overlap_top)
        
        if overlap_h > 0 and espesor_bulbo > 0:
            cond_i = row["Condición"]
            c_i = row["c / cu (kPa)"]
            phi_i = row["phi (grados)"]
            
            if "Corto Plazo" in cond_i:
                qp_i = 9.0 * c_i
            else:
                # LARGO PLAZO: SE DESPRECIA LA COHESIÓN
                phi_rad = math.radians(phi_i)
                N_q = ((1 + math.sin(phi_rad))/(1 - math.sin(phi_rad))) * math.exp(math.pi * math.tan(phi_rad)) if phi_i > 0 else 1.0
                qp_i = f_p * sig_v_eff_punta * N_q
                qp_i = min(qp_i, 20000.0) # Límite CTE estricto: 20 MPa
            
            peso = overlap_h / espesor_bulbo
            qp_eq_acumulado += qp_i * peso
            nombre_estrato = f"{row['Estrato']} (Prolongado)" if is_last and overlap_bot > (z_acum + h_estrato) else row['Estrato']
            
            detalle_bulbo_grafico.append({
                "Estrato": nombre_estrato, "Espesor en bulbo (m)": overlap_h,
                "Participación (%)": peso * 100.0, "q_p individual (kPa)": qp_i
            })
            
        z_acum += h_estrato
        
    qp = qp_eq_acumulado

    # --- COMPROBACIÓN PUNZONAMIENTO CAPAS BLANDAS INFERIORES (CTE 5.3.4.1.2 - Pto 5) ---
    z_acum_inf = 0.0
    aviso_punzonamiento = False
    for idx, row in df.iterrows():
        h_estrato_inf = row["Espesor (m)"]
        z_top_estrato_inf = z_acum_inf  # Cota del techo de la capa
        z_bot_estrato_inf = z_acum_inf + h_estrato_inf
        
        # CORRECCIÓN LÓGICA: El techo del estrato blando debe estar a una cota igual o inferior a la punta
        if z_top_estrato_inf >= L and "Corto Plazo" in row["Condición"]:
            c_u_inf = row["c / cu (kPa)"]
            H_dist = z_top_estrato_inf - L  # Distancia real desde la punta hasta el inicio de la capa blanda
            qp_limite_blando = 6.0 * (1.0 + (H_dist / D)) * c_u_inf
            
            if qp_limite_blando < qp:
                qp = qp_limite_blando
                aviso_punzonamiento = True
                detalle_bulbo_grafico.append({
                    "Estrato": f"⚠️ LÍMITE PUNZONAMIENTO ({row['Estrato']})", 
                    "Espesor en bulbo (m)": 0.0,
                    "Participación (%)": 0.0, 
                    "q_p individual (kPa)": qp_limite_blando
                })
        z_acum_inf += h_estrato_inf

    Area_pilote = (math.pi * D**2) / 4.0
    Q_punta = qp * Area_pilote
    
    auditoria_punta = {
        "Profundidad Punta (m)": L, "σ'_v efectiva base (kPa)": sig_v_eff_punta,
        "Factor f_p": f_p, "Resist. Unitaria q_p (kPa)": qp, "Fuerza Total Punta (kN)": Q_punta
    }
    
    # --- FUSTE CTE ---
    Q_fuste, z_top, Perimetro = 0.0, 0.0, math.pi * D
    auditoria_fuste = [] 
    
    for _, row in df.iterrows():
        z_bot = z_top + row["Espesor (m)"]
        if z_top >= L: break
        
        z_end_tramo = min(z_bot, L)
        if z_end_tramo <= z_top:
            z_top = z_bot; continue
            
        puntos_corte = [z_top]
        if z_top < zw < z_end_tramo: puntos_corte.append(zw)
        puntos_corte.append(z_end_tramo)
        
        for i in range(len(puntos_corte) - 1):
            z_sub_top = puntos_corte[i]
            z_sub_bot = puntos_corte[i+1]
            L_sub = z_sub_bot - z_sub_top
            z_mid = z_sub_top + (L_sub / 2.0)
            sig_v_eff_mid = obtener_tension_a_profundidad(z_mid, z_vals, sig_v_eff)
            
            if "Corto Plazo" in row["Condición"]:
                tau_f = row["c / cu (kPa)"] * (100.0 / (100.0 + row["c / cu (kPa)"]))
                if material == "Acero": tau_f *= 0.8
            else:
                tau_f = sig_v_eff_mid * K_f * f_mat * math.tan(math.radians(row["phi (grados)"]))
                tau_f = min(tau_f, 120.0) # Límite CTE 120 kPa
                
            Q_tramo = tau_f * Perimetro * L_sub
            Q_fuste += Q_tramo
            sufijo = " (Seco)" if z_sub_bot <= zw and len(puntos_corte)>2 else " (Sat.)" if len(puntos_corte)>2 else ""
            
            auditoria_fuste.append({
                "Estrato": row["Estrato"] + sufijo, "Cotas (m)": f"{z_sub_top:.1f} a {z_sub_bot:.1f}",
                "Long. fuste (m)": L_sub, "σ'_v media (kPa)": sig_v_eff_mid,
                "Resist. Unitaria τ_f (kPa)": tau_f, "Fuerza Tramo (kN)": Q_tramo
            })
        z_top = z_bot
        
    Q_total_geo = Q_punta + Q_fuste
    Q_adm_geo = Q_total_geo / fS_val
    
    Q_tope_est = Area_pilote * (sigma_tope_mpa * 1000.0)
    Q_final_diseno = min(Q_adm_geo, Q_tope_est)
    control = "ESTRUCTURAL" if Q_tope_est < Q_adm_geo else "GEOTÉCNICO"

    return {
        "D": D, "L": L, "Q_punta (kN)": Q_punta, "Q_fuste (kN)": Q_fuste, 
        "Q_adm_geo (kN)": Q_adm_geo, "Q_tope_est (kN)": Q_tope_est, "Q_final (kN)": Q_final_diseno,
        "Control": control, "auditoria_punta": auditoria_punta, "auditoria_fuste": auditoria_fuste,
        "auditoria_bulbo": detalle_bulbo_grafico, "z_sup_bulbo": z_sup_bulbo, "z_inf_bulbo": z_inf_bulbo,
        "aviso_punz": aviso_punzonamiento
    }

# ══════════════════════════════════════════════════════════════════════════
# GENERADORES DE INFORMES
# ══════════════════════════════════════════════════════════════════════════
def generar_word_cte(df_estratos, df_pivot_geo, df_pivot_final, fS_val, situacion, zw_val, sigma_tope, datos_unitarios_df, fig_final, ejec, mat):
    doc = Document()
    estilo_tabla = 'Light Grid Accent 1'
    
    doc.add_paragraph('\n\n\n\n')
    title = doc.add_paragraph('ANEJO DE CÁLCULO: CIMENTACIONES PROFUNDAS')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs: run.font.size, run.font.bold = Pt(24), True
    subtitle = doc.add_paragraph('Diseño Analítico de Pilotes según CTE DB-SE-C')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs: run.font.size = Pt(16)
    doc.add_page_break()
    
    doc.add_heading('1. Bases de Cálculo y Parámetros Iniciales', level=1)
    p_bases = doc.add_paragraph()
    p_bases.add_run(f'• Normativa: ').bold = True
    p_bases.add_run('Código Técnico de la Edificación (CTE DB-SE-C)\n')
    p_bases.add_run(f'• Método de Ejecución: ').bold = True
    p_bases.add_run(f'Pilote {ejec} de {mat}\n')
    p_bases.add_run(f'• Factor de Seguridad Global: ').bold = True
    p_bases.add_run(f'{fS_val:.2f}\n')
    p_bases.add_run(f'• Nivel Freático: ').bold = True
    p_bases.add_run(f'Considerado a una profundidad Z = {zw_val:.2f} m\n')
    p_bases.add_run(f'• Tope Estructural Adoptado: ').bold = True
    p_bases.add_run(f'{sigma_tope:.2f} MPa\n')

    doc.add_heading('2. Metodología y Formulación Analítica (CTE)', level=1)
    p_metodo = doc.add_paragraph()
    p_metodo.add_run('Resistencia por Punta (qp):\n').bold = True
    p_metodo.add_run('• Corto Plazo (Arcillas): qp = 9 · cu\n')
    p_metodo.add_run('• Largo Plazo (Arenas): qp = fp · σ\'v,p · Nq (Límite normativo 20 MPa, se desprecia cohesión c\')\n')
    p_metodo.add_run('El valor final de la punta se obtiene promediando en un bulbo de rotura de 6D por encima y 3D por debajo. Además, se verifica el riesgo de punzonamiento en capas arcillosas subyacentes.\n\n')
    p_metodo.add_run('Resistencia por Fuste (τf):\n').bold = True
    p_metodo.add_run('• Corto Plazo: Ley hiperbólica en función de cu\n')
    p_metodo.add_run('• Largo Plazo: τf = σ\'v · Kf · f · tan(δ) (Límite normativo 120 kPa, se desprecia cohesión c\')\n\n')

    doc.add_heading('3. Perfil Estratigráfico', level=1)
    tabla_estratos = doc.add_table(rows=1, cols=len(df_estratos.columns))
    tabla_estratos.style = estilo_tabla
    hdr_cells = tabla_estratos.rows[0].cells
    for i, column in enumerate(df_estratos.columns): hdr_cells[i].text = str(column)
    for index, row in df_estratos.iterrows():
        row_cells = tabla_estratos.add_row().cells
        for i, value in enumerate(row): row_cells[i].text = str(value)
                 
    doc.add_page_break()

    doc.add_heading(f'4. Matriz Admisible TOTAL del Terreno (kN)', level=1)
    tabla_geo = doc.add_table(rows=1, cols=len(df_pivot_geo.columns) + 1)
    tabla_geo.style = estilo_tabla
    hdr_geo = tabla_geo.rows[0].cells
    hdr_geo[0].text = "L / Ø"
    for i, col_name in enumerate(df_pivot_geo.columns): hdr_geo[i+1].text = str(col_name).replace('\n', ' ')
    for index, row in df_pivot_geo.iterrows():
        row_cells = tabla_geo.add_row().cells
        row_cells[0].text = str(index)
        for i, val in enumerate(row): row_cells[i+1].text = f"{val:.0f}"

    doc.add_paragraph('\n')
    doc.add_heading(f'5. Matriz de Diseño FINAL LIMITADA (kN)', level=1)
    tabla_fin = doc.add_table(rows=1, cols=len(df_pivot_final.columns) + 1)
    tabla_fin.style = estilo_tabla
    hdr_fin = tabla_fin.rows[0].cells
    hdr_fin[0].text = "L / Ø"
    for i, col_name in enumerate(df_pivot_final.columns): hdr_fin[i+1].text = str(col_name).replace('\n', ' ')
    for index, row in df_pivot_final.iterrows():
        row_cells = tabla_fin.add_row().cells
        row_cells[0].text = str(index)
        for i, val in enumerate(row): row_cells[i+1].text = str(val)

    if fig_final is not None:
        doc.add_paragraph('\n')
        try:
            img_bytes = fig_final.to_image(format="png", width=800, height=500)
            doc.add_picture(io.BytesIO(img_bytes), width=Inches(6.5))
        except Exception: pass

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def generar_excel_cte(df_estratos, df_unitarios, df_punta, df_fuste, df_geo, df_final, zw_val, fs_val, sit, tope):
    buffer = io.BytesIO()
    with pd.ExcelWriter(buffer, engine='xlsxwriter') as writer:
        df_params = pd.DataFrame({"Parámetro": ["Nivel Freático (m)", "Factor Seguridad", "Situación", "Tope Estructural (MPa)"], "Valor": [zw_val, fs_val, sit, tope]})
        df_params.to_excel(writer, sheet_name='1. Datos Partida', index=False)
        df_estratos.to_excel(writer, sheet_name='1. Datos Partida', index=False, startrow=6)
        df_unitarios.to_excel(writer, sheet_name='2. Unitarias', index=False)
        df_punta.to_excel(writer, sheet_name='3. Matriz Punta')
        df_fuste.to_excel(writer, sheet_name='4. Matriz Fuste')
        df_geo.to_excel(writer, sheet_name='5. Matriz Terreno')
        
        workbook = writer.book
        worksheet = workbook.add_worksheet('6. Matriz Final Limitada')
        writer.sheets['6. Matriz Final Limitada'] = worksheet
        df_final.to_excel(writer, sheet_name='6. Matriz Final Limitada')
        
        formato_verde = workbook.add_format({'bg_color': '#ccffcc'})
        formato_rojo = workbook.add_format({'bg_color': '#ffcccc'})
        formato_naranja = workbook.add_format({'bg_color': '#ffe6cc', 'font_color': '#cc5200', 'bold': True})
        formato_gris = workbook.add_format({'bg_color': '#d3d3d3', 'bold': True, 'bottom': 2})
        
        num_rows, num_cols = len(df_final.index) + 1, len(df_final.columns)
        worksheet.set_column(0, 0, 18) 
        worksheet.set_column(1, num_cols, 15) 
        
        worksheet.conditional_format(2, 1, num_rows, num_cols, {'type': 'text', 'criteria': 'containing', 'value': '[EST]', 'format': formato_rojo})
        worksheet.conditional_format(2, 1, num_rows, num_cols, {'type': 'text', 'criteria': 'containing', 'value': '[PUNZ]', 'format': formato_naranja})
        worksheet.conditional_format(2, 1, num_rows, num_cols, {'type': 'text', 'criteria': 'not containing', 'value': '[', 'format': formato_verde})
        worksheet.set_row(1, None, formato_gris)
    buffer.seek(0)
    return buffer

# ══════════════════════════════════════════════════════════════════════════
# CÁLCULOS PRINCIPALES Y RENDERIZADO DE RESULTADOS
# ══════════════════════════════════════════════════════════════════════════
if st.session_state.calculado:
    z_vals, sig_v, u, sig_v_eff = calcular_perfil_tensiones(df_edit, zw, z_max_total)
    datos_unitarios = []
    
    f_p = 3.0 if tipo_ejecucion == "Hincado" else 2.5
    K_f = 1.0 if tipo_ejecucion == "Hincado" else 0.75
    if material_pilote == "Hormigón in situ": f_mat = 1.0
    elif material_pilote == "Hormigón prefabricado": f_mat = 0.9
    elif material_pilote == "Acero": f_mat = 0.8
    else: f_mat = 1.0
    
    with tab_datos:
        st.markdown("---")
        st.subheader("📊 Resistencias Unitarias CTE")
        z_top_loop = 0.0
        for idx, row_est in df_edit.iterrows():
            espesor = row_est["Espesor (m)"]
            if espesor <= 0: continue
            z_bot_loop = z_top_loop + espesor
            puntos_corte = [z_top_loop]
            if z_top_loop < zw < z_bot_loop: puntos_corte.append(zw)
            puntos_corte.append(z_bot_loop)
            
            for i in range(len(puntos_corte) - 1):
                z_sub_top = puntos_corte[i]
                z_sub_bot = puntos_corte[i+1]
                z_mid_loop = z_sub_top + (z_sub_bot - z_sub_top) / 2.0
                sig_v_eff_mid_loop = obtener_tension_a_profundidad(z_mid_loop, z_vals, sig_v_eff)
                sig_v_eff_base = obtener_tension_a_profundidad(z_sub_bot, z_vals, sig_v_eff)
                
                cond_loop = row_est["Condición"]
                c_loop = row_est["c / cu (kPa)"]
                phi_loop = row_est["phi (grados)"]
                
                if "Corto Plazo" in cond_loop:
                    tau_f_media = c_loop * (100.0 / (100.0 + c_loop))
                    if material_pilote == "Acero": tau_f_media *= 0.8
                    qp_base = 9.0 * c_loop
                else:
                    tau_f_media = min(sig_v_eff_mid_loop * K_f * f_mat * math.tan(math.radians(phi_loop)), 120.0)
                    phi_rad = math.radians(phi_loop)
                    N_q = ((1 + math.sin(phi_rad))/(1 - math.sin(phi_rad))) * math.exp(math.pi * math.tan(phi_rad)) if phi_loop > 0 else 1.0
                    qp_base = min(f_p * sig_v_eff_base * N_q, 20000.0)
                
                sufijo = " (Seco)" if z_sub_bot <= zw and len(puntos_corte)>2 else " (Sat.)" if len(puntos_corte)>2 else ""
                datos_unitarios.append({
                    "Estrato": row_est["Estrato"] + sufijo, "Profundidad (m)": f"de {z_sub_top:.1f} a {z_sub_bot:.1f}",
                    "Fricción media, τ_f (kPa)": round(tau_f_media, 1), "Punta base, q_p (kPa)": round(qp_base, 1)
                })
            z_top_loop = z_bot_loop
            
        df_unitarios = pd.DataFrame(datos_unitarios)
        st.dataframe(df_unitarios.style.format({"Fricción media, τ_f (kPa)": "{:.1f}", "Punta base, q_p (kPa)": "{:.1f}"}).hide(axis="index"), use_container_width=True)

    fig_tens = go.Figure()
    fig_tens.add_trace(go.Scatter(x=sig_v, y=z_vals, name='Total (σ)', line=dict(color='black')))
    fig_tens.add_trace(go.Scatter(x=u, y=z_vals, name='Intersticial (u)', line=dict(color='blue', dash='dash')))
    fig_tens.add_trace(go.Scatter(x=sig_v_eff, y=z_vals, name='Efectiva (σ\')', line=dict(color='red')))
    z_acum = 0
    colores = ['#f0f8ff', '#e6e6fa', '#fff0f5', '#f5fffa', '#ffebcd']
    for i, row in df_edit.iterrows():
        z_next = z_acum + row["Espesor (m)"]
        fig_tens.add_hrect(y0=z_acum, y1=z_next, fillcolor=colores[i%len(colores)], opacity=0.4, line_width=0, annotation_text=row["Estrato"])
        z_acum = z_next
    fig_tens.add_hline(y=zw, line_dash="dot", line_color="blue", annotation_text="N.F.")
    fig_tens.update_yaxes(autorange="reversed", title="Profundidad (m)")
    fig_tens.update_xaxes(title="Presión (kPa)", side="top")
    
    with tab_tensiones:
        st.plotly_chart(fig_tens, use_container_width=True)

    resultados = []
    D_arr = np.arange(D_min, D_max + 1e-5, D_step)
    L_arr = np.arange(L_min, L_max + 1e-5, L_step)

    for D in D_arr:
        for L in L_arr:
            res = calcular_pilote_cte(D, L, df_edit, zw, FS, sigma_tope_mpa, tipo_ejecucion, material_pilote)
            if res is not None: resultados.append(res)

    df_res = pd.DataFrame(resultados)
    df_pivot_geo_global = None
    df_pivot_final_global = None
    df_pivot_punta_global = None
    df_pivot_fuste_global = None

    if not df_res.empty:
        columnas = [f"Ø {d_val:.2f} m" for d_val in df_res['D'].unique()]
        
        df_pivot_punta = df_res.pivot(index="L", columns="D", values="Q_punta (kN)") / FS
        df_pivot_punta.index = [f"L = {idx:.1f} m" for idx in df_pivot_punta.index]
        df_pivot_punta.columns = columnas
        df_pivot_punta_global = df_pivot_punta
        with tab_matriz_punta:
            st.dataframe(df_pivot_punta.style.background_gradient(cmap='Reds', axis=None).format("{:.0f}"), use_container_width=True)

        df_pivot_fuste = df_res.pivot(index="L", columns="D", values="Q_fuste (kN)") / FS
        df_pivot_fuste.index = [f"L = {idx:.1f} m" for idx in df_pivot_fuste.index]
        df_pivot_fuste.columns = columnas
        df_pivot_fuste_global = df_pivot_fuste
        with tab_matriz_fuste:
            st.dataframe(df_pivot_fuste.style.background_gradient(cmap='Oranges', axis=None).format("{:.0f}"), use_container_width=True)

        df_pivot_geo = df_res.pivot(index="L", columns="D", values="Q_adm_geo (kN)")
        df_pivot_geo.index = [f"L = {idx:.1f} m" for idx in df_pivot_geo.index]
        df_pivot_geo.columns = columnas
        df_pivot_geo_global = df_pivot_geo
        with tab_matriz_total:
            st.dataframe(df_pivot_geo.style.background_gradient(cmap='Greens', axis=None).format("{:.0f}"), use_container_width=True)

        df_pivot_final = df_res.pivot(index="L", columns="D", values="Q_final (kN)")
        df_pivot_control = df_res.pivot(index="L", columns="D", values="Control")
        df_pivot_punz = df_res.pivot(index="L", columns="D", values="aviso_punz")
        
        df_final_formateada = pd.DataFrame(index=df_pivot_final.index, columns=df_pivot_final.columns)
        for c in df_pivot_final.columns:
            for r in df_pivot_final.index:
                val = df_pivot_final.loc[r, c]
                ctrl = df_pivot_control.loc[r, c]
                punz = df_pivot_punz.loc[r, c]
                
                if ctrl == "ESTRUCTURAL":
                    df_final_formateada.loc[r, c] = f"{val:.0f} [EST]"
                elif punz:
                    df_final_formateada.loc[r, c] = f"{val:.0f} [PUNZ]"
                else:
                    df_final_formateada.loc[r, c] = f"{val:.0f}"

        df_final_formateada.index = [f"L = {idx:.1f} m" for idx in df_final_formateada.index]
        df_final_formateada.columns = columnas
        
        tope_row = {}
        for d_val in df_res['D'].unique():
            q_tope = df_res[df_res['D'] == d_val]['Q_tope_est (kN)'].iloc[0]
            col_name = [c for c in columnas if f"Ø {d_val:.2f} m" in c][0]
            tope_row[col_name] = f"{q_tope:.0f}"
            
        df_tope = pd.DataFrame([tope_row], index=["🛑 TOPE ESTRUCT. (kN)"])
        df_final_formateada = pd.concat([df_tope, df_final_formateada])
        df_pivot_final_global = df_final_formateada
        
        with tab_matriz_tope:
            st.markdown("💡 **Leyenda:** `[EST]` = Límite Estructural del Hormigón | `[PUNZ]` = Límite por Punzonamiento en capa blanda (CTE 5.3.4.1.2)")
            def style_df(df):
                styles = pd.DataFrame('', index=df.index, columns=df.columns)
                for r in df.index:
                    for c in df.columns:
                        val = str(df.loc[r, c])
                        if r == "🛑 TOPE ESTRUCT. (kN)": 
                            styles.loc[r, c] = 'background-color: #d3d3d3; font-weight: bold; border-bottom: 2px solid black;'
                        elif '[EST]' in val: 
                            styles.loc[r, c] = 'background-color: #ffcccc;'
                        elif '[PUNZ]' in val:
                            styles.loc[r, c] = 'background-color: #ffe6cc; font-weight: bold; color: #cc5200;'
                        else: 
                            styles.loc[r, c] = 'background-color: #ccffcc;'
                return styles
            st.dataframe(df_final_formateada.style.apply(style_df, axis=None), use_container_width=True)
            
            df_plot_final = df_res.copy()
            df_plot_final["Diámetro"] = df_plot_final["D"].apply(lambda x: f"Ø {x:.2f} m")
            fig_final = px.line(df_plot_final, x="L", y="Q_final (kN)", color="Diámetro", markers=True, title="Diseño Final (Límite Estructural CTE)", color_discrete_sequence=px.colors.qualitative.Set1, template="plotly_white")
            st.plotly_chart(fig_final, use_container_width=True)
            st.session_state.fig_final_guardada = fig_final 

        with tab_auditoria:
            col_aud1, col_aud2 = st.columns(2)
            d_aud = col_aud1.selectbox("Diámetro Ø (m):", df_res['D'].unique(), format_func=lambda x: f"{x:.2f}")
            l_aud = col_aud2.selectbox("Longitud L (m):", df_res['L'].unique(), format_func=lambda x: f"{x:.2f}")
            res_auditoria = df_res[(df_res['D'] == d_aud) & (df_res['L'] == l_aud)]
            if not res_auditoria.empty:
                fila_aud = res_auditoria.iloc[0]
                
                if fila_aud['aviso_punz']:
                    st.error("⚠️ **ADVERTENCIA DE PUNZONAMIENTO:** La resistencia por punta ha sido drásticamente limitada por la presencia de un estrato cohesivo blando subyacente según el CTE 5.3.4.1.2.")
                
                st.markdown(f"### ➡️ Carga Final: **{fila_aud['Q_final (kN)']:.0f} kN** ({fila_aud['Control']})")
                st.dataframe(pd.DataFrame(fila_aud['auditoria_fuste']).style.format({"Long. fuste (m)": "{:.2f}", "σ'_v media (kPa)": "{:.1f}", "Resist. Unitaria τ_f (kPa)": "{:.2f}", "Fuerza Tramo (kN)": "{:.0f}"}).hide(axis="index"), use_container_width=True)
                st.dataframe(pd.DataFrame([fila_aud['auditoria_punta']]).style.format({"Profundidad Punta (m)": "{:.2f}", "σ'_v efectiva base (kPa)": "{:.1f}", "Factor f_p": "{:.1f}", "Resist. Unitaria q_p (kPa)": "{:.2f}", "Fuerza Total Punta (kN)": "{:.0f}"}).hide(axis="index"), use_container_width=True)
                
                st.markdown(f"#### 🧅 Composición del Bulbo (Promedio según CTE DB-SE-C 5.3.4.1.2)")
                df_bulbo = pd.DataFrame(fila_aud['auditoria_bulbo'])
                st.dataframe(df_bulbo.style.format({"Espesor en bulbo (m)": "{:.2f}", "Participación (%)": "{:.1f}%", "q_p individual (kPa)": "{:.0f}"}).hide(axis="index"), use_container_width=True)

                st.markdown("---")
                D_val = fila_aud['D']; L_val = fila_aud['L']; z_sup = fila_aud['z_sup_bulbo']; z_inf = fila_aud['z_inf_bulbo']
                
                fig_bulbo = go.Figure()
                z_acum_plt = 0
                colores_plt = ['#f0f8ff', '#e6e6fa', '#fff0f5', '#f5fffa', '#ffebcd']
                for i_plt, row_plt in df_edit.iterrows():
                    is_last = (i_plt == len(df_edit) - 1)
                    z_next_plt = z_acum_plt + row_plt["Espesor (m)"]
                    if is_last: z_next_plt = max(z_next_plt, z_inf + D_val)
                    fig_bulbo.add_hrect(y0=z_acum_plt, y1=z_next_plt, fillcolor=colores_plt[i_plt%len(colores_plt)], opacity=0.7, line_width=1, annotation_text=row_plt["Estrato"], annotation_position="top left")
                    z_acum_plt = z_next_plt

                fig_bulbo.add_shape(type="rect", x0=-D_val*1.2, x1=D_val*1.2, y0=-D_val*0.6, y1=0, fillcolor="#8B8C89", line=dict(color="black", width=2))
                fig_bulbo.add_shape(type="rect", x0=-D_val/2, x1=D_val/2, y0=0, y1=L_val, fillcolor="#A9ACA9", line=dict(color="black", width=2))
                fig_bulbo.add_shape(type="rect", x0=-D_val/2, x1=-D_val/4, y0=0, y1=L_val, fillcolor="black", opacity=0.15, line_width=0)
                fig_bulbo.add_shape(type="rect", x0=D_val/4, x1=D_val/2, y0=0, y1=L_val, fillcolor="white", opacity=0.25, line_width=0)
                fig_bulbo.add_shape(type="line", x0=-D_val/2, x1=D_val/2, y0=L_val, y1=L_val, line=dict(color="black", width=4))
                fig_bulbo.add_shape(type="rect", x0=-D_val*1.5, x1=D_val*1.5, y0=z_sup, y1=z_inf, fillcolor="rgba(255, 0, 0, 0.15)", line=dict(color="red", width=3, dash="dash"))
                
                fig_bulbo.add_annotation(x=D_val*1.6, y=(z_sup+L_val)/2, text=f"6D Arriba<br>({L_val - z_sup:.1f} m)", showarrow=False, xanchor="left", font=dict(color="red", size=13, family="Arial Black"))
                fig_bulbo.add_annotation(x=D_val*1.6, y=(L_val+z_inf)/2, text=f"3D Abajo<br>({z_inf - L_val:.1f} m)", showarrow=False, xanchor="left", font=dict(color="red", size=13, family="Arial Black"))
                
                fig_bulbo.add_hline(y=L_val, line_color="black", line_width=2, opacity=0.8)
                fig_bulbo.add_hline(y=zw, line_dash="dot", line_color="blue", annotation_text="N.F.", annotation_position="bottom right")

                rango_inf_dibujo = max(z_max_total, z_inf + 2*D_val) 
                fig_bulbo.update_yaxes(autorange="reversed", title="Profundidad (m)", range=[rango_inf_dibujo, -D_val], showgrid=True)
                fig_bulbo.update_xaxes(showticklabels=False, range=[-D_val*4, D_val*4], showgrid=False)
                fig_bulbo.update_layout(height=600, margin=dict(l=20, r=20, t=20, b=20), paper_bgcolor="white", plot_bgcolor="white")
                st.plotly_chart(fig_bulbo, use_container_width=True)

        with tab_asientos:
            st.subheader("📉 Cálculo de Asientos Estimados (Pilote Aislado)")
            st.markdown("Estimación analítica del asiento bajo cargas de servicio (ELS) según el CTE DB-SE-C (Apartado F.2.6.1).")
            
            col_as1, col_as2, col_as3 = st.columns(3)
            
            d_as = col_as1.selectbox("Diámetro Ø (m):", df_res['D'].unique(), format_func=lambda x: f"{x:.2f}", key="d_as")
            l_as = col_as2.selectbox("Longitud L (m):", df_res['L'].unique(), format_func=lambda x: f"{x:.2f}", key="l_as")
            
            res_as = df_res[(df_res['D'] == d_as) & (df_res['L'] == l_as)]
            
            if not res_as.empty:
                fila_as = res_as.iloc[0]
                q_diseno = fila_as['Q_final (kN)']
                
                q_servicio = col_as3.number_input("Carga de Servicio a aplicar, P (kN):", min_value=0.0, value=float(q_diseno), step=100.0)
                
                if q_servicio > q_diseno:
                    st.warning(f"⚠️ La carga introducida supera la capacidad de diseño final del pilote ({q_diseno:.0f} kN).")

                st.markdown("---")
                st.markdown("#### Parámetros de Deformación (Inputs)")
                col_mod1, col_mod2 = st.columns(2)
                
                ep_def = 30000.0 if "Hormigón" in material_pilote else 210000.0 if material_pilote == "Acero" else 10000.0
                Ep = col_mod1.number_input("Módulo de elasticidad del Pilote, E (MPa):", value=ep_def, step=1000.0)
                l1_as = col_mod2.number_input("Longitud fuera del terreno, l1 (m):", value=0.0, step=0.5)

                # --- MOTOR DE ASIENTOS (CTE DB-SE-C F.2.6.1 - Fórmulas F.44 y F.45) ---
                Area_p = (math.pi * d_as**2) / 4.0
                
                R_pk = fila_as['Q_punta (kN)']
                R_fk = fila_as['Q_fuste (kN)']
                R_ck = R_pk + R_fk
                
                alpha = (0.5 * R_fk + R_pk) / R_ck if R_ck > 0 else 1.0
                l2_as = l_as
                
                s1_m = q_servicio * (l1_as + alpha * l2_as) / (Area_p * (Ep * 1000.0))
                s1_mm = s1_m * 1000.0
                
                s2_m = q_servicio * (d_as / (40.0 * R_ck)) if R_ck > 0 else 0.0
                s2_mm = s2_m * 1000.0
                
                s_total_mm = s1_mm + s2_mm

                st.markdown("#### 📊 Resultados de Asiento (CTE F.2.6.1)")
                st.info(f"💡 **Dato interno:** El pilote trabaja con un factor de transmisión $\\alpha$ = **{alpha:.3f}** (0.5 = 100% fuste | 1.0 = 100% punta).")
                
                col_res1, col_res2, col_res3 = st.columns(3)
                col_res1.metric("Acortamiento Pilote", f"{s1_mm:.2f} mm")
                col_res2.metric("Asiento del Terreno", f"{s2_mm:.2f} mm")
                col_res3.metric("ASIENTO TOTAL (s)", f"{s_total_mm:.2f} mm", delta_color="inverse")
                
                if s_total_mm > 25.0:
                    st.error(f"⚠️ **Atención:** El asiento estimado ({s_total_mm:.2f} mm) supera los 25 mm. Revisar compatibilidad con la estructura (ELS).")
                else:
                    st.success(f"✅ Asiento estimado dentro de los rangos admisibles típicos (< 25 mm).")

        with tab_formulacion:
            st.subheader("📖 Formulación Analítica (CTE DB-SE-C)")
            st.markdown("Resumen de las ecuaciones implementadas según el Documento Básico SE-C para el cálculo de la resistencia al hundimiento y asientos.")

            col_form1, col_form2 = st.columns(2)

            with col_form1:
                st.markdown("### 1. Suelos Finos (Corto Plazo)")
                st.info("Condición no drenada. Depende de la cohesión sin drenaje ($c_u$).")
                
                st.markdown("**1.1. Resistencia por Punta ($q_p$)**")
                st.latex(r"q_p = 9 \cdot c_u")
                st.markdown("*Referencia: CTE DB-SE-C, Apartado F.2.1.2. (Ecuación F.33).*")
                
                st.markdown("**1.2. Resistencia por Fuste ($\tau_f$)**")
                st.latex(r"\tau_f = c_u \cdot \frac{100}{100 + c_u}")
                st.markdown("*Referencia: CTE DB-SE-C, Apartado F.2.1.2.*")
                if material_pilote == "Acero":
                    st.warning("⚠️ **Nota:** Al ser un pilote de ACERO, la norma exige aplicar un coeficiente reductor de **0.80** a la fricción hiperbólica.")

                st.markdown("---")
                st.markdown("### 3. Comprobación de Punzonamiento")
                st.markdown("Si existe una capa arcillosa blanda por debajo de la punta, el CTE limita la capacidad para evitar el hundimiento en bloque.")
                st.latex(r"q_p \le 6 \cdot \left[1 + \frac{H}{D}\right] \cdot c_u")
                st.markdown("*Referencia: CTE DB-SE-C, Apartado 5.3.4.1.2. (Ecuación 5.11).*")

            with col_form2:
                st.markdown("### 2. Suelos Granulares (Largo Plazo)")
                st.info("Condición drenada. Se **desprecia la cohesión ($c'$ = 0)** y depende del ángulo de rozamiento ($\phi'$).")
                
                st.markdown("**2.1. Resistencia por Punta ($q_p$)**")
                st.latex(r"q_p = f_p \cdot \sigma'_{vp} \cdot N_q \le 20 \text{ MPa}")
                st.markdown(f"*Factor $f_p$ aplicado actual: **{f_p}** (Pilote {tipo_ejecucion}).*")
                
                st.markdown("**2.2. Resistencia por Fuste ($\tau_f$)**")
                st.latex(r"\tau_f = \sigma'_v \cdot K_f \cdot f \cdot \tan(\phi') \le 120 \text{ kPa}")
                st.markdown(f"*Coeficiente de empuje $K_f$: **{K_f}** (Pilote {tipo_ejecucion}).*")
                st.markdown(f"*Factor de material $f$: **{f_mat}** ({material_pilote}).*")
                
                st.markdown("---")
                st.markdown("### 4. Asientos (Pilote Aislado)")
                st.latex(r"s_i = \left( \frac{D}{40 R_{ck}} + \frac{l_1 + \alpha l_2}{A E} \right) P")
                st.latex(r"\alpha = \frac{1}{R_{ck}} (0.5 R_{fk} + R_{pk})")
                st.markdown("*Referencia: CTE DB-SE-C, Apartado F.2.6.1. (Ecuaciones F.44 y F.45).*")

    st.sidebar.markdown("---")
    col_ex1, col_ex2 = st.sidebar.columns(2)
    if col_ex1.button("📑 Word", type="primary", use_container_width=True):
        if not df_res.empty:
            st.session_state.word_buffer = generar_word_cte(df_edit, df_pivot_geo_global, df_pivot_final_global, FS, sit_str, zw, sigma_tope_mpa, df_unitarios, st.session_state.fig_final_guardada, tipo_ejecucion, material_pilote)
            st.sidebar.success("✅ Word listo!")
    if col_ex2.button("📊 Excel", type="primary", use_container_width=True):
        if not df_res.empty:
            st.session_state.excel_buffer = generar_excel_cte(df_edit, df_unitarios, df_pivot_punta_global, df_pivot_fuste_global, df_pivot_geo_global, df_pivot_final_global, zw, FS, sit_str, sigma_tope_mpa)
            st.sidebar.success("✅ Excel listo!")

    if st.session_state.word_buffer:
        st.sidebar.download_button("⬇️ Bajar Word", data=st.session_state.word_buffer, file_name="Memoria_CTE.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
    if st.session_state.excel_buffer:
        st.sidebar.download_button("⬇️ Bajar Excel", data=st.session_state.excel_buffer, file_name="Matrices_CTE.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True)