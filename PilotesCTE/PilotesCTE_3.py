import streamlit as st
import numpy as np
import pandas as pd
import math
import io
import plotly.graph_objects as go
import plotly.express as px
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date

st.set_page_config(page_title="Cálculo de Pilotes - CTE DB-SE-C", layout="wide", page_icon="🏢")

st.title("🏢 Diseño de Pilotes CTE DB-SE-C")
st.markdown("Cálculo de cimentaciones profundas según el Documento Básico SE-C del Código Técnico de la Edificación.")

# ══════════════════════════════════════════════════════════════════════════
# INICIALIZACIÓN DE LA TABLA BASE Y CLAVE DINÁMICA
# ══════════════════════════════════════════════════════════════════════════
if 'df_base' not in st.session_state:
    st.session_state.df_base = pd.DataFrame({
        "Estrato": ["UG-01", "UG-02", "UG-03", "UG-04"],
        "Espesor (m)": [2.0, 5.0, 8.0, 10.0],
        "Gamma Seco (kN/m3)": [18.0, 19.0, 18.5, 20.0],
        "Gamma Sat. (kN/m3)": [20.0, 20.0, 20.5, 21.0],
        "Condición": ["Corto Plazo", "Corto Plazo", "Largo Plazo", "Largo Plazo"],
        "c / cu (kPa)": [30.0, 150.0, 0.0, 0.0],
        "phi (grados)": [0.0, 0.0, 30.0, 35.0]
    })

if 'table_key' not in st.session_state: 
    st.session_state.table_key = 0

if 'calculado' not in st.session_state: st.session_state.calculado = False
if 'word_buffer' not in st.session_state: st.session_state.word_buffer = None
if 'fig_final_guardada' not in st.session_state: st.session_state.fig_final_guardada = None

# ══════════════════════════════════════════════════════════════════════════
# BARRA LATERAL: CONFIGURACIÓN Y TABLA 5.1 CTE
# ══════════════════════════════════════════════════════════════════════════
st.sidebar.header("⚙️ Ejecución y Tope (Tabla 5.1)")

metodo_inst = st.sidebar.selectbox("Método de Ejecución", ["Perforados", "Hincados"])

if metodo_inst == "Hincados":
    fp_val, Kf_val = 3.0, 1.0
    tipo_pilote_hincado = st.sidebar.selectbox("Tipo de pilote", ["Hormigón armado", "Hormigón pretensado o postesado", "Metálicos", "Madera"])

    if tipo_pilote_hincado == "Hormigón armado":
        fck = st.sidebar.number_input("Resistencia fck (MPa)", value=30.0, step=5.0)
        sigma_tope_mpa = 0.30 * fck
        f_rug, is_steel = 0.9, False
    elif tipo_pilote_hincado == "Hormigón pretensado o postesado":
        fck = st.sidebar.number_input("Resistencia fck (MPa)", value=40.0, step=5.0)
        fp_tension = st.sidebar.number_input("Tensión fp (MPa)", value=10.0, step=1.0)
        sigma_tope_mpa = 0.30 * (fck - 0.9 * fp_tension)
        f_rug, is_steel = 0.9, False
    elif tipo_pilote_hincado == "Metálicos":
        fyk = st.sidebar.number_input("Límite elástico fyk (MPa)", value=275.0, step=10.0)
        sigma_tope_mpa = 0.30 * fyk
        f_rug, is_steel = 0.8, True
    else:
        sigma_tope_mpa = 5.0
        f_rug, is_steel = 1.0, False
        
    desc_tope = f"Hincado - {tipo_pilote_hincado}"

else: # Perforados
    fp_val, Kf_val, f_rug, is_steel = 2.5, 0.75, 1.0, False
    tipo_pilote_perf = st.sidebar.selectbox("Procedimiento específico", ["Entubados", "Lodos", "En seco", "Barrenados con control de parámetros", "Barrenados sin control de parámetros"])
    apoyo = st.sidebar.radio("Tipo de apoyo en punta", ["Suelo firme", "Roca"])

    if tipo_pilote_perf == "Entubados":
        sigma_tope_mpa = 6.0 if apoyo == "Roca" else 5.0
    elif tipo_pilote_perf in ["Lodos", "En seco"]:
        sigma_tope_mpa = 5.0 if apoyo == "Roca" else 4.0
    elif tipo_pilote_perf == "Barrenados con control de parámetros":
        sigma_tope_mpa = 4.0
    elif tipo_pilote_perf == "Barrenados sin control de parámetros":
        sigma_tope_mpa = 3.5
        if apoyo == "Roca": st.sidebar.warning("Tabla 5.1 no define apoyo en roca para este método. Se asume 3.5 MPa.")

    control_integridad = st.sidebar.checkbox("Control de integridad (+25% Tope)")
    if control_integridad:
        sigma_tope_mpa *= 1.25
        
    desc_tope = f"Perforado ({tipo_pilote_perf}) en {apoyo}" + (" [+25%]" if control_integridad else "")

st.sidebar.success(f"🛑 **Tope Estructural:** {sigma_tope_mpa:.2f} MPa")

st.sidebar.markdown("---")
st.sidebar.header("💧 Datos Geotécnicos")
zw = st.sidebar.number_input("Prof. Nivel Freático, zw (m)", min_value=0.0, value=3.0, step=0.5)
z_nulo = st.sidebar.number_input("Fuste Nulo en Cabeza (m)", min_value=0.0, value=1.5, step=0.5, help="Zona de desecación superficial donde no se cuenta rozamiento.")
gamma_r = st.sidebar.number_input("Coef. Parcial Resistencia (γ_R):", min_value=1.0, value=3.00, step=0.1, help="Según Tabla 2.1 del CTE DB-SE-C (Valor habitual 3.0)")

st.sidebar.header("📐 Matriz Geométrica")
col_d1, col_d2 = st.sidebar.columns(2)
D_min = col_d1.number_input("Ø min (m)", value=0.4, min_value=0.2, step=0.1)
D_max = col_d2.number_input("Ø max (m)", value=1.0, min_value=0.2, step=0.1)
D_step = st.sidebar.number_input("Paso Ø (m)", value=0.2, min_value=0.1)

z_max_total_sidebar = st.session_state.df_base["Espesor (m)"].sum()
col_l1, col_l2 = st.sidebar.columns(2)

val_def_L_min = max(1.0, min(10.0, float(z_max_total_sidebar)))
L_min = col_l1.number_input("L min (m)", value=val_def_L_min, min_value=1.0, step=1.0)

val_def_L_max = max(float(L_min), min(20.0, float(z_max_total_sidebar)))
L_max = col_l2.number_input("L max (m)", value=val_def_L_max, min_value=float(L_min), step=1.0)
L_step = st.sidebar.number_input("Paso L (m)", value=2.5, min_value=0.5)


# ══════════════════════════════════════════════════════════════════════════
# MOTORES DE CÁLCULO CTE DB-SE-C Y PUNZONAMIENTO
# ══════════════════════════════════════════════════════════════════════════
def calcular_perfil_tensiones(df, zw, z_max):
    z_max_extendido = z_max + 10.0 
    z_vals = np.arange(0, z_max_extendido + 0.1, 0.1)
    sigma_v, u, sigma_v_eff = [], [], []
    sv_acum = 0.0
    for z in z_vals:
        z_acum, gamma_actual = 0.0, 18.0
        for idx, row in df.iterrows():
            is_last = (idx == len(df) - 1)
            if z <= z_acum + row["Espesor (m)"] + 1e-5 or is_last:
                gamma_actual = row["Gamma Seco (kN/m3)"] if z <= zw + 1e-5 else row["Gamma Sat. (kN/m3)"]
                break
            z_acum += row["Espesor (m)"]
            
        sv_acum += gamma_actual * 0.1 if z > 0 else 0.0
        u_val = max(0, (z - zw) * 9.81)
        sigma_v.append(sv_acum)
        u.append(u_val)
        sigma_v_eff.append(sv_acum - u_val)
    return z_vals, np.array(sigma_v), np.array(u), np.array(sigma_v_eff)

def obtener_tension_a_profundidad(z_target, z_vals, sigma_array):
    return np.interp(z_target, z_vals, sigma_array)

def calcular_pilote_cte(D, L, df, zw, z_nulo, gamma_r_val, sigma_tope, fp, Kf, f_rug, is_steel):
    z_max = df["Espesor (m)"].sum()
    if L > z_max: return None 
    
    z_vals, sig_v, u, sig_v_eff = calcular_perfil_tensiones(df, zw, z_max)
    z_eval_punta = L
    sig_v_eff_punta = obtener_tension_a_profundidad(z_eval_punta, z_vals, sig_v_eff)
    
    z_sup_bulbo, z_inf_bulbo = max(0.0, L - (6 * D)), L + (3 * D)
    tipo_bulbo = "6D/3D"

    espesor_bulbo = z_inf_bulbo - z_sup_bulbo
    qp_eq_acumulado, z_acum, detalle_bulbo_grafico = 0.0, 0.0, []
    
    # 1. Cálculo base de punta ponderada (CTE)
    for idx, row in df.iterrows():
        is_last = (idx == len(df) - 1)
        h_estrato = row["Espesor (m)"]
        z_top_estrato = z_acum
        z_bot_estrato = max(z_acum + h_estrato, z_inf_bulbo) if is_last else z_acum + h_estrato
            
        overlap_top = max(z_sup_bulbo, z_top_estrato)
        overlap_bot = min(z_inf_bulbo, z_bot_estrato)
        overlap_h = max(0.0, overlap_bot - overlap_top)
        
        if overlap_h > 0 and espesor_bulbo > 0:
            cond_i, c_i, phi_i = row["Condición"], row["c / cu (kPa)"], row["phi (grados)"]
            
            if "Corto Plazo" in cond_i:
                qp_i = 9.0 * c_i
            else:
                phi_rad = math.radians(phi_i)
                Nq = 1.0 if phi_i == 0 else ((1 + math.sin(phi_rad))/(1 - math.sin(phi_rad))) * math.exp(math.pi * math.tan(phi_rad))
                qp_i = min(fp * sig_v_eff_punta * Nq, 20000.0) 
            
            peso = overlap_h / espesor_bulbo
            qp_eq_acumulado += qp_i * peso
            
            detalle_bulbo_grafico.append({
                "Estrato": f"{row['Estrato']} (Prol.)" if is_last and overlap_bot > (z_acum + h_estrato) else row['Estrato'],
                "Espesor en bulbo (m)": overlap_h, "Participación (%)": peso * 100.0, "q_p individual (kPa)": qp_i
            })
        z_acum += h_estrato

    # 2. FILTRO DE SEGURIDAD (Punzonamiento de estrato inferior blando para justificar CTE 5.2.2)
    qp_original_ponderada = qp_eq_acumulado
    alerta_punzonamiento = None
    z_acum_temp = 0.0
    
    for idx, row in df.iterrows():
        h_estrato = row["Espesor (m)"]
        z_top_estrato = z_acum_temp
        z_acum_temp += h_estrato
        
        if z_top_estrato >= L - 1e-5 and "Corto Plazo" in row["Condición"]:
            H_distancia = z_top_estrato - L
            cu_inferior = row["c / cu (kPa)"]
            
            qp_lim_punzonamiento = 6.0 * ((1.0 + (H_distancia / D))**2) * cu_inferior
            
            if qp_lim_punzonamiento < qp_eq_acumulado:
                qp_eq_acumulado = qp_lim_punzonamiento
                alerta_punzonamiento = f"Limitado por punzonamiento hacia estrato '{row['Estrato']}' situado a H = {H_distancia:.2f} m."

    Area_pilote = (math.pi * D**2) / 4.0
    Q_punta = qp_eq_acumulado * Area_pilote
    
    auditoria_punta = {
        "Profundidad Punta (m)": L, "σ'_v efectiva base (kPa)": sig_v_eff_punta,
        "q_p original ponderada (kPa)": qp_original_ponderada,
        "Resist. Unitaria q_p FINAL (kPa)": qp_eq_acumulado, 
        "Fuerza Total Punta (kN)": Q_punta,
        "Alerta Punzonamiento": alerta_punzonamiento
    }
    
    # 3. Cálculo de Fuste
    Q_fuste, z_top, Perimetro, auditoria_fuste = 0.0, 0.0, math.pi * D, []
    for _, row in df.iterrows():
        z_bot = z_top + row["Espesor (m)"]
        if z_top >= L: break
        
        z_start_fuste = max(z_top, z_nulo)
        z_end_tramo = min(z_bot, L)
        
        if z_end_tramo <= z_start_fuste:
            z_top = z_bot; continue
            
        puntos_corte = [z_start_fuste]
        if z_start_fuste < zw < z_end_tramo: puntos_corte.append(zw)
        puntos_corte.append(z_end_tramo)
        
        for i in range(len(puntos_corte) - 1):
            z_sub_top, z_sub_bot = puntos_corte[i], puntos_corte[i+1]
            L_sub = z_sub_bot - z_sub_top
            sig_v_eff_mid = obtener_tension_a_profundidad(z_sub_top + (L_sub / 2.0), z_vals, sig_v_eff)
            
            if "Corto Plazo" in row["Condición"]:
                tau_f = (100.0 * row["c / cu (kPa)"]) / (100.0 + row["c / cu (kPa)"])
                if is_steel: tau_f *= 0.8
            else:
                tau_f_raw = sig_v_eff_mid * Kf * f_rug * math.tan(math.radians(row["phi (grados)"]))
                tau_f = min(tau_f_raw, 100.0 if row["c / cu (kPa)"] > 0 else 120.0)
                
            Q_tramo = tau_f * Perimetro * L_sub
            Q_fuste += Q_tramo
            
            sufijo = " (Seco)" if len(puntos_corte) > 2 and z_sub_bot <= zw else (" (Sat.)" if len(puntos_corte) > 2 else "")
            auditoria_fuste.append({
                "Estrato": row["Estrato"] + sufijo, "Cotas (m)": f"{z_sub_top:.1f} a {z_sub_bot:.1f}",
                "Long. fuste (m)": L_sub, "σ'_v media (kPa)": sig_v_eff_mid,
                "Resist. Unitaria τ_f (kPa)": tau_f, "Fuerza Tramo (kN)": Q_tramo
            })
        z_top = z_bot
        
    Q_adm_geo = (Q_punta + Q_fuste) / gamma_r_val
    Q_tope_est = Area_pilote * (sigma_tope * 1000.0)

    return {
        "D": D, "L": L, "Q_punta (kN)": Q_punta, "Q_fuste (kN)": Q_fuste, 
        "Q_adm_geo (kN)": Q_adm_geo, "Q_tope_est (kN)": Q_tope_est, "Q_final (kN)": min(Q_adm_geo, Q_tope_est),
        "Control": "ESTRUCTURAL" if Q_tope_est < Q_adm_geo else "GEOTÉCNICA",
        "auditoria_punta": auditoria_punta, "auditoria_fuste": auditoria_fuste,
        "auditoria_bulbo": detalle_bulbo_grafico, "z_sup_bulbo": z_sup_bulbo, "z_inf_bulbo": z_inf_bulbo,
        "tipo_bulbo": tipo_bulbo,
        "alerta_punzonamiento": alerta_punzonamiento
    }

# ══════════════════════════════════════════════════════════════════════════
# INTERFAZ PRINCIPAL DE PESTAÑAS 
# ══════════════════════════════════════════════════════════════════════════
tab_datos, tab_tensiones, tab_matriz_punta, tab_matriz_fuste, tab_matriz_total, tab_matriz_tope, tab_auditoria, tab_formulacion = st.tabs([
    "📋 1. Estratigrafía", "🌊 2. Tensiones", "🔻 3. Punta", "🟫 4. Fuste", "🌍 5. Total", "🛑 6. Tope Estruct.", "🔍 7. Auditoría", "📖 8. Formulación"
])

with tab_datos:
    st.subheader("Definición de Unidades Geotécnicas")
    
    df_edit = st.data_editor(
        st.session_state.df_base, 
        key=f"tabla_estratos_{st.session_state.table_key}", 
        num_rows="dynamic", 
        use_container_width=True,
        column_config={
            "Condición": st.column_config.SelectboxColumn(options=["Corto Plazo", "Largo Plazo"], required=True)
        }
    )
    
    # 🛡️ AUTOCORRECCIÓN INFALIBLE
    df_modificado = False
    
    mascara_corto = df_edit["Condición"] == "Corto Plazo"
    if (df_edit.loc[mascara_corto, "phi (grados)"] != 0.0).any():
        df_edit.loc[mascara_corto, "phi (grados)"] = 0.0
        df_modificado = True
        st.toast("💡 Corto Plazo detectado: Ángulo de rozamiento forzado a 0º", icon="⚙️")

    mascara_largo = df_edit["Condición"] == "Largo Plazo"
    if (df_edit.loc[mascara_largo, "c / cu (kPa)"] != 0.0).any():
        df_edit.loc[mascara_largo, "c / cu (kPa)"] = 0.0
        df_modificado = True
        st.toast("💡 Largo Plazo detectado: Cohesión forzada a 0 kPa", icon="⚙️")

    if df_modificado:
        st.session_state.df_base = df_edit.copy()
        st.session_state.table_key += 1 
        st.rerun()

    espesores_invalidos = (df_edit["Espesor (m)"] <= 0).any()
    columnas_geomecanicas = ["Gamma Seco (kN/m3)", "Gamma Sat. (kN/m3)", "c / cu (kPa)", "phi (grados)"]
    valores_negativos = (df_edit[columnas_geomecanicas] < 0).any().any()

    if espesores_invalidos: st.error("⛔ **Error:** Los estratos deben tener un espesor > 0 m.")
    if valores_negativos: st.error("⛔ **Error:** No se permiten valores negativos en las propiedades del terreno.")

    errores_presentes = espesores_invalidos or valores_negativos
    z_max_total = df_edit["Espesor (m)"].sum()
    st.info(f"Profundidad máxima actual del sondeo: **{z_max_total:.2f} m**.")
    
    if not st.session_state.calculado and not errores_presentes:
        st.warning("👈 Haz clic en **'Calcular Pilotes'** en el menú izquierdo.")

st.sidebar.markdown("---")
if st.sidebar.button("⚙️ Calcular Pilotes", type="primary", use_container_width=True, disabled=errores_presentes):
    st.session_state.calculado = True
    st.session_state.word_buffer = None

# ══════════════════════════════════════════════════════════════════════════
# CÁLCULOS PRINCIPALES Y RENDERIZADO
# ══════════════════════════════════════════════════════════════════════════
if st.session_state.calculado:
    z_vals, sig_v, u, sig_v_eff = calcular_perfil_tensiones(df_edit, zw, z_max_total)
    
    with tab_datos:
        st.markdown("---")
        st.subheader("📊 Resistencias Unitarias del Terreno (CTE DB-SE-C)")
        
        datos_unitarios = []
        z_top_loop = 0.0
        for idx, row_est in df_edit.iterrows():
            espesor = row_est["Espesor (m)"]
            if espesor <= 0: continue
            z_bot_loop = z_top_loop + espesor
            
            puntos_corte = [z_top_loop]
            if z_top_loop < zw < z_bot_loop: puntos_corte.append(zw)
            puntos_corte.append(z_bot_loop)
            
            for i in range(len(puntos_corte) - 1):
                z_sub_top = puntos_corte[i]
                z_sub_bot = puntos_corte[i+1]
                z_mid_loop = z_sub_top + (z_sub_bot - z_sub_top) / 2.0
                
                sig_v_eff_mid_loop = obtener_tension_a_profundidad(z_mid_loop, z_vals, sig_v_eff)
                sig_v_eff_base = obtener_tension_a_profundidad(z_sub_bot, z_vals, sig_v_eff)
                cond_loop, c_loop, phi_loop = row_est["Condición"], row_est["c / cu (kPa)"], row_est["phi (grados)"]
                
                if "Corto Plazo" in cond_loop:
                    tau_f_media = (100.0 * c_loop) / (100.0 + c_loop)
                    if is_steel: tau_f_media *= 0.8
                    qp_base = 9.0 * c_loop
                else:
                    tau_f_raw = sig_v_eff_mid_loop * Kf_val * f_rug * math.tan(math.radians(phi_loop))
                    tau_f_media = min(tau_f_raw, 100.0 if c_loop > 0 else 120.0)
                    
                    phi_rad = math.radians(phi_loop)
                    Nq_star = 1.0 if phi_loop == 0 else ((1 + math.sin(phi_rad))/(1 - math.sin(phi_rad))) * math.exp(math.pi * math.tan(phi_rad))
                    qp_base = min(fp_val * sig_v_eff_base * Nq_star, 20000.0)
                
                sufijo = " (Seco)" if len(puntos_corte) > 2 and z_sub_bot <= zw else (" (Sat.)" if len(puntos_corte) > 2 else "")
                datos_unitarios.append({
                    "Estrato": row_est["Estrato"] + sufijo,
                    "Profundidad (m)": f"de {z_sub_top:.1f} a {z_sub_bot:.1f}",
                    "Fricción media, τ_f (kPa)": round(tau_f_media, 1) if z_mid_loop > z_nulo else 0.0,
                    "Punta base, q_p (kPa)": round(qp_base, 1)
                })
            z_top_loop = z_bot_loop
            
        df_unitarios = pd.DataFrame(datos_unitarios)
        st.dataframe(df_unitarios.style.format({"Fricción media, τ_f (kPa)": "{:.1f}", "Punta base, q_p (kPa)": "{:.1f}"}).hide(axis="index"), use_container_width=True)

    fig_tens = go.Figure()
    fig_tens.add_trace(go.Scatter(x=sig_v, y=z_vals, name='Tensión Total (σ)', line=dict(color='black', width=2)))
    fig_tens.add_trace(go.Scatter(x=u, y=z_vals, name='Pres. Intersticial (u)', line=dict(color='blue', dash='dash')))
    fig_tens.add_trace(go.Scatter(x=sig_v_eff, y=z_vals, name='Tensión Efectiva (σ\')', line=dict(color='red', width=3)))
    z_acum = 0
    colores = ['#f0f8ff', '#e6e6fa', '#fff0f5', '#f5fffa', '#ffebcd']
    for i, row in df_edit.iterrows():
        z_next = z_acum + row["Espesor (m)"]
        fig_tens.add_hrect(y0=z_acum, y1=z_next, fillcolor=colores[i%len(colores)], opacity=0.4, line_width=0, annotation_text=row["Estrato"])
        z_acum = z_next
    fig_tens.add_hline(y=zw, line_dash="dot", line_color="blue", annotation_text="Nivel Freático")
    fig_tens.update_yaxes(autorange="reversed", title="Profundidad (m)")
    fig_tens.update_xaxes(title="Presión (kPa)", side="top")
    
    with tab_tensiones: 
        st.plotly_chart(fig_tens, use_container_width=True)
        st.markdown("---")
        st.subheader("🔎 Consulta puntual de tensiones")
        z_consulta = st.number_input("Introduce una profundidad Z (m):", min_value=0.0, max_value=float(z_max_total + 10.0), value=float(zw), step=0.1)
        sig_v_z = obtener_tension_a_profundidad(z_consulta, z_vals, sig_v)
        u_z = obtener_tension_a_profundidad(z_consulta, z_vals, u)
        sig_v_eff_z = obtener_tension_a_profundidad(z_consulta, z_vals, sig_v_eff)
        
        col_t1, col_t2, col_t3 = st.columns(3)
        col_t1.metric(label="Tensión Total (σ)", value=f"{sig_v_z:.2f} kPa")
        col_t2.metric(label="Pres. Intersticial (u)", value=f"{u_z:.2f} kPa")
        col_t3.metric(label="Tensión Efectiva (σ')", value=f"{sig_v_eff_z:.2f} kPa")

    resultados = []
    for D in np.arange(D_min, D_max + 1e-5, D_step):
        for L in np.arange(L_min, L_max + 1e-5, L_step):
            res = calcular_pilote_cte(D, L, df_edit, zw, z_nulo, gamma_r, sigma_tope_mpa, fp_val, Kf_val, f_rug, is_steel)
            if res is not None: resultados.append(res)

    df_res = pd.DataFrame(resultados)
    df_pivot_geo_global, df_pivot_final_global = None, None

    if not df_res.empty:
        columnas_con_d = [f"Ø {d_val:.2f} m" for d_val in df_res['D'].unique()]
        
        df_pivot_punta = df_res.pivot(index="L", columns="D", values="Q_punta (kN)") / gamma_r
        df_pivot_punta.index, df_pivot_punta.columns = [f"L = {idx:.1f} m" for idx in df_pivot_punta.index], columnas_con_d
        with tab_matriz_punta:
            st.subheader(f"🔻 Resistencia Diseño SOLO por Punta, R_cd (kN)")
            st.dataframe(df_pivot_punta.style.background_gradient(cmap='Reds', axis=None).format("{:.0f}"), use_container_width=True)
            st.markdown("---")
            df_plot_punta = df_res.copy()
            df_plot_punta["Diámetro"] = df_plot_punta["D"].apply(lambda x: f"Ø {x:.2f} m")
            df_plot_punta["R_cd_punta (kN)"] = df_plot_punta["Q_punta (kN)"] / gamma_r
            fig_punta = px.line(df_plot_punta, x="L", y="R_cd_punta (kN)", color="Diámetro", markers=True, title=f"Capacidad de Diseño por Punta vs. Longitud (γ_R = {gamma_r:.2f})", template="plotly_white")
            st.plotly_chart(fig_punta, use_container_width=True)

        df_pivot_fuste = df_res.pivot(index="L", columns="D", values="Q_fuste (kN)") / gamma_r
        df_pivot_fuste.index, df_pivot_fuste.columns = [f"L = {idx:.1f} m" for idx in df_pivot_fuste.index], columnas_con_d
        with tab_matriz_fuste:
            st.subheader(f"🟫 Resistencia Diseño SOLO por Fuste, R_cd (kN)")
            st.dataframe(df_pivot_fuste.style.background_gradient(cmap='Oranges', axis=None).format("{:.0f}"), use_container_width=True)
            st.markdown("---")
            df_plot_fuste = df_res.copy()
            df_plot_fuste["Diámetro"] = df_plot_fuste["D"].apply(lambda x: f"Ø {x:.2f} m")
            df_plot_fuste["R_cd_fuste (kN)"] = df_plot_fuste["Q_fuste (kN)"] / gamma_r
            fig_fuste = px.line(df_plot_fuste, x="L", y="R_cd_fuste (kN)", color="Diámetro", markers=True, title=f"Capacidad de Diseño por Fuste vs. Longitud (γ_R = {gamma_r:.2f})", template="plotly_white")
            st.plotly_chart(fig_fuste, use_container_width=True)

        df_pivot_geo = df_res.pivot(index="L", columns="D", values="Q_adm_geo (kN)")
        df_pivot_geo.index, df_pivot_geo.columns = [f"L = {idx:.1f} m" for idx in df_pivot_geo.index], columnas_con_d
        df_pivot_geo_global = df_pivot_geo
        with tab_matriz_total:
            st.subheader(f"🌍 Resistencia de Diseño TOTAL del Terreno, R_cd (kN)")
            st.dataframe(df_pivot_geo.style.background_gradient(cmap='Greens', axis=None).format("{:.0f}"), use_container_width=True)

        df_pivot_final, df_pivot_control = df_res.pivot(index="L", columns="D", values="Q_final (kN)"), df_res.pivot(index="L", columns="D", values="Control")
        df_final_formateada = pd.DataFrame(index=df_pivot_final.index, columns=df_pivot_final.columns)
        for c in df_pivot_final.columns:
            for r in df_pivot_final.index:
                df_final_formateada.loc[r, c] = f"{df_pivot_final.loc[r, c]:.0f} [EST]" if df_pivot_control.loc[r, c] == "ESTRUCTURAL" else f"{df_pivot_final.loc[r, c]:.0f}"
        df_final_formateada.index, df_final_formateada.columns = [f"L = {idx:.1f} m" for idx in df_final_formateada.index], columnas_con_d
        
        tope_row = {col: f"{df_res[df_res['D'] == d_val]['Q_tope_est (kN)'].iloc[0]:.0f}" for d_val, col in zip(df_res['D'].unique(), columnas_con_d)}
        df_final_formateada = pd.concat([pd.DataFrame([tope_row], index=["🛑 TOPE ESTRUCT. (kN)"]), df_final_formateada])
        df_pivot_final_global = df_final_formateada
        
        with tab_matriz_tope:
            st.subheader(f"🛑 Matriz de Carga de Diseño FINAL LIMITADA (kN)")
            def style_df(df):
                styles = pd.DataFrame('', index=df.index, columns=df.columns)
                for r in df.index:
                    for c in df.columns:
                        if r == "🛑 TOPE ESTRUCT. (kN)": styles.loc[r, c] = 'background-color: #d3d3d3; font-weight: bold;'
                        elif '[EST]' in str(df.loc[r, c]): styles.loc[r, c] = 'background-color: #ffcccc;'
                        else: styles.loc[r, c] = 'background-color: #ccffcc;'
                return styles
            st.dataframe(df_final_formateada.style.apply(style_df, axis=None), use_container_width=True)
            
            df_plot_final = df_res.copy()
            df_plot_final["Diámetro"] = df_plot_final["D"].apply(lambda x: f"Ø {x:.2f} m")
            fig_final = px.line(df_plot_final, x="L", y="Q_final (kN)", color="Diámetro", markers=True, title="Curvas de Diseño Final (Intersección con Tope Estructural)", template="plotly_white")
            
            for trace in fig_final.data:
                diam_str = trace.name
                q_tope = df_plot_final[df_plot_final["Diámetro"] == diam_str]["Q_tope_est (kN)"].iloc[0]
                fig_final.add_hline(
                    y=q_tope, 
                    line_dash="dot", 
                    line_color=trace.line.color, 
                    annotation_text=f"Tope {diam_str}", 
                    annotation_position="bottom right",
                    opacity=0.6
                )
            
            st.plotly_chart(fig_final, use_container_width=True)
            st.session_state.fig_final_guardada = fig_final 

        with tab_auditoria:
            st.subheader("🔍 Auditoría (CTE DB-SE-C)")
            col_aud1, col_aud2 = st.columns(2)
            d_aud = col_aud1.selectbox("Diámetro Ø (m):", df_res['D'].unique(), format_func=lambda x: f"{x:.2f}")
            l_aud = col_aud2.selectbox("Longitud L (m):", df_res['L'].unique(), format_func=lambda x: f"{x:.2f}")
            res_aud = df_res[(df_res['D'] == d_aud) & (df_res['L'] == l_aud)].iloc[0]
            
            st.markdown(f"**Carga Final:** {res_aud['Q_final (kN)']:.0f} kN | **R_cd Punta:** {res_aud['Q_punta (kN)']/gamma_r:.0f} kN | **R_cd Fuste:** {res_aud['Q_fuste (kN)']/gamma_r:.0f} kN")
            
            if res_aud['alerta_punzonamiento']:
                st.error(f"⚠️ **Atención:** {res_aud['alerta_punzonamiento']} La resistencia en punta calculada por el bulbo ({res_aud['auditoria_punta']['q_p original ponderada (kPa)']:.0f} kPa) ha sido recortada a {res_aud['auditoria_punta']['Resist. Unitaria q_p FINAL (kPa)']:.0f} kPa para cumplir el CTE.")
            
            st.dataframe(pd.DataFrame(res_aud['auditoria_fuste']).style.format({"Long. fuste (m)": "{:.2f}", "σ'_v media (kPa)": "{:.1f}", "Resist. Unitaria τ_f (kPa)": "{:.2f}", "Fuerza Tramo (kN)": "{:.0f}"}).hide(axis="index"), use_container_width=True)
            st.dataframe(pd.DataFrame(res_aud['auditoria_bulbo']).style.format({"Espesor en bulbo (m)": "{:.2f}", "Participación (%)": "{:.1f}%", "q_p individual (kPa)": "{:.0f}"}).hide(axis="index"), use_container_width=True)

            st.markdown("---")
            st.markdown(f"#### 📐 Esquema Gráfico del Pilote (Bulbo: {res_aud['tipo_bulbo']})")
            
            D_val = res_aud['D']
            L_val = res_aud['L']
            z_sup = res_aud['z_sup_bulbo']
            z_inf = res_aud['z_inf_bulbo']
            
            fig_bulbo = go.Figure()
            z_acum_plt = 0
            for i_plt, row_plt in df_edit.iterrows():
                is_last = (i_plt == len(df_edit) - 1)
                z_next_plt = z_acum_plt + row_plt["Espesor (m)"]
                if is_last: z_next_plt = max(z_next_plt, z_inf + D_val)
                fig_bulbo.add_hrect(y0=z_acum_plt, y1=z_next_plt, fillcolor=colores[i_plt%len(colores)], opacity=0.7, line_width=1, annotation_text=row_plt["Estrato"], annotation_position="top left")
                z_acum_plt = z_next_plt

            fig_bulbo.add_shape(type="rect", x0=-D_val*1.2, x1=D_val*1.2, y0=-D_val*0.6, y1=0, fillcolor="#8B8C89", line=dict(color="black", width=2))
            fig_bulbo.add_shape(type="rect", x0=-D_val/2, x1=D_val/2, y0=0, y1=L_val, fillcolor="#A9ACA9", line=dict(color="black", width=2))
            
            if z_nulo > 0:
                fig_bulbo.add_shape(type="rect", x0=-D_val/2, x1=D_val/2, y0=0, y1=z_nulo, fillcolor="rgba(255, 255, 0, 0.4)", line=dict(color="orange", width=2, dash="dash"))
                fig_bulbo.add_annotation(x=D_val*0.8, y=z_nulo/2, text="Fuste<br>Nulo", showarrow=False, xanchor="left", font=dict(color="orange"))

            fig_bulbo.add_shape(type="rect", x0=-D_val/2, x1=-D_val/4, y0=0, y1=L_val, fillcolor="black", opacity=0.15, line_width=0)
            fig_bulbo.add_shape(type="rect", x0=D_val/4, x1=D_val/2, y0=0, y1=L_val, fillcolor="white", opacity=0.25, line_width=0)
            fig_bulbo.add_shape(type="line", x0=-D_val/2, x1=D_val/2, y0=L_val, y1=L_val, line=dict(color="black", width=4))

            fig_bulbo.add_shape(type="rect", x0=-D_val*1.5, x1=D_val*1.5, y0=z_sup, y1=z_inf, fillcolor="rgba(255, 0, 0, 0.15)", line=dict(color="red", width=3, dash="dash"))
            fig_bulbo.add_annotation(x=D_val*1.6, y=(z_sup+L_val)/2, text=f"6D Arriba<br>({L_val - z_sup:.1f} m)", showarrow=False, xanchor="left", font=dict(color="red", size=13, family="Arial Black"))
            fig_bulbo.add_annotation(x=D_val*1.6, y=(L_val+z_inf)/2, text=f"3D Abajo<br>({z_inf - L_val:.1f} m)", showarrow=False, xanchor="left", font=dict(color="red", size=13, family="Arial Black"))
            fig_bulbo.add_hline(y=L_val, line_color="black", line_width=2, opacity=0.8)
            fig_bulbo.add_hline(y=zw, line_dash="dot", line_color="blue", annotation_text="N.F.", annotation_position="bottom right")

            fig_bulbo.update_yaxes(autorange="reversed", title="Profundidad (m)", range=[max(z_max_total, z_inf + 2*D_val), -D_val], showgrid=True)
            fig_bulbo.update_xaxes(showticklabels=False, range=[-D_val*4, D_val*4], showgrid=False)
            fig_bulbo.update_layout(height=600, margin=dict(l=20, r=20, t=20, b=20), paper_bgcolor="white", plot_bgcolor="white")
            st.plotly_chart(fig_bulbo, use_container_width=True)

        with tab_formulacion:
            st.subheader("📖 Ecuaciones y Topes (CTE DB-SE-C)")
            st.markdown(f"**Tope Estructural Aplicado:** {sigma_tope_mpa:.2f} MPa")
            st.markdown(f"*(Definido por procedimiento: {desc_tope})*")
            st.markdown("### 1. Punta en Suelos Granulares (F.30)")
            st.markdown(r"$q_p = f_p \cdot \sigma'_{vp} \cdot N_q \le 20 \text{ MPa}$")
            st.markdown(f"*En este cálculo: $f_p = {fp_val}$. La evaluación de la punta se promedia ponderadamente en un entorno de **6D por encima y 3D por debajo**.*")
            st.markdown("### 2. Fuste en Suelos Granulares (F.31)")
            st.markdown(r"$\tau_f = \sigma'_v \cdot K_f \cdot f \cdot \tan\phi \le 120 \text{ kPa}$")
            st.markdown(f"*En este cálculo: $K_f = {Kf_val}$ y $f = {f_rug}$*")
            st.markdown("### 3. Punta en Suelos Finos (F.32)")
            st.markdown(r"$q_p = N_p \cdot c_u$ (Con $N_p = 9$)")
            st.markdown("### 4. Punzonamiento de Estrato Inferior (Apoyo CTE 5.2.2)")
            st.markdown(r"$q_p \le 6 \cdot \left( 1 + \frac{H}{D} \right)^2 \cdot c_u$")
            st.markdown("*Si bajo la punta existe una capa de suelo blando cohesivo a una distancia H, la resistencia de punta se recorta a este valor máximo para evitar la rotura por punzonamiento y garantizar el cumplimiento del Art. 5.2.2 del Código Técnico.*")
            st.markdown("### 5. Fuste en Suelos Finos (F.33)")
            st.markdown(r"$\tau_f = \frac{100 \cdot c_u}{100 + c_u}$")

# ══════════════════════════════════════════════════════════════════════════
# GENERADOR DEL INFORME WORD CTE ENRIQUECIDO CON FOTOGRAFÍAS KALEIDO
# ══════════════════════════════════════════════════════════════════════════
def generar_word_cte(df_estratos, fig_tens, df_pivot_geo, df_pivot_final, gamma_val, z_nul, zw_val, sigma_tope, fig_final, desc_t):
    doc = Document()
    estilo_tabla = 'Light Grid Accent 1'
    
    doc.add_paragraph('\n\n\n\n')
    title = doc.add_paragraph('ANEJO DE CÁLCULO: CIMENTACIONES PROFUNDAS')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs: run.font.size, run.font.bold = Pt(24), True
    subtitle = doc.add_paragraph('Diseño Analítico de Pilotes según CTE DB-SE-C')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs: run.font.size = Pt(16)
        
    doc.add_paragraph('\n\n\n')
    doc.add_paragraph(f'Fecha: {date.today().strftime("%d/%m/%Y")}').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
    
    doc.add_heading('1. Bases de Cálculo y Parámetros Iniciales', level=1)
    p_bases = doc.add_paragraph()
    p_bases.add_run(f'• Descripción del Pilote (Tabla 5.1): ').bold = True
    p_bases.add_run(f'{desc_t}\n')
    p_bases.add_run(f'• Tope Estructural Aplicado: ').bold = True
    p_bases.add_run(f'{sigma_tope:.2f} MPa\n')
    p_bases.add_run(f'• Nivel Freático: ').bold = True
    p_bases.add_run(f'Z = {zw_val:.2f} m\n')
    p_bases.add_run(f'• Zonas Excluidas: ').bold = True
    p_bases.add_run(f'Se descuenta rozamiento en los primeros {z_nul:.2f} m de fuste.\n')
    p_bases.add_run(f'• Factor de Seguridad (Estados Límite): ').bold = True
    p_bases.add_run(f'Coeficiente Parcial de Resistencia γ_R = {gamma_val:.2f}\n')

    doc.add_heading('2. Perfil Geotécnico y Tensional', level=1)
    doc.add_paragraph('Definición de las unidades geotécnicas y sus propiedades resistentes características:')
    
    tabla_est = doc.add_table(rows=1, cols=len(df_estratos.columns))
    tabla_est.style = estilo_tabla
    for i, col in enumerate(df_estratos.columns): tabla_est.rows[0].cells[i].text = str(col)
    for _, row in df_estratos.iterrows():
        row_cells = tabla_est.add_row().cells
        for i, val in enumerate(row): 
            if isinstance(val, float): row_cells[i].text = f"{val:.2f}"
            else: row_cells[i].text = str(val)

    if fig_tens is not None:
        doc.add_paragraph('\nEsquema de tensiones verticales del terreno (totales, efectivas e intersticiales):')
        # --- CONTROL DE ERRORES PARA KALEIDO (GRÁFICA 1) ---
        try:
            img_tens = fig_tens.to_image(format="png", width=700, height=450)
            doc.add_picture(io.BytesIO(img_tens), width=Inches(6.0))
        except Exception as e:
            p_err = doc.add_paragraph(f"⚠️ AVISO: No se pudo insertar la gráfica. Para exportar imágenes es obligatorio instalar 'kaleido' (ejecuta: pip install kaleido en la terminal). \nDetalle del error: {e}")
            p_err.runs[0].font.color.rgb = RGBColor(255, 0, 0) 

    doc.add_heading('3. Metodología Analítica (CTE DB-SE-C)', level=1)
    p_metodo = doc.add_paragraph()
    p_metodo.add_run('Resistencia por Punta (qp):\n').bold = True
    p_metodo.add_run('• Suelos Finos (Corto Plazo): qp = Np · cu.\n')
    p_metodo.add_run('• Suelos Granulares (Largo Plazo): qp = fp · σ\'vp · Nq ≤ 20 MPa.\n')
    p_metodo.add_run('• Punzonamiento (Justificación CTE 5.2.2): Si existen estratos arcillosos por debajo de la punta, la resistencia máxima se evalúa mediante formulaciones complementarias acotando el valor a qp ≤ 6(1+H/D)²·cu.\n')
    p_metodo.add_run('NOTA: El entorno de evaluación general de la punta se promedia ponderadamente en un bulbo de 6 diámetros por encima y 3 diámetros por debajo para todo tipo de suelos, respetando los topes por punzonamiento inferiores si los hubiese.\n\n')
    
    p_metodo.add_run('Resistencia por Fuste (τf):\n').bold = True
    p_metodo.add_run('• Corto Plazo: τf = 100·cu / (100 + cu). Afectado por coef. reductor 0.8 si el fuste es de acero.\n')
    p_metodo.add_run('• Largo Plazo: τf = σ\'v · Kf · f · tg(φ) ≤ 120 kPa.\n\n')
    
    doc.add_heading('4. Matriz de Capacidad Geotécnica Estricta (kN)', level=1)
    doc.add_paragraph('Valores de hundimiento geotécnico minorados por el Coeficiente Parcial de Resistencia (sin limitar por tope estructural).')
    tabla_geo = doc.add_table(rows=1, cols=len(df_pivot_geo.columns) + 1)
    tabla_geo.style = estilo_tabla
    hdr_geo = tabla_geo.rows[0].cells
    hdr_geo[0].text = "L / Ø"
    for i, col_name in enumerate(df_pivot_geo.columns): hdr_geo[i+1].text = str(col_name)
    for index, row in df_pivot_geo.iterrows():
        row_cells = tabla_geo.add_row().cells
        row_cells[0].text = str(index)
        for i, val in enumerate(row): row_cells[i+1].text = f"{val:.0f}"

    doc.add_heading('5. Matriz de Diseño Final Limitada (kN)', level=1)
    doc.add_paragraph('Valores finales de diseño. Si la carga está limitada por la resistencia del material del pilote, se indica con la etiqueta [EST].')
    tabla_fin = doc.add_table(rows=1, cols=len(df_pivot_final.columns) + 1)
    tabla_fin.style = estilo_tabla
    hdr_fin = tabla_fin.rows[0].cells
    hdr_fin[0].text = "L / Ø"
    for i, col_name in enumerate(df_pivot_final.columns): hdr_fin[i+1].text = str(col_name)
    for index, row in df_pivot_final.iterrows():
        row_cells = tabla_fin.add_row().cells
        row_cells[0].text = str(index)
        for i, val in enumerate(row): row_cells[i+1].text = str(val)

    if fig_final is not None:
        doc.add_paragraph('\n')
        doc.add_heading('Gráfico: Curvas de Diseño Final', level=2)
        # --- CONTROL DE ERRORES PARA KALEIDO (GRÁFICA 2) ---
        try:
            img_bytes = fig_final.to_image(format="png", width=800, height=500)
            doc.add_picture(io.BytesIO(img_bytes), width=Inches(6.5))
        except Exception as e:
            p_err2 = doc.add_paragraph(f"⚠️ AVISO: No se pudo insertar la gráfica. Para exportar imágenes es obligatorio instalar 'kaleido' (ejecuta: pip install kaleido en la terminal). \nDetalle del error: {e}")
            p_err2.runs[0].font.color.rgb = RGBColor(255, 0, 0) 

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ══════════════════════════════════════════════════════════════════════════
# INTERFAZ DE DESCARGA
# ══════════════════════════════════════════════════════════════════════════
if st.session_state.calculado:
    st.sidebar.markdown("---")
    st.sidebar.subheader("📄 Memoria de Cálculo")
    
    if st.sidebar.button("🛠️ Generar Informe (.docx)", type="primary", use_container_width=True):
        if not df_res.empty:
            with st.spinner("Generando documento Word..."):
                buffer = generar_word_cte(
                    df_edit, fig_tens, df_pivot_geo_global, df_pivot_final_global, 
                    gamma_r, z_nulo, zw, sigma_tope_mpa, st.session_state.fig_final_guardada, desc_tope
                )
                st.session_state.word_buffer = buffer
            st.sidebar.success("✅ ¡Informe ejecutivo generado!")

    if st.session_state.word_buffer is not None:
        st.sidebar.download_button(
            label="⬇️ Descargar Informe", 
            data=st.session_state.word_buffer, 
            file_name="Anejo_Pilotes_CTE.docx", 
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
            use_container_width=True
        )