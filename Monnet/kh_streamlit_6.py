# calculadora de Kh 
# Ábaco de Chadeisson
# Método de Granados https://share.google/VcqCQNKYtkJAQnOXR
# Método de iterpolación mediante las lineas del ábaco

# Desarrollado por Germán López Pineda
# ICCP, MIT y MMC
# Bajo criterios del ITQ



import streamlit as st
import pandas as pd
import numpy as np
import math
import matplotlib.pyplot as plt
from io import BytesIO
from docx import Document
from docx.shared import Inches

# --- CONFIGURACIÓN DE PÁGINA ---
st.set_page_config(page_title="Estimación de Kh Ábaco de Chadeisson", layout="wide")

# --- CONSTANTES DE CONVERSIÓN ---
KPA_TO_TM2 = 1.0 / 9.80665
TM3_TO_KNM3 = 9.80665

# --- FUNCIONES DE FORMATO ESPAÑOL (Punto para miles, Coma para decimales) ---
def formato_espanol(x):
    """Convierte un número a formato 15.000,00"""
    return f"{x:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

def formato_porcentaje_es(x):
    """Convierte un porcentaje a formato 15,00%"""
    if x == 0: return "-"
    return f"{x:.2f}".replace(".", ",") + "%"


# =====================================================================
# Puntos de las rectas del ábaco (Datos Empíricos)
# =====================================================================
datos_usuario = [
    {'Kh': 700,   'c1': 0,    'phi1': 7.0,  'c2': 1.64, 'phi2': 0.0},
    {'Kh': 800,   'c1': 0,    'phi1': 10.0, 'c2': 2.47, 'phi2': 0.0},
    {'Kh': 900,   'c1': 0,    'phi1': 12.0, 'c2': 3.19, 'phi2': 0.0},
    {'Kh': 1000,  'c1': 0,    'phi1': 14.0, 'c2': 4.0,  'phi2': 0.0},
    {'Kh': 1100,  'c1': 0,    'phi1': 16.0, 'c2': 4.84, 'phi2': 0.0},
    {'Kh': 1200,  'c1': 0,    'phi1': 18.0, 'c2': 5.69, 'phi2': 0.0},
    {'Kh': 1300,  'c1': 0,    'phi1': 19.2, 'c2': 6.46, 'phi2': 0.0},
    {'Kh': 1400,  'c1': 0,    'phi1': 20.0, 'c2': 7.29, 'phi2': 0.0},
    {'Kh': 1500,  'c1': 0,    'phi1': 21.5, 'c2': 8.0,  'phi2': 0.0},
    {'Kh': 2000,  'c1': 0,    'phi1': 26.0, 'c2': 9.0,  'phi2': 7.0},
    {'Kh': 3000,  'c1': 0,    'phi1': 31.5, 'c2': 9.0,  'phi2': 18.0},
    {'Kh': 4000,  'c1': 0,    'phi1': 35.0, 'c2': 9.0,  'phi2': 24.74},
    {'Kh': 5000,  'c1': 0,    'phi1': 38.0, 'c2': 9.0,  'phi2': 28.4},
    {'Kh': 6000,  'c1': 0,    'phi1': 40.0, 'c2': 9.0,  'phi2': 31.5},
    {'Kh': 7000,  'c1': 0,    'phi1': 41.5, 'c2': 9.0,  'phi2': 34.0},
    {'Kh': 8000,  'c1': 0,    'phi1': 43.0, 'c2': 9.0,  'phi2': 36.0},
    {'Kh': 9000,  'c1': 0,    'phi1': 44.0, 'c2': 9.0,  'phi2': 38.0},
    {'Kh': 10000, 'c1': 0,    'phi1': 45.0, 'c2': 9.0,  'phi2': 39.2},
    {'Kh': 12000, 'c1': 3.17, 'phi1': 45.0, 'c2': 9.0,  'phi2': 41.5},
    {'Kh': 14000, 'c1': 6.0,  'phi1': 45.0, 'c2': 9.0,  'phi2': 43.2},
    {'Kh': 16000, 'c1': 8.38, 'phi1': 46.0, 'c2': 9.0,  'phi2': 45.0}
]

rectas_chadeisson_empiricas = {}
for fila in datos_usuario:
    kh = fila['Kh']
    m = (fila['phi2'] - fila['phi1']) / (fila['c2'] - fila['c1'])
    n = fila['phi1'] - (m * fila['c1'])
    rectas_chadeisson_empiricas[kh] = (m, n)

# =====================================================================
# MOTORES DE CÁLCULO
# =====================================================================
def calc_chadeisson_granados_tm3(phi_d, c_tm2):
    phi_rad = math.radians(phi_d)
    a = {1: 0.2780567713112, 2: 5.57961020316, 3: -35.301673725713, 4: 147.66061817668381, 5: -245.6043814457953, 6: 164.0109185222536, 7: 0.1188816874005, 8: 0.3480332606123, 9: -1.0388796245679, 10: 3.6912996582687, 11: -5.0917695015669, 12: 3.657829914242, 13: 0.0000376350382, 14: -0.0084577071902, 15: 0.1117470768513, 16: -0.4859679131769, 17: 0.8440710781142, 18: -0.5066994114313}
    kh_sum = sum(a[1 + j + 6 * k] * (phi_rad**j) * (c_tm2**k) for k in range(3) for j in range(6))
    return max(kh_sum * 1000.0, 0.0)

def calc_chadeisson_granados(phi_d, c_kpa):
    return calc_chadeisson_granados_tm3(phi_d, c_kpa * KPA_TO_TM2) * TM3_TO_KNM3

def calc_chadeisson_geometrico_tm3(phi_d, c_tm2):
    alturas_phi = sorted([(m * c_tm2 + n, kh) for kh, (m, n) in rectas_chadeisson_empiricas.items()])
    if phi_d < alturas_phi[0][0]: return 0.0 
    if phi_d >= alturas_phi[-1][0]: return alturas_phi[-1][1]
    for i in range(len(alturas_phi) - 1):
        if alturas_phi[i][0] <= phi_d <= alturas_phi[i+1][0]:
            if alturas_phi[i+1][0] == alturas_phi[i][0]: return alturas_phi[i][1]
            peso = (phi_d - alturas_phi[i][0]) / (alturas_phi[i+1][0] - alturas_phi[i][0])
            return alturas_phi[i][1] + peso * (alturas_phi[i+1][1] - alturas_phi[i][1])
    return 0.0

def calc_chadeisson_geometrico(phi_d, c_kpa):
    return calc_chadeisson_geometrico_tm3(phi_d, c_kpa * KPA_TO_TM2) * TM3_TO_KNM3

# =====================================================================
# FUNCIONES DE DIBUJO (Reutilizables)
# =====================================================================
c_vals = np.linspace(0, 9, 250)
phi_vals = np.linspace(0, 45, 250)
C_grid, Phi_grid = np.meshgrid(c_vals, phi_vals)
niveles_completos = list(rectas_chadeisson_empiricas.keys())

Kh_grid_granados = np.zeros_like(C_grid)
Kh_grid_geom = np.zeros_like(C_grid)
for i in range(C_grid.shape[0]):
    for j in range(C_grid.shape[1]):
        Kh_grid_granados[i, j] = calc_chadeisson_granados_tm3(Phi_grid[i, j], C_grid[i, j])
        Kh_grid_geom[i, j] = calc_chadeisson_geometrico_tm3(Phi_grid[i, j], C_grid[i, j])

def formatear_ejes(ax, df_input):
    ax.set_xlim(0, 9); ax.set_ylim(0, 45)
    ax.set_xticks(np.arange(0, 10, 1)); ax.set_yticks(np.arange(0, 50, 5))
    ax.xaxis.grid(True, linestyle='-', color='#e0e0e0', linewidth=0.5)
    ax.yaxis.grid(True, linestyle='--', color='#e0e0e0', linewidth=0.5)
    ax.set_xlabel('C (t/m²) (cohesión)', fontsize=9)
    ax.set_ylabel('Ángulo de rozamiento (φ)', fontsize=9)
    ax.tick_params(axis='both', which='major', labelsize=8)
    
    m_700, n_700 = rectas_chadeisson_empiricas[700]
    corte_x = -n_700 / m_700 
    ax.add_patch(plt.Polygon([[0,0], [corte_x, 0], [0, n_700]], color='gray', alpha=0.15, hatch='///'))
    
    if not df_input.empty:
        for _, row in df_input.iterrows():
            phi_p = row['phi [°]']; c_p = row['c [kPa]'] * KPA_TO_TM2
            if 0 <= phi_p <= 45 and 0 <= c_p <= 9:
                ax.scatter(c_p, phi_p, color='#e74c3c', s=45, edgecolors='black', zorder=5)
                ax.annotate(row['Estrato'], (c_p, phi_p), xytext=(6, 6), textcoords='offset points', fontsize=8, bbox=dict(boxstyle="round,pad=0.2", fc="white", ec="gray", alpha=0.8), zorder=6)

def dibujar_granados(ax, df_input):
    contornos_g = ax.contour(C_grid, Phi_grid, Kh_grid_granados, levels=niveles_completos, colors='#2b59c3', linewidths=1.0, alpha=0.8)
    ax.clabel(contornos_g, inline=True, fontsize=7, fmt='%1.0f')
    ax.set_title("Modelo Granados ", fontweight='bold', fontsize=10)
    formatear_ejes(ax, df_input)

def dibujar_geometrico(ax, df_input):
    contornos_m = ax.contour(C_grid, Phi_grid, Kh_grid_geom, levels=niveles_completos, colors='#e67e22', linewidths=1.0, alpha=0.9)
    ax.clabel(contornos_m, inline=True, fontsize=7, fmt='%1.0f')
    ax.set_title("Modelo Geométrico Empírico", fontweight='bold', fontsize=10)
    formatear_ejes(ax, df_input)


# =====================================================================
# INTERFAZ STREAMLIT
# =====================================================================
st.title("🧮 Calculadora de Kh (Ábaco de Chadeisson)")

tab_calc, tab_abacos = st.tabs(["📝 Tabla de Cálculo", "📊 Gráfica"])

with tab_calc:
    st.markdown("Calculadora basada en el método de Granados y en el método empírico.")
    
    df_init = pd.DataFrame([
        {"Estrato": "UG-01", "phi [°]": 15.0, "c [kPa]": 30.0},
        {"Estrato": "UG-02", "phi [°]": 35.0, "c [kPa]": 0.0},
        {"Estrato": "UG-03", "phi [°]": 20.0, "c [kPa]": 18}
    ])

    df_input = st.data_editor(
        df_init, num_rows="dynamic", use_container_width=True,
        column_config={
            "Estrato": st.column_config.TextColumn("Nombre del estrato"),
            "phi [°]": st.column_config.NumberColumn("Áng. Rozamiento φ (°)", min_value=0.0, max_value=45.0, step=0.1, format="%.2f"),
            "c [kPa]": st.column_config.NumberColumn("Cohesión c' (kPa)", min_value=0.0, step=0.1, format="%.2f")
        }
    )

    if not df_input.empty:
        df_input['kh Granados [kN/m³]'] = df_input.apply(lambda r: calc_chadeisson_granados(r['phi [°]'], r['c [kPa]']), axis=1)
        df_input['kh Geométrico [kN/m³]'] = df_input.apply(lambda r: calc_chadeisson_geometrico(r['phi [°]'], r['c [kPa]']), axis=1)
        
        df_input['Desviación (%)'] = np.where(df_input['kh Geométrico [kN/m³]'] > 0, abs(df_input['kh Granados [kN/m³]'] - df_input['kh Geométrico [kN/m³]']) / df_input['kh Geométrico [kN/m³]'] * 100, 0.0)
        
        # --- TABLA EN STREAMLIT: APLICANDO FORMATO ESPAÑOL A TODAS LAS COLUMNAS ---
        st.dataframe(
            df_input.style.format({
                'phi [°]': formato_espanol,
                'c [kPa]': formato_espanol,
                'kh Granados [kN/m³]': lambda x: "⚠️ Suelos muy blandos" if x < 1 else formato_espanol(x), 
                'kh Geométrico [kN/m³]': lambda x: "⚠️ Suelos muy blandos" if x == 0 else formato_espanol(x), 
                'Desviación (%)': formato_porcentaje_es
            }).background_gradient(subset=['Desviación (%)'], cmap='OrRd', vmin=0, vmax=15), 
            use_container_width=True
        )

        # --- GENERADOR DE INFORME WORD (.DOCX) ---
        st.markdown("---")
        st.subheader("📄 Exportar Resultados")
        
        if st.button("Preparar Informe Word (.docx)", type="primary"):
            with st.spinner("Compilando documento Word..."):
                doc = Document()
                doc.add_heading('Estimación de Kh (Ábaco de Chadeisson)', 0)
                
                # 1. Añadir Tabla
                doc.add_heading('1. Datos de Entrada y Valores Resumen', level=1)
                
                # Preparamos el DataFrame para Word aplicando las funciones de formato español
                df_report = df_input.copy()
                df_report['phi [°]'] = df_report['phi [°]'].apply(formato_espanol)
                df_report['c [kPa]'] = df_report['c [kPa]'].apply(formato_espanol)
                df_report['kh Granados [kN/m³]'] = df_report['kh Granados [kN/m³]'].apply(lambda x: "Suelo muy blando" if x < 1 else formato_espanol(x))
                df_report['kh Geométrico [kN/m³]'] = df_report['kh Geométrico [kN/m³]'].apply(lambda x: "Suelo muy blando" if x == 0 else formato_espanol(x))
                df_report['Desviación (%)'] = df_report['Desviación (%)'].apply(formato_porcentaje_es)
                
                t = doc.add_table(rows=1, cols=len(df_report.columns))
                t.style = 'Table Grid'
                hdr_cells = t.rows[0].cells
                for i, col_name in enumerate(df_report.columns):
                    hdr_cells[i].text = col_name
                for index, row in df_report.iterrows():
                    row_cells = t.add_row().cells
                    for i, val in enumerate(row):
                        row_cells[i].text = str(val)

                # 2. Añadir Gráficas Individuales
                doc.add_heading('2. Gráficas Individuales', level=1)
                
                fig_g, ax_g = plt.subplots(figsize=(7, 7))
                dibujar_granados(ax_g, df_input)
                buf_g = BytesIO()
                fig_g.savefig(buf_g, format='png', bbox_inches='tight', dpi=150)
                buf_g.seek(0)
                doc.add_picture(buf_g, width=Inches(5.5))
                doc.add_paragraph("Figura 1. Interpolación mediante Polinomio de Granados.")
                
                fig_m, ax_m = plt.subplots(figsize=(7, 7))
                dibujar_geometrico(ax_m, df_input)
                buf_m = BytesIO()
                fig_m.savefig(buf_m, format='png', bbox_inches='tight', dpi=150)
                buf_m.seek(0)
                doc.add_picture(buf_m, width=Inches(5.5))
                doc.add_paragraph("Figura 2. Interpolación mediante Modelo Geométrico Empírico.")
                
                # 3. Añadir Gráfica Conjunta
                doc.add_heading('3. Comparativa Conjunta', level=1)
                fig_j, (ax_j1, ax_j2) = plt.subplots(1, 2, figsize=(10, 6))
                dibujar_granados(ax_j1, df_input)
                dibujar_geometrico(ax_j2, df_input)
                buf_j = BytesIO()
                fig_j.savefig(buf_j, format='png', bbox_inches='tight', dpi=150)
                buf_j.seek(0)
                doc.add_picture(buf_j, width=Inches(6.0))
                doc.add_paragraph("Figura 3. Comparativa visual de ambos métodos.")
                
                doc_io = BytesIO()
                doc.save(doc_io)
                doc_io.seek(0)
                
                plt.close('all')
                
                st.success("¡Informe preparado con éxito!")
                st.download_button(
                    label="📥 Descargar Informe en Word (.docx)",
                    data=doc_io,
                    file_name="Informe_Chadeisson.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary"
                )

with tab_abacos:
    st.subheader("Comparativa de resultados")
    fig_pantalla, (ax1, ax2) = plt.subplots(1, 2, figsize=(8, 4.5))
    
    dibujar_granados(ax1, df_input)
    dibujar_geometrico(ax2, df_input)
    
    plt.tight_layout()
    st.pyplot(fig_pantalla)