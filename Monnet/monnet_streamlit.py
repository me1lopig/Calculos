import streamlit as st
import numpy as np
import pandas as pd
import io 
from docx import Document
from scipy.interpolate import griddata

# --- INICIALIZACIÓN DE LA MEMORIA DE LA APP ---
if 'documentacion_generada' not in st.session_state:
    st.session_state.documentacion_generada = False

if 'material' not in st.session_state: st.session_state['material'] = "HA-30 (E ≈ 29 GPa)"
if 'modulo_e' not in st.session_state: st.session_state['modulo_e'] = 29000000.0
if 'espesor' not in st.session_state: st.session_state['espesor'] = 0.60
if 'dro' not in st.session_state: st.session_state['dro'] = 0.015
if 'c0' not in st.session_state: st.session_state['c0'] = 30.0
if 'ap' not in st.session_state: st.session_state['ap'] = 10.0

# --- BASE DE DATOS DEL ÁBACO ORIGINAL DE CHADEISSON ---
puntos_chadeisson = np.array([
    [14, 0], [10, 16], [4, 40], [0, 58],                            # Curva 10.000
    [24, 0], [18.5, 20], [14, 38], [11, 50], [5, 80],               # Curva 20.000
    [30, 0], [25, 20], [21, 40], [17, 60], [13, 85],                # Curva 30.000
    [35, 0], [30, 22], [26, 40], [23, 60], [18, 90],                # Curva 40.000
    [38, 0], [34, 20], [31, 40], [27, 65], [23, 90]                 # Curva 50.000
])
valores_kh_chadeisson = np.array([
    10000, 10000, 10000, 10000,
    20000, 20000, 20000, 20000, 20000,
    30000, 30000, 30000, 30000, 30000,
    40000, 40000, 40000, 40000, 40000,
    50000, 50000, 50000, 50000, 50000
])

# --- FUNCIONES CALLBACK ---
def update_e():
    mat = st.session_state['material']
    if "HA-25" in mat: st.session_state['modulo_e'] = 27000000.0
    elif "HA-30" in mat: st.session_state['modulo_e'] = 29000000.0
    elif "HA-35" in mat: st.session_state['modulo_e'] = 30000000.0

def cargar_muro_chadeisson():
    """Carga los datos de la pantalla original con la que se hizo el ábaco de los años 60"""
    st.session_state['material'] = "Otro (Manual)"
    st.session_state['modulo_e'] = 20000000.0  
    st.session_state['espesor'] = 0.80         
    st.session_state['dro'] = 0.015
    st.session_state['c0'] = 30.0
    st.session_state['ap'] = 10.0

# Configuración de la página
st.set_page_config(page_title="Balasto Pantallas: Monnet vs Chadeisson", layout="wide")

st.title("🏗️ Coeficiente de Balasto ($K_h$) - Monnet vs. Chadeisson")

# --- BARRA LATERAL ---
st.sidebar.header("1. Propiedades de la Pantalla")

st.sidebar.button(
    "🎯 Cargar Muro Chadeisson (Original)", 
    on_click=cargar_muro_chadeisson, 
    type="secondary",
    help="Carga los parámetros del muro original (E=20 GPa, e=0.80 m) para igualar la fórmula al ábaco."
)

tipo_hormigon = st.sidebar.selectbox(
    "Tipo de Hormigón:", 
    ["HA-25 (E ≈ 27 GPa)", "HA-30 (E ≈ 29 GPa)", "HA-35 (E ≈ 30 GPa)", "Otro (Manual)"],
    key='material',
    on_change=update_e
)

E = st.sidebar.number_input("Módulo de Elasticidad E (kPa):", step=1000000.0, format="%.0f", key='modulo_e')
espesor = st.sidebar.number_input("Espesor del muro (m):", step=0.05, min_value=0.30, key='espesor')

I = (1.0 * espesor**3) / 12.0
EI = E * I
st.sidebar.info(f"**Inercia (I):** {I:.5f} m⁴/m")
st.sidebar.success(f"**Rigidez a flexión (EI):** {EI:,.0f} kN·m²/m")

st.sidebar.header("2. Parámetros Monnet (Fórmula)")
with st.sidebar.expander("⚙️ Configuración Avanzada"):
    dro = st.number_input("Desplazamiento ref. d_ro (m)", format="%.3f", key='dro')
    c0 = st.number_input("Cohesión ref. c_0 (kPa)", key='c0')
    Ap = st.number_input("Coef. ajuste A_p", key='ap')


# ==========================================
# CREACIÓN DE PESTAÑAS (TABS)
# ==========================================
tab_calculadora, tab_teoria = st.tabs(["🧮 Calculadora Comparativa", "📚 Formulación y Teoría"])

# --- PESTAÑA 1: LA CALCULADORA ---
with tab_calculadora:
    st.header("Estratigrafía del Terreno")
    st.write("Edita la tabla inferior. El cálculo base se actualiza al instante utilizando ambos métodos.")

    datos_iniciales = {
        "Capa": ["Rellenos", "Arcilla Firme", "Arena Densa"],
        "Profundidad Base (m)": [2.0, 6.0, 12.0],
        "Peso Específico γ (kN/m³)": [18.0, 20.0, 21.0],
        "Ángulo Rozamiento φ (º)": [25.0, 15.0, 35.0],
        "Cohesión c (kPa)": [5.0, 40.0, 0.0]
    }
    df_editado = st.data_editor(pd.DataFrame(datos_iniciales), num_rows="dynamic", use_container_width=True)

    resultados = df_editado.copy()
    ka_lista, k0_lista, kp_lista = [], [], []
    kh_monnet_lista, kh_chadeisson_lista = [], []

    for index, row in resultados.iterrows():
        gamma, phi_deg, c = row["Peso Específico γ (kN/m³)"], row["Ángulo Rozamiento φ (º)"], row["Cohesión c (kPa)"]
        phi_rad = np.radians(phi_deg)
        
        # 1. Empujes de Rankine
        Ka = (1 - np.sin(phi_rad)) / (1 + np.sin(phi_rad))
        K0 = 1 - np.sin(phi_rad)
        Kp = (1 + np.sin(phi_rad)) / (1 - np.sin(phi_rad))
        
        ka_lista.append(round(Ka, 3))
        k0_lista.append(round(K0, 3))
        kp_lista.append(round(Kp, 3))
        
        # 2. Método Analítico: MONNET
        termino_friccion = ((20 * EI * gamma * (Kp - K0)) / (dro**4))**(1/5)
        termino_cohesion = (Ap * c * np.tanh(c / c0)) / dro
        kh_monnet_lista.append(round(termino_friccion + termino_cohesion, 2))

        # 3. Método Empírico: CHADEISSON (Interpolación del dibujo)
        try:
            kh_int = griddata(puntos_chadeisson, valores_kh_chadeisson, (phi_deg, c), method='cubic')
            if np.isnan(kh_int):
                kh_int = griddata(puntos_chadeisson, valores_kh_chadeisson, (phi_deg, c), method='nearest')
            kh_chadeisson_lista.append(round(float(kh_int), 2))
        except Exception:
            kh_chadeisson_lista.append(0.0)

    resultados["K_a"] = ka_lista
    resultados["K_0"] = k0_lista
    resultados["K_p"] = kp_lista
    resultados["K_h Monnet (kN/m³)"] = kh_monnet_lista
    resultados["K_h Chadeisson (kN/m³)"] = kh_chadeisson_lista

    st.subheader("📊 Resultados Numéricos Comparativos")
    columnas_finales = ["Capa", "Profundidad Base (m)", "K_a", "K_0", "K_p", "K_h Monnet (kN/m³)", "K_h Chadeisson (kN/m³)"]
    st.dataframe(resultados[columnas_finales], use_container_width=True)

    st.divider()
    st.header("📄 Exportación de Resultados")

    if st.button("Generar Documentación", type="primary"):
        st.session_state.documentacion_generada = True
        st.success("¡Documentación generada con éxito!")

    if st.session_state.documentacion_generada:
        col_dl1, col_dl2 = st.columns(2)

        buffer_excel = io.BytesIO()
        with pd.ExcelWriter(buffer_excel, engine='openpyxl') as writer:
            resultados[columnas_finales].to_excel(writer, index=False, sheet_name='Balasto_Comparativo')

        with col_dl1:
            st.download_button("📊 Descargar Tabla en Excel", data=buffer_excel.getvalue(), file_name="Balasto_Comparativo.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

        doc = Document()
        doc.add_heading('MEMORIA DE CÁLCULO: BALASTO EN PANTALLAS', 0)
        doc.add_heading('1. Geometría y Material (Solo aplica a Monnet)', level=1)
        doc.add_paragraph(f"Material: {st.session_state['material']}", style='List Bullet')
        doc.add_paragraph(f"Espesor de pantalla: {espesor:.2f} m", style='List Bullet')
        doc.add_paragraph(f"Módulo de Elasticidad (E): {E:,.0f} kPa", style='List Bullet')
        doc.add_paragraph(f"Rigidez a flexión (EI): {EI:,.0f} kN·m²/m", style='List Bullet')

        doc.add_heading('2. Cuadro Comparativo de Estratos', level=1)
        doc.add_paragraph("Se contrastan los resultados de la fórmula analítica (interacción suelo-estructura real) frente a los obtenidos por lectura directa del ábaco original de Chadeisson.")
        
        tabla_word = doc.add_table(rows=1, cols=len(columnas_finales))
        tabla_word.style = 'Table Grid'
        
        hdr_cells = tabla_word.rows[0].cells
        for i, nombre_col in enumerate(columnas_finales):
            hdr_cells[i].text = nombre_col

        for index, row in resultados.iterrows():
            row_cells = tabla_word.add_row().cells
            for i, nombre_col in enumerate(columnas_finales):
                row_cells[i].text = str(row[nombre_col])

        buffer_word = io.BytesIO()
        doc.save(buffer_word)
        buffer_word.seek(0)

        with col_dl2:
            st.download_button("📝 Descargar Memoria (Word)", data=buffer_word, file_name="Memoria_Balasto_Comparativa.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# --- PESTAÑA 2: LA TEORÍA Y FORMULACIÓN ---
with tab_teoria:
    st.header("Formulación Analítica y Empírica Empleada")
    
    st.markdown("""
    Esta herramienta calcula el coeficiente de interacción suelo-estructura utilizando simultáneamente dos aproximaciones clásicas de la geotecnia, permitiendo contrastar la metodología analítica moderna con la bibliografía empírica original.
    """)
    
    st.subheader("1. Método Analítico parametrizado (Fórmula de Monnet)")
    st.markdown("A diferencia del ábaco estático, esta expresión analítica desarrollada por A. Monnet (1995) permite evaluar el balasto teniendo en cuenta la verdadera rigidez a flexión ($EI$) del elemento estructural:")
    st.latex(r"""
    K_h = \left( \frac{20 \cdot EI \cdot \gamma \cdot (K_p - K_0)}{d_{ro}^4} \right)^{1/5} + \frac{A_p \cdot c \cdot \tanh(c/c_0)}{d_{ro}}
    """)
    
    st.markdown("""
    **Donde:**
    * **$EI$:** Rigidez a flexión del muro por metro lineal (kN·m²/m).
    * **$\gamma$:** Peso específico del terreno (kN/m³).
    * **$K_p, K_0$:** Coeficientes de empuje pasivo y reposo de Rankine.
    * **$c$:** Cohesión efectiva del terreno (kPa).
    * **$d_{ro}$, $c_0$, $A_p$:** Constantes matemáticas de calibración del modelo.
    """)
    
    st.divider()

    st.subheader("2. Método Empírico Original (Ábaco de Chadeisson)")
    st.markdown("""
    El método original desarrollado por Chadeisson en los años 60 en Francia no obedece a una fórmula matemática, sino a una familia de curvas empíricas deducidas de la experiencia en excavaciones de muros pantalla continuos reales de la época.
    
    Para obtener el valor **$K_h \text{ Chadeisson}$**, este programa lee y digitaliza la matriz de coordenadas exacta del ábaco original y realiza una **interpolación espacial bicúbica** cruzando el ángulo de rozamiento interno ($\phi$) y la cohesión ($c$). 
    
    *Nota: El ábaco original lleva implícita una rigidez estructural altísima, equivalente a pantallas de 0.80m de espesor con hormigones de 20 GPa.*
    """)

    st.divider()

    st.subheader("3. Empujes de Tierras (Teoría de Rankine)")
    col_k1, col_k2, col_k3 = st.columns(3)
    with col_k1:
        st.markdown("**Empuje Activo ($K_a$)**")
        st.latex(r"K_a = \frac{1 - \sin\phi}{1 + \sin\phi}")
    with col_k2:
        st.markdown("**Empuje al Reposo ($K_0$)**")
        st.latex(r"K_0 = 1 - \sin\phi")
    with col_k3:
        st.markdown("**Empuje Pasivo ($K_p$)**")
        st.latex(r"K_p = \frac{1 + \sin\phi}{1 - \sin\phi}")