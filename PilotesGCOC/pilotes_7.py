import streamlit as st
import numpy as np
import pandas as pd
import math
import io
import plotly.graph_objects as go
import plotly.express as px
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date

st.set_page_config(page_title="Cálculo de Pilotes - GCOC", layout="wide", page_icon="🏗️")

st.title("🏗️ Diseño de Pilotes GCOC ")
st.markdown("Cálculo según el método analítico.")

# ══════════════════════════════════════════════════════════════════════════
# INICIALIZACIÓN DE LA TABLA BASE 
# ══════════════════════════════════════════════════════════════════════════
if 'df_base' not in st.session_state:
    st.session_state.df_base = pd.DataFrame({
        "Estrato": ["UG-01", "UG-02", "UG-03", "UG-04"],
        "Espesor (m)": [2.0, 5.0, 8.0, 10.0],
        "Gamma Seco (kN/m3)": [18.0, 17.0, 19.0, 20.0],
        "Gamma Sat. (kN/m3)": [20.0, 18.0, 21.0, 21.0],
        "Condición": ["Largo Plazo", "Corto Plazo", "Largo Plazo", "Corto Plazo"],
        "c / cu (kPa)": [0.0, 100.0, 0.0, 150.0],
        "phi (grados)": [28.0, 0.0, 35.0, 0.0]
    })

if 'calculado' not in st.session_state: st.session_state.calculado = False
if 'word_buffer' not in st.session_state: st.session_state.word_buffer = None
if 'fig_auditoria_guardada' not in st.session_state: st.session_state.fig_auditoria_guardada = None
if 'fig_final_guardada' not in st.session_state: st.session_state.fig_final_guardada = None

# ══════════════════════════════════════════════════════════════════════════
# INTERFAZ PRINCIPAL DE PESTAÑAS 
# ══════════════════════════════════════════════════════════════════════════
tab_datos, tab_tensiones, tab_matriz_punta, tab_matriz_fuste, tab_matriz_total, tab_matriz_tope, tab_auditoria, tab_formulacion = st.tabs([
    "📋 1. Estratigrafía", "🌊 2. Tensiones", "🔻 3. Punta", "🟫 4. Fuste", "🌍 5. Total", "🛑 6. Tope Estruct.", "🔍 7. Auditoría", "📖 8. Formulación"
])

with tab_datos:
    st.subheader("Definición de Unidades")
    df_edit = st.data_editor(
        st.session_state.df_base, 
        key="tabla_estratos",
        num_rows="dynamic", 
        use_container_width=True,
        column_config={
            "Condición": st.column_config.SelectboxColumn(
                options=["Corto Plazo", "Largo Plazo"],
                required=True
            )
        }
    )
    
    mascara_corto_plazo = df_edit["Condición"] == "Corto Plazo"
    if (df_edit.loc[mascara_corto_plazo, "phi (grados)"] != 0.0).any():
        df_edit.loc[mascara_corto_plazo, "phi (grados)"] = 0.0
        st.session_state.df_base = df_edit.copy() 
        st.toast("💡 φ ajustado automáticamente a 0º", icon="⚙️")
        st.rerun()

    # --- VALIDACIÓN ESTRICTA DE ESPESORES ---
    espesores_invalidos = (df_edit["Espesor (m)"] <= 0).any()
    if espesores_invalidos:
        estratos_nulos = df_edit[df_edit["Espesor (m)"] <= 0]["Estrato"].tolist()
        st.error(
            f"⛔ **Error de datos:** Los estratos **{', '.join(estratos_nulos)}** tienen "
            f"espesor ≤ 0 m. Corrige la tabla (debe ser > 0) antes de calcular."
        )
        st.session_state.calculado = False

    z_max_total = df_edit["Espesor (m)"].sum()
    st.info(f"Profundidad máxima actual del sondeo: **{z_max_total:.2f} m**.")
    st.info("ℹ️ **Nota:** Si la zona de influencia de la punta (3D) sobrepasa la profundidad máxima definida, el software asume que el último estrato se prolonga indefinidamente.")
    st.info("Después de cada modificación hay que pulsar el botón 'Calcular Pilotes' para actualizar los resultados.")
    if not st.session_state.calculado and not espesores_invalidos:
        st.warning("👈 Haz clic en **'Calcular Pilotes'** en el menú izquierdo para procesar los datos.")

# ══════════════════════════════════════════════════════════════════════════
# BARRA LATERAL: CONFIGURACIÓN 
# ══════════════════════════════════════════════════════════════════════════
st.sidebar.header("💧 Nivel Freático")
zw = st.sidebar.number_input("Prof. Nivel Freático, zw (m)", min_value=0.0, value=3.0, step=0.5)

st.sidebar.header("🪂 Factor de Seguridad")
opcion_fs = st.sidebar.selectbox(
    "Situación de Proyecto (FS)",
    ["Casi permanente (FS ≥ 3.00)", "Característica / Transitoria (FS ≥ 2.60)", "Accidental / Sísmica (FS ≥ 2.20)", "Valor Personalizado"]
)

if opcion_fs == "Casi permanente (FS ≥ 3.00)": FS, sit_str = 3.00, "Casi Permanente"
elif opcion_fs == "Característica / Transitoria (FS ≥ 2.60)": FS, sit_str = 2.60, "Transitoria"
elif opcion_fs == "Accidental / Sísmica (FS ≥ 2.20)": FS, sit_str = 2.20, "Accidental"
else:
    FS = st.sidebar.number_input("Introduce FS personalizado:", min_value=1.0, value=3.00, step=0.1)
    sit_str = "Personalizada"

st.sidebar.header("🛑 Tope Estructural (GCOC 5.15.1)")
opcion_tope = st.sidebar.selectbox(
    "Método de Ejecución del Pilote",
    ["In situ: Seco / Lodos / Barrena (4 MPa)", 
     "In situ: Entubado (5 MPa)", 
     "In situ: Apoyo en Roca Dura (6 MPa)",
     "Valor Personalizado (MPa)"]
)

if "4 MPa" in opcion_tope: sigma_tope_mpa = 4.0
elif "5 MPa" in opcion_tope: sigma_tope_mpa = 5.0
elif "6 MPa" in opcion_tope: sigma_tope_mpa = 6.0
else: sigma_tope_mpa = st.sidebar.number_input("Tensión Tope (MPa):", min_value=1.0, value=4.0, step=0.5)

st.sidebar.header("📐 Matriz Geométrica del Pilote")
col_d1, col_d2 = st.sidebar.columns(2)
D_min = col_d1.number_input("Ø min (m)", value=0.6, min_value=0.3, step=0.1)
D_max = col_d2.number_input("Ø max (m)", value=1.5, min_value=0.3, step=0.1)
D_step = st.sidebar.number_input("Paso Ø (m)", value=0.3, min_value=0.1)

col_l1, col_l2 = st.sidebar.columns(2)
L_min_default = min(10.0, float(z_max_total)) 
L_min = col_l1.number_input("L min (m)", value=L_min_default, min_value=1.0, max_value=float(z_max_total), step=1.0)
L_max_default = min(20.0, float(z_max_total))
L_max_default = max(L_max_default, float(L_min))
L_max = col_l2.number_input("L max (m)", value=L_max_default, min_value=float(L_min), max_value=float(z_max_total), step=1.0)
L_step = st.sidebar.number_input("Paso L (m)", value=2.5, min_value=0.5)

st.sidebar.markdown("---")
# BOTÓN BLOQUEADO SI HAY ERRORES EN LA TABLA
if st.sidebar.button("⚙️ Calcular Pilotes", type="primary", use_container_width=True, disabled=espesores_invalidos):
    st.session_state.calculado = True
    st.session_state.word_buffer = None

# ══════════════════════════════════════════════════════════════════════════
# MOTORES DE CÁLCULO
# ══════════════════════════════════════════════════════════════════════════
def calcular_perfil_tensiones(df, zw, z_max):
    z_max_extendido = z_max + 10.0 
    z_vals = np.arange(0, z_max_extendido + 0.1, 0.1)
    sigma_v, u, sigma_v_eff = [], [], []
    sv_acum = 0.0
    for z in z_vals:
        z_acum, gamma_actual = 0.0, 18.0
        for idx, row in df.iterrows():
            is_last = (idx == len(df) - 1)
            if z <= z_acum + row["Espesor (m)"] + 1e-5 or is_last:
                if z <= zw + 1e-5:
                    gamma_actual = row["Gamma Seco (kN/m3)"]
                else:
                    gamma_actual = row["Gamma Sat. (kN/m3)"]
                break
            z_acum += row["Espesor (m)"]
            
        sv_acum = sv_acum + gamma_actual * 0.1 if z > 0 else 0.0
        u_val = max(0, (z - zw) * 9.81)
        sigma_v.append(sv_acum)
        u.append(u_val)
        sigma_v_eff.append(sv_acum - u_val)
        
    return z_vals, np.array(sigma_v), np.array(u), np.array(sigma_v_eff)

def obtener_tension_a_profundidad(z_target, z_vals, sigma_array):
    return np.interp(z_target, z_vals, sigma_array)

def calcular_pilote(D, L, df, zw, fS_val, sigma_tope_mpa):
    z_max = df["Espesor (m)"].sum()
    if L > z_max: return None 
    
    fD_calc = max(2.0/3.0, 1.0 - (D / 3.0))
    z_vals, sig_v, u, sig_v_eff = calcular_perfil_tensiones(df, zw, z_max)
    
    # --- PUNTA GCOC (ESTRICTA LEY DE ESPESORES 6D/3D) ---
    z_eval_punta = min(L, 20 * D)
    sig_v_eff_punta = obtener_tension_a_profundidad(z_eval_punta, z_vals, sig_v_eff)
    
    z_sup_bulbo = max(0.0, L - (6 * D))
    z_inf_bulbo = L + (3 * D)
    espesor_bulbo = z_inf_bulbo - z_sup_bulbo
    
    qp_eq_acumulado = 0.0
    z_acum = 0.0
    detalle_bulbo_grafico = [] 
    
    for idx, row in df.iterrows():
        is_last = (idx == len(df) - 1)
        h_estrato = row["Espesor (m)"]
        z_top_estrato = z_acum
        
        z_bot_estrato = max(z_acum + h_estrato, z_inf_bulbo) if is_last else z_acum + h_estrato
            
        overlap_top = max(z_sup_bulbo, z_top_estrato)
        overlap_bot = min(z_inf_bulbo, z_bot_estrato)
        overlap_h = max(0.0, overlap_bot - overlap_top)
        
        if overlap_h > 0 and espesor_bulbo > 0:
            cond_i = row["Condición"]
            c_i = row["c / cu (kPa)"]
            phi_i = row["phi (grados)"]
            
            if "Corto Plazo" in cond_i:
                qp_i = 9.0 * c_i * fD_calc
            else:
                phi_rad = math.radians(phi_i)
                if phi_i == 0:
                    Nq_star, Nc_star = 1.0, 9.0 * fD_calc
                else:
                    Nq_star = 1.5 * ((1 + math.sin(phi_rad))/(1 - math.sin(phi_rad))) * math.exp(math.pi * math.tan(phi_rad)) * fD_calc
                    Nc_star = (Nq_star - 1) / math.tan(phi_rad)
                qp_i = (sig_v_eff_punta * Nq_star) + (c_i * Nc_star)
            
            peso = overlap_h / espesor_bulbo
            qp_eq_acumulado += qp_i * peso
            
            nombre_estrato = f"{row['Estrato']} (Prolongado)" if is_last and overlap_bot > (z_acum + h_estrato) else row['Estrato']
            
            detalle_bulbo_grafico.append({
                "Estrato": nombre_estrato,
                "Espesor en bulbo (m)": overlap_h,
                "Participación (%)": peso * 100.0,
                "q_p individual (kPa)": qp_i
            })
            
        z_acum += h_estrato
        
    qp = qp_eq_acumulado
    Area_pilote = (math.pi * D**2) / 4.0
    Q_punta = qp * Area_pilote
    
    auditoria_punta = {
        "Profundidad Punta (m)": L,
        "σ'_v efectiva base (kPa)": sig_v_eff_punta,
        "Factor de escala (fD)": fD_calc,
        "Resist. Unitaria q_p (kPa)": qp,
        "Fuerza Total Punta (kN)": Q_punta
    }
    
    # --- FUSTE (MEJORADO PARA CORTES DEL NIVEL FREÁTICO) ---
    Q_fuste, z_top, Perimetro = 0.0, 0.0, math.pi * D
    auditoria_fuste = [] 
    k0_tan_delta_fijo = 0.30 
    
    for _, row in df.iterrows():
        z_bot = z_top + row["Espesor (m)"]
        if z_top >= L: break
        
        z_end_tramo = min(z_bot, L)
        if z_end_tramo <= z_top:
            z_top = z_bot
            continue
            
        # Identificar si el Nivel Freático cruza por medio de este estrato
        puntos_corte = [z_top]
        if z_top < zw < z_end_tramo:
            puntos_corte.append(zw)
        puntos_corte.append(z_end_tramo)
        
        # Calcular los sub-tramos (secos y saturados por separado)
        for i in range(len(puntos_corte) - 1):
            z_sub_top = puntos_corte[i]
            z_sub_bot = puntos_corte[i+1]
            L_sub = z_sub_bot - z_sub_top
            
            z_mid = z_sub_top + (L_sub / 2.0)
            sig_v_eff_mid = obtener_tension_a_profundidad(z_mid, z_vals, sig_v_eff)
            
            if "Corto Plazo" in row["Condición"]:
                tau_f = min(row["c / cu (kPa)"] * (100.0 / (100.0 + row["c / cu (kPa)"])), 70.0)
            else:
                tau_f = min(row["c / cu (kPa)"] + k0_tan_delta_fijo * sig_v_eff_mid, 90.0)
                
            Q_tramo = tau_f * Perimetro * L_sub
            Q_fuste += Q_tramo
            
            # Etiquetamos el estrato si ha sufrido una división freática
            sufijo = ""
            if len(puntos_corte) > 2:
                sufijo = " (Seco)" if z_sub_bot <= zw else " (Sat.)"
            
            auditoria_fuste.append({
                "Estrato": row["Estrato"] + sufijo,
                "Cotas (m)": f"{z_sub_top:.1f} a {z_sub_bot:.1f}",
                "Long. fuste (m)": f"{L_sub:.2f}",
                "σ'_v media (kPa)": sig_v_eff_mid,
                "Resist. Unitaria τ_f (kPa)": tau_f,
                "Fuerza Tramo (kN)": Q_tramo
            })
            
        z_top = z_bot
        
    Q_total_geo = Q_punta + Q_fuste
    Q_adm_geo = Q_total_geo / fS_val
    
    Q_tope_est = Area_pilote * (sigma_tope_mpa * 1000.0)
    Q_final_diseno = min(Q_adm_geo, Q_tope_est)
    control = "ESTRUCTURAL" if Q_tope_est < Q_adm_geo else "GEOTÉCNICO"

    return {
        "D": D, "L": L, "fD": fD_calc, 
        "Q_punta (kN)": Q_punta, "Q_fuste (kN)": Q_fuste, 
        "Q_adm_geo (kN)": Q_adm_geo,
        "Q_tope_est (kN)": Q_tope_est,
        "Q_final (kN)": Q_final_diseno,
        "Control": control,
        "auditoria_punta": auditoria_punta,
        "auditoria_fuste": auditoria_fuste,
        "auditoria_bulbo": detalle_bulbo_grafico,
        "z_sup_bulbo": z_sup_bulbo,
        "z_inf_bulbo": z_inf_bulbo
    }

# ══════════════════════════════════════════════════════════════════════════
# GENERADOR DEL INFORME EN WORD (VERSIÓN EJECUTIVA SIMPLIFICADA)
# ══════════════════════════════════════════════════════════════════════════
def generar_word_pilotes(df_estratos, fig_tens, df_pivot_geo, df_pivot_final, fS_val, situacion, zw_val, sigma_tope, datos_unitarios_df, fig_final):
    doc = Document()
    estilo_tabla = 'Light Grid Accent 1' # Estilo elegante de Word
    
    # --- 1. PORTADA PROFESIONAL ---
    doc.add_paragraph('\n\n\n\n')
    title = doc.add_paragraph('ANEJO DE CÁLCULO: CIMENTACIONES PROFUNDAS')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(24)
        run.font.bold = True
        
    subtitle = doc.add_paragraph('Diseño Analítico de Pilotes según GCOC')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs: run.font.size = Pt(16)
        
    doc.add_paragraph('\n\n\n')
    doc.add_paragraph(f'Fecha de Generación: {date.today().strftime("%d/%m/%Y")}').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Proyecto: ______________________________________________________').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Calculista: _____________________________________________________').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
    
    # --- 2. BASES DE CÁLCULO Y METODOLOGÍA (CON FÓRMULAS) ---
    doc.add_heading('1. Bases de Cálculo y Parámetros Iniciales', level=1)
    p_bases = doc.add_paragraph()
    p_bases.add_run(f'• Situación de Proyecto: ').bold = True
    p_bases.add_run(f'{situacion} (Factor de Seguridad Global = {fS_val:.2f})\n')
    p_bases.add_run(f'• Nivel Freático: ').bold = True
    p_bases.add_run(f'Considerado a una profundidad Z = {zw_val:.2f} m\n')
    p_bases.add_run(f'• Tope Estructural del Hormigón: ').bold = True
    p_bases.add_run(f'{sigma_tope:.1f} MPa (GCOC Tabla 5.5)\n')
    p_bases.add_run(f'• Condición de Contorno en Punta: ').bold = True
    p_bases.add_run(f'Si el bulbo de influencia 3D supera la profundidad explorada, se asume prolongación indefinida del estrato base.')

    doc.add_heading('2. Metodología y Formulación Analítica (GCOC)', level=1)
    p_metodo = doc.add_paragraph()
    p_metodo.add_run('Resistencia por Punta (qp):\n').bold = True
    p_metodo.add_run('Se aplica un factor de reducción por efecto de escala para pilotes de gran diámetro: fD = 1 - (D/3) ≥ 2/3.\n')
    p_metodo.add_run('• Corto Plazo (Arcillas): qp = 9 · cu · fD\n')
    p_metodo.add_run('• Largo Plazo (Arenas): qp = σ\'v,p · Nq* + c\' · Nc*\n')
    p_metodo.add_run('El valor final de la punta se obtiene promediando estrictamente por espesores interceptados en un bulbo de rotura de 6D por encima y 3D por debajo de la base.\n\n')
    
    p_metodo.add_run('Resistencia por Fuste (τf):\n').bold = True
    p_metodo.add_run('• Corto Plazo: τf = cu · [100 / (100 + cu)] ≤ 70 kPa\n')
    p_metodo.add_run('• Largo Plazo: τf = c\' + σ\'v · K0 · tan(δ) ≤ 90 kPa. (Se adopta normativamente K0·tan(δ) = 0.30).\n')
    p_metodo.add_run('Los estratos atravesados por el nivel freático se discretizan en sub-tramos independientes para integrar correctamente el quiebro de la tensión vertical efectiva.\n\n')
    
    p_metodo.add_run('Diseño Final:\n').bold = True
    p_metodo.add_run('La carga de diseño es el mínimo entre (Q_geotécnica / FS) y el Tope Estructural geométrico del hormigón.')

    # --- 3. ESTRATIGRAFÍA Y RESISTENCIAS UNITARIAS ---
    doc.add_heading('3. Perfil Estratigráfico', level=1)
    tabla_estratos = doc.add_table(rows=1, cols=len(df_estratos.columns))
    tabla_estratos.style = estilo_tabla
    hdr_cells = tabla_estratos.rows[0].cells
    for i, column in enumerate(df_estratos.columns): hdr_cells[i].text = str(column)
    for index, row in df_estratos.iterrows():
        row_cells = tabla_estratos.add_row().cells
        for i, value in enumerate(row): row_cells[i].text = str(value)
            
    doc.add_heading('4. Resistencias Unitarias Base del Terreno', level=1)
    doc.add_paragraph('Fricción media por estrato y presión en punta evaluada en la base de cada capa (sin afectar por fD).')
    if not datos_unitarios_df.empty:
        tabla_unit = doc.add_table(rows=1, cols=len(datos_unitarios_df.columns))
        tabla_unit.style = estilo_tabla
        hdr_unit = tabla_unit.rows[0].cells
        for i, col in enumerate(datos_unitarios_df.columns): hdr_unit[i].text = str(col)
        for index, row in datos_unitarios_df.iterrows():
            row_cells = tabla_unit.add_row().cells
            for i, value in enumerate(row):
                 if isinstance(value, float): row_cells[i].text = f"{value:.1f}"
                 else: row_cells[i].text = str(value)
                 
    doc.add_page_break()

    # --- 4. MATRICES DE DISEÑO Y CURVAS ---
    doc.add_heading(f'5. Matriz Admisible TOTAL del Terreno (kN)', level=1)
    doc.add_paragraph('Capacidad portante exclusiva del terreno (Punta + Fuste) dividida por el Factor de Seguridad.')
    tabla_geo = doc.add_table(rows=1, cols=len(df_pivot_geo.columns) + 1)
    tabla_geo.style = estilo_tabla
    hdr_geo = tabla_geo.rows[0].cells
    hdr_geo[0].text = "L / Ø"
    for i, col_name in enumerate(df_pivot_geo.columns): hdr_geo[i+1].text = str(col_name).replace('\n', ' ')
    for index, row in df_pivot_geo.iterrows():
        row_cells = tabla_geo.add_row().cells
        row_cells[0].text = str(index)
        for i, val in enumerate(row): row_cells[i+1].text = f"{val:.0f}"

    doc.add_paragraph('\n')
    
    doc.add_heading(f'6. Matriz de Diseño FINAL LIMITADA (kN)', level=1)
    doc.add_paragraph('Mínimo entre la Matriz Geotécnica y el Tope Estructural. Los valores marcados con [EST] indican colapso estructural antes que geotécnico.')
    tabla_fin = doc.add_table(rows=1, cols=len(df_pivot_final.columns) + 1)
    tabla_fin.style = estilo_tabla
    hdr_fin = tabla_fin.rows[0].cells
    hdr_fin[0].text = "L / Ø"
    for i, col_name in enumerate(df_pivot_final.columns): hdr_fin[i+1].text = str(col_name).replace('\n', ' ')
    for index, row in df_pivot_final.iterrows():
        row_cells = tabla_fin.add_row().cells
        row_cells[0].text = str(index)
        for i, val in enumerate(row): row_cells[i+1].text = str(val)

    if fig_final is not None:
        doc.add_paragraph('\n')
        doc.add_heading('Gráfico: Curvas de Diseño Final', level=2)
        try:
            img_bytes = fig_final.to_image(format="png", width=800, height=500)
            doc.add_picture(io.BytesIO(img_bytes), width=Inches(6.5))
        except Exception: pass

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ══════════════════════════════════════════════════════════════════════════
# CÁLCULOS PRINCIPALES Y RENDERIZADO DE RESULTADOS
# ══════════════════════════════════════════════════════════════════════════
if st.session_state.calculado:
    z_vals, sig_v, u, sig_v_eff = calcular_perfil_tensiones(df_edit, zw, z_max_total)
    
    # --- Guardamos esta tabla unitaria mejorada (con cortes freáticos) para el Word ---
    datos_unitarios = []
    
    with tab_datos:
        st.markdown("---")
        st.subheader("📊 Resistencias Unitarias Base del Terreno (GCOC)")
        st.markdown("*Nota: Fricción $\\tau_f$ media por estrato. La presión en punta $q_p$ base se muestra **sin afectar por el factor de escala** ($f_D = 1.0$).*")
        
        z_top_loop = 0.0
        for idx, row_est in df_edit.iterrows():
            espesor = row_est["Espesor (m)"]
            if espesor <= 0: continue
            z_bot_loop = z_top_loop + espesor
            
            # Identificar si el Nivel Freático cruza por medio de este estrato
            puntos_corte = [z_top_loop]
            if z_top_loop < zw < z_bot_loop:
                puntos_corte.append(zw)
            puntos_corte.append(z_bot_loop)
            
            for i in range(len(puntos_corte) - 1):
                z_sub_top = puntos_corte[i]
                z_sub_bot = puntos_corte[i+1]
                z_mid_loop = z_sub_top + (z_sub_bot - z_sub_top) / 2.0
                
                sig_v_eff_mid_loop = obtener_tension_a_profundidad(z_mid_loop, z_vals, sig_v_eff)
                sig_v_eff_base = obtener_tension_a_profundidad(min(z_sub_bot, 20*D_min), z_vals, sig_v_eff)
                cond_loop = row_est["Condición"]
                c_loop = row_est["c / cu (kPa)"]
                phi_loop = row_est["phi (grados)"]
                
                if "Corto Plazo" in cond_loop:
                    tau_f_media = min(c_loop * (100.0 / (100.0 + c_loop)), 70.0)
                    qp_base = 9.0 * c_loop
                else:
                    tau_f_media = min(c_loop + 0.30 * sig_v_eff_mid_loop, 90.0)
                    phi_rad = math.radians(phi_loop)
                    if phi_loop == 0: Nq_star, Nc_star = 1.0, 9.0
                    else:
                        Nq_star = 1.5 * ((1 + math.sin(phi_rad))/(1 - math.sin(phi_rad))) * math.exp(math.pi * math.tan(phi_rad))
                        Nc_star = (Nq_star - 1) / math.tan(phi_rad)
                    qp_base = (sig_v_eff_base * Nq_star) + (c_loop * Nc_star)
                
                sufijo = ""
                if len(puntos_corte) > 2:
                    sufijo = " (Seco)" if z_sub_bot <= zw else " (Sat.)"
                    
                datos_unitarios.append({
                    "Estrato": row_est["Estrato"] + sufijo,
                    "Profundidad (m)": f"de {z_sub_top:.1f} a {z_sub_bot:.1f}",
                    "Fricción media, τ_f (kPa)": round(tau_f_media, 1),
                    "Punta base, q_p (kPa)": round(qp_base, 1)
                })
            
            z_top_loop = z_bot_loop
            
        df_unitarios = pd.DataFrame(datos_unitarios)
        st.dataframe(df_unitarios.style.format({"Fricción media, τ_f (kPa)": "{:.1f}", "Punta base, q_p (kPa)": "{:.1f}"}).hide(axis="index"), use_container_width=True)

    fig_tens = go.Figure()
    fig_tens.add_trace(go.Scatter(x=sig_v, y=z_vals, name='Tensión Total (σ)', line=dict(color='black', width=2)))
    fig_tens.add_trace(go.Scatter(x=u, y=z_vals, name='Pres. Intersticial (u)', line=dict(color='blue', dash='dash')))
    fig_tens.add_trace(go.Scatter(x=sig_v_eff, y=z_vals, name='Tensión Efectiva (σ\')', line=dict(color='red', width=3)))
    
    z_acum = 0
    colores = ['#f0f8ff', '#e6e6fa', '#fff0f5', '#f5fffa', '#ffebcd']
    for i, row in df_edit.iterrows():
        z_next = z_acum + row["Espesor (m)"]
        fig_tens.add_hrect(y0=z_acum, y1=z_next, fillcolor=colores[i%len(colores)], opacity=0.4, line_width=0, annotation_text=row["Estrato"])
        z_acum = z_next
    fig_tens.add_hline(y=zw, line_dash="dot", line_color="blue", annotation_text="Nivel Freático")
    fig_tens.update_yaxes(autorange="reversed", title="Profundidad (m)")
    fig_tens.update_xaxes(title="Presión (kPa)", side="top")
    
    with tab_tensiones:
        st.plotly_chart(fig_tens, use_container_width=True)
        
        # --- NUEVO: INSPECTOR PUNTUAL DE TENSIONES ---
        st.markdown("---")
        st.subheader("🔎 Consulta puntual de tensiones")
        z_consulta = st.number_input("Introduce una profundidad Z (m):", min_value=0.0, max_value=float(z_max_total + 10.0), value=float(zw), step=0.1)
        
        sig_v_z = obtener_tension_a_profundidad(z_consulta, z_vals, sig_v)
        u_z = obtener_tension_a_profundidad(z_consulta, z_vals, u)
        sig_v_eff_z = obtener_tension_a_profundidad(z_consulta, z_vals, sig_v_eff)
        
        col_t1, col_t2, col_t3 = st.columns(3)
        col_t1.metric(label="Tensión Total (σ)", value=f"{sig_v_z:.2f} kPa")
        col_t2.metric(label="Pres. Intersticial (u)", value=f"{u_z:.2f} kPa")
        col_t3.metric(label="Tensión Efectiva (σ')", value=f"{sig_v_eff_z:.2f} kPa")

    resultados = []
    D_arr = np.arange(D_min, D_max + 1e-5, D_step)
    L_arr = np.arange(L_min, L_max + 1e-5, L_step)

    for D in D_arr:
        for L in L_arr:
            res = calcular_pilote(D, L, df_edit, zw, FS, sigma_tope_mpa)
            if res is not None: resultados.append(res)

    df_res = pd.DataFrame(resultados)
    df_pivot_geo_global = None
    df_pivot_final_global = None
    res_auditoria_seleccionada = None

    if not df_res.empty:
        columnas_con_fd = [f"Ø {d_val:.2f} m\n(fD={df_res[df_res['D'] == d_val]['fD'].iloc[0]:.2f})" for d_val in df_res['D'].unique()]
        
        df_pivot_punta = df_res.pivot(index="L", columns="D", values="Q_punta (kN)") / FS
        df_pivot_punta.index = [f"L = {idx:.1f} m" for idx in df_pivot_punta.index]
        df_pivot_punta.columns = columnas_con_fd
        with tab_matriz_punta:
            st.subheader(f"🔻 Carga Admisible SOLO por Punta (kN) - {sit_str} (FS = {FS:.2f})")
            st.dataframe(df_pivot_punta.style.background_gradient(cmap='Reds', axis=None).format("{:.0f}"), use_container_width=True)

        df_pivot_fuste = df_res.pivot(index="L", columns="D", values="Q_fuste (kN)") / FS
        df_pivot_fuste.index = [f"L = {idx:.1f} m" for idx in df_pivot_fuste.index]
        df_pivot_fuste.columns = columnas_con_fd
        with tab_matriz_fuste:
            st.subheader(f"🟫 Carga Admisible SOLO por Fuste (kN) - {sit_str} (FS = {FS:.2f})")
            st.dataframe(df_pivot_fuste.style.background_gradient(cmap='Oranges', axis=None).format("{:.0f}"), use_container_width=True)

        df_pivot_geo = df_res.pivot(index="L", columns="D", values="Q_adm_geo (kN)")
        df_pivot_geo.index = [f"L = {idx:.1f} m" for idx in df_pivot_geo.index]
        df_pivot_geo.columns = columnas_con_fd
        df_pivot_geo_global = df_pivot_geo
        with tab_matriz_total:
            st.subheader(f"🌍 Carga Admisible TOTAL del Terreno (kN) - {sit_str}")
            st.markdown("*Esta es la capacidad exclusiva del terreno (Punta + Fuste) dividida por el FS, sin importar de qué material esté hecho el pilote.*")
            st.dataframe(df_pivot_geo.style.background_gradient(cmap='Greens', axis=None).format("{:.0f}"), use_container_width=True)

        df_pivot_final = df_res.pivot(index="L", columns="D", values="Q_final (kN)")
        df_pivot_control = df_res.pivot(index="L", columns="D", values="Control")
        
        df_final_formateada = pd.DataFrame(index=df_pivot_final.index, columns=df_pivot_final.columns)
        for c in df_pivot_final.columns:
            for r in df_pivot_final.index:
                val = df_pivot_final.loc[r, c]
                ctrl = df_pivot_control.loc[r, c]
                if ctrl == "ESTRUCTURAL":
                    df_final_formateada.loc[r, c] = f"{val:.0f} [EST]"
                else:
                    df_final_formateada.loc[r, c] = f"{val:.0f}"

        df_final_formateada.index = [f"L = {idx:.1f} m" for idx in df_final_formateada.index]
        df_final_formateada.columns = columnas_con_fd
        df_pivot_final_global = df_final_formateada
        
        with tab_matriz_tope:
            st.subheader(f"🛑 Matriz de Carga de Diseño FINAL (kN) - Limitada a {sigma_tope_mpa} MPa")
            st.markdown("*Muestra el valor mínimo entre la Resistencia Geotécnica y el Tope Estructural del Hormigón.*")
            
            def color_tope(val):
                color = '#ffcccc' if '[EST]' in str(val) else '#ccffcc'
                return f'background-color: {color}'
            
            st.dataframe(df_final_formateada.style.applymap(color_tope), use_container_width=True)
            
            st.markdown("---")
            st.subheader("📈 Curvas de Diseño Final (Geotécnico + Estructural)")
            df_plot_final = df_res.copy()
            df_plot_final["Diámetro"] = df_plot_final["D"].apply(lambda x: f"Ø {x:.2f} m")
            
            # --- MEJORA: GRÁFICA OBLIGATORIAMENTE A COLOR Y FONDO BLANCO ---
            fig_final = px.line(df_plot_final, x="L", y="Q_final (kN)", color="Diámetro", markers=True,
                                title="Capacidad de Diseño Final vs. Longitud (Limitada por Tope Estructural)",
                                color_discrete_sequence=px.colors.qualitative.Set1, # Fuerza una paleta de colores viva
                                template="plotly_white") # Fuerza el fondo blanco para el Word
            
            fig_final.update_layout(xaxis_title="Longitud L (m)", yaxis_title="Carga de Diseño Final (kN)", hovermode="x unified")
            fig_final.update_traces(line=dict(width=3), marker=dict(size=8)) # Líneas más gruesas para que luzcan en el informe
            
            st.plotly_chart(fig_final, use_container_width=True)
            st.session_state.fig_final_guardada = fig_final # Guardamos para el Word

        with tab_auditoria:
            st.subheader("🔍 Inspector de Resistencias Unitarias y Bulbo de Punta")
            col_aud1, col_aud2 = st.columns(2)
            
            d_aud = col_aud1.selectbox("Selecciona Diámetro Ø (m):", df_res['D'].unique(), format_func=lambda x: f"{x:.2f}", key="d_aud")
            l_aud = col_aud2.selectbox("Selecciona Longitud L (m):", df_res['L'].unique(), format_func=lambda x: f"{x:.2f}", key="l_aud")
            
            res_auditoria = df_res[(df_res['D'] == d_aud) & (df_res['L'] == l_aud)]
            
            if not res_auditoria.empty:
                fila_aud = res_auditoria.iloc[0]
                
                st.markdown(f"### ➡️ Carga de Diseño Final: **{fila_aud['Q_final (kN)']:.0f} kN** (Controlado por: **{fila_aud['Control']}**)")
                st.markdown(f"- Admisible del Terreno: {fila_aud['Q_adm_geo (kN)']:.0f} kN")
                st.markdown(f"- Tope Estructural del Hormigón: {fila_aud['Q_tope_est (kN)']:.0f} kN")
                st.markdown("---")
                
                st.markdown(f"#### 🟫 Desglose por Fuste (sin minorar)")
                st.dataframe(pd.DataFrame(fila_aud['auditoria_fuste']).style.format({"Long. fuste (m)": "{:.2f}", "σ'_v media (kPa)": "{:.1f}", "Resist. Unitaria τ_f (kPa)": "{:.2f}", "Fuerza Tramo (kN)": "{:.0f}"}).hide(axis="index"), use_container_width=True)
                
                st.markdown("---")
                st.markdown(f"#### 🔻 Resumen de la Punta (sin minorar)")
                st.dataframe(pd.DataFrame([fila_aud['auditoria_punta']]).style.format({
                    "Profundidad Punta (m)": "{:.2f}", 
                    "σ'_v efectiva base (kPa)": "{:.1f}", 
                    "Factor de escala (fD)": "{:.3f}", 
                    "Resist. Unitaria q_p (kPa)": "{:.2f}", 
                    "Fuerza Total Punta (kN)": "{:.0f}"
                }).hide(axis="index"), use_container_width=True)
                
                st.markdown(f"#### 🧅 Composición del Bulbo (Promedio según GCOC)")
                df_bulbo = pd.DataFrame(fila_aud['auditoria_bulbo'])
                st.dataframe(df_bulbo.style.format({
                    "Espesor en bulbo (m)": "{:.2f}", 
                    "Participación (%)": "{:.1f}%", 
                    "q_p individual (kPa)": "{:.0f}"
                }).hide(axis="index"), use_container_width=True)

                st.markdown("---")
                st.markdown(f"#### 📐 Esquema Gráfico del Pilote y Estratigrafía")
                
                D_val = fila_aud['D']
                L_val = fila_aud['L']
                z_sup = fila_aud['z_sup_bulbo']
                z_inf = fila_aud['z_inf_bulbo']
                
                fig_bulbo = go.Figure()
                
                z_acum_plt = 0
                colores_plt = ['#f0f8ff', '#e6e6fa', '#fff0f5', '#f5fffa', '#ffebcd']
                for i_plt, row_plt in df_edit.iterrows():
                    is_last = (i_plt == len(df_edit) - 1)
                    z_next_plt = z_acum_plt + row_plt["Espesor (m)"]
                    
                    if is_last:
                        z_next_plt = max(z_next_plt, z_inf + D_val)
                        
                    fig_bulbo.add_hrect(y0=z_acum_plt, y1=z_next_plt, fillcolor=colores_plt[i_plt%len(colores_plt)], opacity=0.7, line_width=1, annotation_text=row_plt["Estrato"], annotation_position="top left")
                    z_acum_plt = z_next_plt

                fig_bulbo.add_shape(type="rect", x0=-D_val*1.2, x1=D_val*1.2, y0=-D_val*0.6, y1=0, fillcolor="#8B8C89", line=dict(color="black", width=2))
                fig_bulbo.add_shape(type="rect", x0=-D_val/2, x1=D_val/2, y0=0, y1=L_val, fillcolor="#A9ACA9", line=dict(color="black", width=2))
                fig_bulbo.add_shape(type="rect", x0=-D_val/2, x1=-D_val/4, y0=0, y1=L_val, fillcolor="black", opacity=0.15, line_width=0)
                fig_bulbo.add_shape(type="rect", x0=D_val/4, x1=D_val/2, y0=0, y1=L_val, fillcolor="white", opacity=0.25, line_width=0)
                fig_bulbo.add_shape(type="line", x0=-D_val/2, x1=D_val/2, y0=L_val, y1=L_val, line=dict(color="black", width=4))

                fig_bulbo.add_shape(type="rect", x0=-D_val*1.5, x1=D_val*1.5, y0=z_sup, y1=z_inf, fillcolor="rgba(255, 0, 0, 0.15)", line=dict(color="red", width=3, dash="dash"))
                
                fig_bulbo.add_annotation(x=D_val*1.6, y=(z_sup+L_val)/2, text=f"6D Arriba<br>({L_val - z_sup:.1f} m)", showarrow=False, xanchor="left", font=dict(color="red", size=13, family="Arial Black"))
                fig_bulbo.add_annotation(x=D_val*1.6, y=(L_val+z_inf)/2, text=f"3D Abajo<br>({z_inf - L_val:.1f} m)", showarrow=False, xanchor="left", font=dict(color="red", size=13, family="Arial Black"))
                fig_bulbo.add_hline(y=L_val, line_color="black", line_width=2, opacity=0.8)
                
                fig_bulbo.add_hline(y=zw, line_dash="dot", line_color="blue", annotation_text="N.F.", annotation_position="bottom right")

                rango_inf_dibujo = max(z_max_total, z_inf + 2*D_val) 
                
                fig_bulbo.update_yaxes(autorange="reversed", title="Profundidad (m)", range=[rango_inf_dibujo, -D_val], showgrid=True)
                fig_bulbo.update_xaxes(showticklabels=False, range=[-D_val*4, D_val*4], showgrid=False)
                fig_bulbo.update_layout(height=600, margin=dict(l=20, r=20, t=20, b=20), paper_bgcolor="white", plot_bgcolor="white")
                
                st.plotly_chart(fig_bulbo, use_container_width=True)

        with tab_formulacion:
            st.subheader("📖 Ecuaciones y Formulación Geotécnica")
            st.markdown("Las fórmulas empleadas en este software se rigen por la **Guía de Cimentaciones en Obras de Carretera (GCOC)**.")
            
            st.markdown("### 1. Resistencia por Punta ($q_p$)")
            st.markdown("**1.1. Factor de Efecto de Escala ($f_D$)**")
            st.markdown("Aplica una reducción a la resistencia en punta para pilotes de gran tamaño:")
            st.latex(r"f_D = 1 - \frac{D}{3} \ge \frac{2}{3}")
            
            st.markdown("**1.2. Punta a Corto Plazo (Condición No Drenada)**")
            st.latex(r"q_p = 9 \cdot c_u \cdot f_D")
            
            st.markdown("**1.3. Punta a Largo Plazo (Condición Drenada)**")
            st.latex(r"q_p = \sigma'_{v,p} \cdot N_q^* + c' \cdot N_c^*")
            st.markdown("Donde los factores de capacidad de carga, afectados por la escala, son:")
            st.latex(r"N_q^* = 1.5 \cdot \frac{1 + \sin(\phi)}{1 - \sin(\phi)} \cdot e^{\pi \cdot \tan(\phi)} \cdot f_D")
            st.latex(r"N_c^* = \frac{N_q^* - 1}{\tan(\phi)}")
            st.markdown("*Nota: La tensión vertical efectiva en punta $\sigma'_{v,p}$ está limitada normativamente a una profundidad máxima de $20D$. Además, el cálculo emplea una resistencia $q_p$ promediada en función de los espesores de cada estrato en una zona de influencia de $6D$ por encima y $3D$ por debajo de la punta.*")
            
            st.markdown("---")
            st.markdown(r"### 2. Resistencia por Fuste ($\tau_f$)")
            st.markdown("**2.1. Fuste a Corto Plazo (Ley Hiperbólica)**")
            st.latex(r"\tau_f = c_u \cdot \frac{100}{100 + c_u} \le 70 \text{ kPa}")
            
            st.markdown("**2.2. Fuste a Largo Plazo (Ley Mohr-Coulomb)**")
            st.latex(r"\tau_f = c' + \sigma'_v \cdot K_0 \cdot \tan(\delta) \le 90 \text{ kPa}")
            st.markdown("*Nota: Para pilotes perforados y hormigonados in situ, el programa adopta de forma automática el valor recomendado por la GCOC de $K_0 \tan(\delta) = 0.30$.*")
            
            st.markdown("---")
            st.markdown("### 3. Tope Estructural")
            st.markdown("La capacidad final del pilote se limita por la resistencia del material constitutivo (hormigón).")
            st.latex(r"Q_{tope} = A \cdot \sigma_{tope} \cdot 1000")
            st.markdown(r"Donde $A$ es el área transversal del pilote ($\text{m}^2$) y $\sigma_{tope}$ es la resistencia admisible del hormigón impuesta por el método de ejecución (en MPa).")

    # ══════════════════════════════════════════════════════════════════════════
    # GENERACIÓN Y DESCARGA DE INFORME 
    # ══════════════════════════════════════════════════════════════════════════
    st.sidebar.markdown("---")
    st.sidebar.subheader("📄 Memoria de Cálculo")
    
    if st.sidebar.button("🛠️ Generar Informe (.docx)", type="primary", use_container_width=True):
        if not df_res.empty and df_pivot_geo_global is not None and df_pivot_final_global is not None:
            with st.spinner("Generando documento Word..."):
                # Hemos eliminado auditoria_data y fig_auditoria de los parámetros pasados
                buffer = generar_word_pilotes(
                    df_edit, fig_tens, df_pivot_geo_global, df_pivot_final_global, 
                    FS, sit_str, zw, sigma_tope_mpa, df_unitarios,
                    st.session_state.fig_final_guardada
                )
                st.session_state.word_buffer = buffer
            st.sidebar.success("✅ ¡Informe ejecutivo generado!")
        else:
            st.sidebar.error("No hay resultados calculados.")

    if st.session_state.word_buffer is not None:
        st.sidebar.download_button(
            label="⬇️ Descargar Informe", 
            data=st.session_state.word_buffer, 
            file_name="Anejo_Pilotes_GCOC.docx", 
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )