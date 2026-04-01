# calculo de pilotes GCOC


import streamlit as st
import numpy as np
import pandas as pd
import math
import io
import plotly.graph_objects as go
import plotly.express as px
from docx import Document
from docx.shared import Inches

st.set_page_config(page_title="Cálculo de Pilotes - GCOC", layout="wide", page_icon="🏗️")

st.title("🏗️ Diseño de Pilotes: Método Analítico (GCOC 5.10.2.5 y 5.15.1)")
st.markdown("Cálculo estricto según la normativa española (topes 20D, fD por diámetro, promedio de Resistencias q_p en bulbo 6D/3D y Tope Estructural).")

# ══════════════════════════════════════════════════════════════════════════
# INICIALIZACIÓN DE LA TABLA BASE (SIN K0·tan(delta) PARA MAYOR LIMPIEZA)
# ══════════════════════════════════════════════════════════════════════════
if 'df_base' not in st.session_state:
    st.session_state.df_base = pd.DataFrame({
        "Estrato": ["Rellenos", "Arcilla Blanda", "Arena Densa", "Arcilla Firme"],
        "Espesor (m)": [2.0, 5.0, 8.0, 10.0],
        "Gamma (kN/m3)": [18.0, 17.0, 20.0, 21.0],
        "Condición": ["Largo Plazo (Granular)", "Corto Plazo (Cohesivo)", "Largo Plazo (Granular)", "Corto Plazo (Cohesivo)"],
        "c o cu (kPa)": [0.0, 100.0, 0.0, 150.0],
        "phi (grados)": [28.0, 0.0, 35.0, 0.0]
    })

if 'calculado' not in st.session_state: st.session_state.calculado = False
if 'word_buffer' not in st.session_state: st.session_state.word_buffer = None

# ══════════════════════════════════════════════════════════════════════════
# INTERFAZ PRINCIPAL DE PESTAÑAS (DEFINICIÓN INICIAL)
# ══════════════════════════════════════════════════════════════════════════
tab_datos, tab_tensiones, tab_matriz_punta, tab_matriz_fuste, tab_matriz_geo, tab_matriz_tope, tab_graficas, tab_auditoria, tab_formulacion, tab_informe = st.tabs([
    "📋 1. Estratigrafía", "🌊 2. Tensiones", "🔻 3. Punta", "🟫 4. Fuste", "🌍 5. Geotécnico", "🛑 6. Tope Estruct.", "📈 7. Gráficas", "🔍 8. Auditoría", "📖 9. Formulación", "📄 10. Informe"
])

# Procesamos la tabla de entrada ANTES de la barra lateral
with tab_datos:
    st.subheader("Definición de Estratos (Datos de Entrada Puros)")
    df_edit = st.data_editor(
        st.session_state.df_base, 
        key="tabla_estratos",
        num_rows="dynamic", 
        use_container_width=True,
        column_config={
            "Condición": st.column_config.SelectboxColumn(
                options=["Corto Plazo (Cohesivo)", "Largo Plazo (Granular)"],
                required=True
            )
        }
    )
    
    # --- GUARDIÁN DE SEGURIDAD GEOTÉCNICA ---
    mascara_corto_plazo = df_edit["Condición"].str.contains("Corto Plazo")
    if (df_edit.loc[mascara_corto_plazo, "phi (grados)"] != 0.0).any():
        st.warning("⚠️ **Corrección Automática:** En condiciones a Corto Plazo (no drenadas), el ángulo de rozamiento no tiene significado físico. El sistema ha forzado internamente **φ = 0º** para esos estratos.")
        df_edit.loc[mascara_corto_plazo, "phi (grados)"] = 0.0
    # -----------------------------------------------

    z_max_total = df_edit["Espesor (m)"].sum()
    st.info(f"Profundidad máxima actual del sondeo: **{z_max_total:.2f} m**.")
    
    if not st.session_state.calculado:
        st.warning("👈 Haz clic en **'Calcular Pilotes'** en el menú izquierdo para procesar los datos.")

# ══════════════════════════════════════════════════════════════════════════
# BARRA LATERAL: CONFIGURACIÓN 
# ══════════════════════════════════════════════════════════════════════════
st.sidebar.header("💧 Nivel Freático y Seguridad")
zw = st.sidebar.number_input("Prof. Nivel Freático, zw (m)", min_value=0.0, value=3.0, step=0.5)

opcion_fs = st.sidebar.selectbox(
    "Situación de Proyecto (FS)",
    ["Casi permanente (FS ≥ 3.00)", "Característica / Transitoria (FS ≥ 2.60)", "Accidental / Sísmica (FS ≥ 2.20)", "Valor Personalizado"]
)

if opcion_fs == "Casi permanente (FS ≥ 3.00)": FS, sit_str = 3.00, "Casi Permanente"
elif opcion_fs == "Característica / Transitoria (FS ≥ 2.60)": FS, sit_str = 2.60, "Transitoria"
elif opcion_fs == "Accidental / Sísmica (FS ≥ 2.20)": FS, sit_str = 2.20, "Accidental"
else:
    FS = st.sidebar.number_input("Introduce FS personalizado:", min_value=1.0, value=3.00, step=0.1)
    sit_str = "Personalizada"

st.sidebar.header("🛑 Tope Estructural (GCOC 5.15.1)")
opcion_tope = st.sidebar.selectbox(
    "Método de Ejecución del Pilote",
    ["In situ: Seco / Lodos / Barrena (4 MPa)", 
     "In situ: Entubado (5 MPa)", 
     "In situ: Apoyo en Roca Dura (6 MPa)",
     "Valor Personalizado (MPa)"]
)

if "4 MPa" in opcion_tope: sigma_tope_mpa = 4.0
elif "5 MPa" in opcion_tope: sigma_tope_mpa = 5.0
elif "6 MPa" in opcion_tope: sigma_tope_mpa = 6.0
else: sigma_tope_mpa = st.sidebar.number_input("Tensión Tope (MPa):", min_value=1.0, value=4.0, step=0.5)

st.sidebar.header("📐 Matriz Geométrica del Pilote")
col_d1, col_d2 = st.sidebar.columns(2)
D_min = col_d1.number_input("Ø min (m)", value=0.6, min_value=0.3, step=0.1)
D_max = col_d2.number_input("Ø max (m)", value=1.5, min_value=0.3, step=0.1)
D_step = st.sidebar.number_input("Paso Ø (m)", value=0.3, min_value=0.1)

col_l1, col_l2 = st.sidebar.columns(2)
L_min_default = min(5.0, float(z_max_total)) 
L_min = col_l1.number_input("L min (m)", value=L_min_default, min_value=1.0, max_value=float(z_max_total), step=1.0)
L_max_default = min(20.0, float(z_max_total))
L_max_default = max(L_max_default, float(L_min))
L_max = col_l2.number_input("L max (m)", value=L_max_default, min_value=float(L_min), max_value=float(z_max_total), step=1.0)
L_step = st.sidebar.number_input("Paso L (m)", value=2.5, min_value=0.5)

st.sidebar.markdown("---")
if st.sidebar.button("⚙️ Calcular Pilotes", type="primary", use_container_width=True):
    st.session_state.calculado = True
    st.session_state.word_buffer = None

# ══════════════════════════════════════════════════════════════════════════
# MOTORES DE CÁLCULO
# ══════════════════════════════════════════════════════════════════════════
def calcular_perfil_tensiones(df, zw, z_max):
    z_vals = np.arange(0, z_max + 0.1, 0.1)
    sigma_v, u, sigma_v_eff = [], [], []
    sv_acum = 0.0
    for z in z_vals:
        z_acum, gamma_actual = 0.0, 18.0
        for _, row in df.iterrows():
            if z <= z_acum + row["Espesor (m)"] + 1e-5:
                gamma_actual = row["Gamma (kN/m3)"]
                break
            z_acum += row["Espesor (m)"]
        sv_acum = sv_acum + gamma_actual * 0.1 if z > 0 else 0.0
        u_val = max(0, (z - zw) * 9.81)
        sigma_v.append(sv_acum)
        u.append(u_val)
        sigma_v_eff.append(sv_acum - u_val)
    return z_vals, np.array(sigma_v), np.array(u), np.array(sigma_v_eff)

def obtener_tension_a_profundidad(z_target, z_vals, sigma_array):
    return np.interp(z_target, z_vals, sigma_array)

def calcular_pilote(D, L, df, zw, fS_val, sigma_tope_mpa):
    z_max = df["Espesor (m)"].sum()
    if L > z_max: return None 
    
    fD_calc = max(2.0/3.0, 1.0 - (D / 3.0))
    z_vals, sig_v, u, sig_v_eff = calcular_perfil_tensiones(df, zw, z_max)
    
    # --- PUNTA GCOC ---
    z_eval_punta = min(L, 20 * D)
    sig_v_eff_punta = obtener_tension_a_profundidad(z_eval_punta, z_vals, sig_v_eff)
    
    z_sup_bulbo = max(0.0, L - (6 * D))
    z_inf_bulbo = min(z_max, L + (3 * D))
    espesor_bulbo = z_inf_bulbo - z_sup_bulbo
    
    qp_eq_acumulado = 0.0
    z_acum = 0.0
    detalle_bulbo = []
    
    for _, row in df.iterrows():
        h_estrato = row["Espesor (m)"]
        z_top_estrato = z_acum
        z_bot_estrato = z_acum + h_estrato
            
        overlap_top = max(z_sup_bulbo, z_top_estrato)
        overlap_bot = min(z_inf_bulbo, z_bot_estrato)
        overlap_h = max(0.0, overlap_bot - overlap_top)
        
        if overlap_h > 0 and espesor_bulbo > 0:
            cond_i = row["Condición"]
            c_i = row["c o cu (kPa)"]
            phi_i = row["phi (grados)"]
            
            if "Corto Plazo" in cond_i:
                qp_i = 9.0 * c_i * fD_calc
            else:
                phi_rad = math.radians(phi_i)
                if phi_i == 0:
                    Nq_star, Nc_star = 1.0, 9.0 * fD_calc
                else:
                    Nq_star = 1.5 * ((1 + math.sin(phi_rad))/(1 - math.sin(phi_rad))) * math.exp(math.pi * math.tan(phi_rad)) * fD_calc
                    Nc_star = (Nq_star - 1) / math.tan(phi_rad)
                qp_i = (sig_v_eff_punta * Nq_star) + (c_i * Nc_star)
            
            peso = overlap_h / espesor_bulbo
            qp_eq_acumulado += qp_i * peso
            
            detalle_bulbo.append(f"{row['Estrato']} ({peso*100:.0f}%) -> {qp_i:.0f} kPa")
            
        z_acum += h_estrato
        
    qp = qp_eq_acumulado
    Area_pilote = (math.pi * D**2) / 4.0
    Q_punta = qp * Area_pilote
    
    auditoria_punta = {
        "Profundidad Punta (m)": L,
        "σ'_v efectiva base (kPa)": sig_v_eff_punta,
        "Factor de escala (fD)": fD_calc,
        "Resist. Unitaria q_p (kPa)": qp,
        "Fuerza Total Punta (kN)": Q_punta,
        "Composición Bulbo": " | ".join(detalle_bulbo)
    }
    
    # --- FUSTE ---
    Q_fuste, z_top, Perimetro = 0.0, 0.0, math.pi * D
    auditoria_fuste = [] 
    k0_tan_delta_fijo = 0.30 # Valor normativo inyectado automáticamente
    
    for _, row in df.iterrows():
        z_bot = z_top + row["Espesor (m)"]
        if z_top >= L: break
        
        L_tramo = min(z_bot, L) - z_top
        if L_tramo > 0:
            z_mid = z_top + (L_tramo / 2.0)
            sig_v_eff_mid = obtener_tension_a_profundidad(z_mid, z_vals, sig_v_eff)
            
            if "Corto Plazo" in row["Condición"]:
                tau_f = min(row["c o cu (kPa)"] * (100.0 / (100.0 + row["c o cu (kPa)"])), 70.0)
            else:
                # Usamos el 0.30 automático en lugar de buscarlo en la tabla
                tau_f = min(row["c o cu (kPa)"] + k0_tan_delta_fijo * sig_v_eff_mid, 90.0)
                
            Q_tramo = tau_f * Perimetro * L_tramo
            Q_fuste += Q_tramo
            
            auditoria_fuste.append({
                "Estrato": row["Estrato"],
                "Cotas (m)": f"{z_top:.1f} a {min(z_bot, L):.1f}",
                "Long. Roce (m)": L_tramo,
                "σ'_v media (kPa)": sig_v_eff_mid,
                "Resist. Unitaria τ_f (kPa)": tau_f,
                "Fuerza Tramo (kN)": Q_tramo
            })
            
        z_top = z_bot
        
    Q_total_geo = Q_punta + Q_fuste
    Q_adm_geo = Q_total_geo / fS_val
    
    # --- TOPE ESTRUCTURAL ---
    Q_tope_est = Area_pilote * (sigma_tope_mpa * 1000.0)
    Q_final_diseno = min(Q_adm_geo, Q_tope_est)
    control = "ESTRUCTURAL" if Q_tope_est < Q_adm_geo else "GEOTÉCNICO"

    return {
        "D": D, "L": L, "fD": fD_calc, 
        "Q_punta (kN)": Q_punta, "Q_fuste (kN)": Q_fuste, 
        "Q_adm_geo (kN)": Q_adm_geo,
        "Q_tope_est (kN)": Q_tope_est,
        "Q_final (kN)": Q_final_diseno,
        "Control": control,
        "auditoria_punta": auditoria_punta,
        "auditoria_fuste": auditoria_fuste 
    }

# ══════════════════════════════════════════════════════════════════════════
# GENERADOR DEL INFORME PROFESIONAL EN WORD
# ══════════════════════════════════════════════════════════════════════════
def generar_word_pilotes(df_estratos, fig_tens, df_pivot_geo, df_pivot_final, df_res, fS_val, situacion, D_array, zw_val, sigma_tope):
    doc = Document()
    
    doc.add_heading('Anejo de Cálculo: Cimentaciones Profundas', 0)
    doc.add_paragraph('Documento de cálculo analítico para el diseño de pilotes bajo cargas verticales, evaluado estrictamente según la metodología de la Guía de Cimentaciones en Obras de Carretera (GCOC).')
    
    doc.add_heading('1. Bases de Cálculo y Parámetros Iniciales', level=1)
    p_bases = doc.add_paragraph()
    p_bases.add_run(f'• Situación de Proyecto: ').bold = True
    p_bases.add_run(f'{situacion} (Factor de Seguridad Global de la Resistencia = {fS_val:.2f})\n')
    p_bases.add_run(f'• Nivel Freático: ').bold = True
    p_bases.add_run(f'Considerado a una profundidad Z = {zw_val:.2f} m\n')
    p_bases.add_run(f'• Tope Estructural del Hormigón: ').bold = True
    p_bases.add_run(f'{sigma_tope:.1f} MPa (GCOC Tabla 5.5)')

    doc.add_heading('2. Perfil Estratigráfico', level=1)
    tabla_estratos = doc.add_table(rows=1, cols=len(df_estratos.columns))
    tabla_estratos.style = 'Table Grid'
    hdr_cells = tabla_estratos.rows[0].cells
    for i, column in enumerate(df_estratos.columns): hdr_cells[i].text = str(column)
    for index, row in df_estratos.iterrows():
        row_cells = tabla_estratos.add_row().cells
        for i, value in enumerate(row): row_cells[i].text = str(value)
            
    doc.add_heading('3. Metodología Analítica Empleada', level=1)
    p_metodo = doc.add_paragraph()
    p_metodo.add_run('Resistencia por Fuste (GCOC 5.10.2.5.2): ').bold = True
    p_metodo.add_run('Aplicación de formulación hiperbólica para suelos a corto plazo y modelo de Mohr-Coulomb para largo plazo (asumiendo normativamente un K0·tan(δ) = 0.30). Se imponen topes normativos máximos de 70 kPa y 90 kPa respectivamente.\n\n')
    p_metodo.add_run('Resistencia por Punta (GCOC 5.10.2.5.1): ').bold = True
    p_metodo.add_run('Cálculo de resistencia equivalente ponderada evaluando un bulbo de rotura extendido 6D por encima y 3D por debajo de la base del pilote. Se aplica factor de reducción por efecto de escala (fD).\n\n')
    p_metodo.add_run('Diseño Final (GCOC 5.15.1): ').bold = True
    p_metodo.add_run('La carga admisible de diseño de cada pilote corresponde al valor mínimo entre la capacidad admisible del terreno y el tope estructural de la sección geométrica de hormigón.')

    doc.add_heading('4. Estado Tensional del Terreno', level=1)
    if fig_tens is not None:
        try:
            img_bytes = fig_tens.to_image(format="png", width=750, height=450)
            doc.add_picture(io.BytesIO(img_bytes), width=Inches(6.0))
        except Exception: pass
        
    doc.add_page_break()
    
    doc.add_heading(f'5. Matriz Admisible GEOTÉCNICA (kN)', level=1)
    doc.add_paragraph('Capacidad portante exclusiva del terreno (Punta + Fuste) dividida por el Factor de Seguridad.')
    tabla_geo = doc.add_table(rows=1, cols=len(df_pivot_geo.columns) + 1)
    tabla_geo.style = 'Table Grid'
    hdr_geo = tabla_geo.rows[0].cells
    hdr_geo[0].text = "L / Ø"
    for i, col_name in enumerate(df_pivot_geo.columns): hdr_geo[i+1].text = str(col_name).replace('\n', ' ')
    for index, row in df_pivot_geo.iterrows():
        row_cells = tabla_geo.add_row().cells
        row_cells[0].text = str(index)
        for i, val in enumerate(row): row_cells[i+1].text = f"{val:.0f}"

    doc.add_paragraph('\n')
    
    doc.add_heading(f'6. Matriz de Diseño FINAL LIMITADA (kN)', level=1)
    doc.add_paragraph('Mínimo entre la Matriz Geotécnica y el Tope Estructural. Los valores marcados con [EST] indican que el pilote alcanza el colapso estructural del material antes que el colapso del terreno.')
    tabla_fin = doc.add_table(rows=1, cols=len(df_pivot_final.columns) + 1)
    tabla_fin.style = 'Table Grid'
    hdr_fin = tabla_fin.rows[0].cells
    hdr_fin[0].text = "L / Ø"
    for i, col_name in enumerate(df_pivot_final.columns): hdr_fin[i+1].text = str(col_name).replace('\n', ' ')
    for index, row in df_pivot_final.iterrows():
        row_cells = tabla_fin.add_row().cells
        row_cells[0].text = str(index)
        for i, val in enumerate(row): row_cells[i+1].text = str(val)

    doc.add_page_break()
    
    doc.add_heading('7. Desgloses Gráficos por Diámetro', level=1)
    for D_val in D_array:
        df_fil = df_res[df_res["D"] == D_val]
        if df_fil.empty: continue
        fd_aplicado = df_fil["fD"].iloc[0]
        tope_val = df_fil["Q_tope_est (kN)"].iloc[0]
        
        fig_b = go.Figure()
        fig_b.add_trace(go.Bar(x=df_fil["L"], y=df_fil["Q_punta (kN)"]/fS_val, name="Por Punta", marker_color='indianred'))
        fig_b.add_trace(go.Bar(x=df_fil["L"], y=df_fil["Q_fuste (kN)"]/fS_val, name="Por Fuste", marker_color='lightsalmon'))
        fig_b.add_trace(go.Scatter(x=df_fil["L"], y=[tope_val]*len(df_fil["L"]), mode="lines", name=f"Tope Estruct. ({tope_val:.0f} kN)", line=dict(color='red', width=3, dash='dash')))
        fig_b.update_layout(barmode='stack', title=f"Desglose (Ø {D_val:.2f} m | fD = {fd_aplicado:.2f})", xaxis_title="Longitud L (m)", yaxis_title="Carga Admisible (kN)")
        
        doc.add_heading(f'Pilotes de Ø {D_val:.2f} m', level=2)
        try:
            img_bytes_break = fig_b.to_image(format="png", width=700, height=350)
            doc.add_picture(io.BytesIO(img_bytes_break), width=Inches(6.0))
        except Exception: pass

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ══════════════════════════════════════════════════════════════════════════
# CÁLCULOS PRINCIPALES Y RENDERIZADO DE RESULTADOS
# ══════════════════════════════════════════════════════════════════════════
if st.session_state.calculado:
    z_vals, sig_v, u, sig_v_eff = calcular_perfil_tensiones(df_edit, zw, z_max_total)
    
    with tab_datos:
        st.markdown("---")
        st.subheader("📊 Resistencias Unitarias Base del Terreno (GCOC)")
        st.markdown("*Nota: Fricción $\\tau_f$ media por estrato. La presión en punta $q_p$ base se muestra **sin afectar por el factor de escala** ($f_D = 1.0$).*")
        
        datos_unitarios = []
        z_top_loop = 0.0
        for idx, row_est in df_edit.iterrows():
            espesor = row_est["Espesor (m)"]
            if espesor <= 0: continue
            z_bot_loop = z_top_loop + espesor
            z_mid_loop = z_top_loop + (espesor / 2.0)
            sig_v_eff_mid_loop = obtener_tension_a_profundidad(z_mid_loop, z_vals, sig_v_eff)
            sig_v_eff_base = obtener_tension_a_profundidad(min(z_bot_loop, 20*D_min), z_vals, sig_v_eff)
            cond_loop = row_est["Condición"]
            c_loop = row_est["c o cu (kPa)"]
            phi_loop = row_est["phi (grados)"]
            
            if "Corto Plazo" in cond_loop:
                tau_f_media = min(c_loop * (100.0 / (100.0 + c_loop)), 70.0)
                qp_base = 9.0 * c_loop
            else:
                tau_f_media = min(c_loop + 0.30 * sig_v_eff_mid_loop, 90.0)
                phi_rad = math.radians(phi_loop)
                if phi_loop == 0: Nq_star, Nc_star = 1.0, 9.0
                else:
                    Nq_star = 1.5 * ((1 + math.sin(phi_rad))/(1 - math.sin(phi_rad))) * math.exp(math.pi * math.tan(phi_rad))
                    Nc_star = (Nq_star - 1) / math.tan(phi_rad)
                qp_base = (sig_v_eff_base * Nq_star) + (c_loop * Nc_star)
            
            datos_unitarios.append({
                "Estrato": row_est["Estrato"],
                "Profundidad (m)": f"de {z_top_loop:.1f} a {z_bot_loop:.1f}",
                "Fricción Fuste, τ_f media (kPa)": tau_f_media,
                "Resist. Punta Base, q_p (kPa) | fD=1": qp_base
            })
            z_top_loop = z_bot_loop
            
        st.dataframe(pd.DataFrame(datos_unitarios).style.format({
            "Fricción Fuste, τ_f media (kPa)": "{:.1f}",
            "Resist. Punta Base, q_p (kPa) | fD=1": "{:.0f}"
        }), use_container_width=True)

    # --- PERFIL TENSIONES ---
    fig_tens = go.Figure()
    fig_tens.add_trace(go.Scatter(x=sig_v, y=z_vals, name='Tensión Total (σ)', line=dict(color='black', width=2)))
    fig_tens.add_trace(go.Scatter(x=u, y=z_vals, name='Pres. Intersticial (u)', line=dict(color='blue', dash='dash')))
    fig_tens.add_trace(go.Scatter(x=sig_v_eff, y=z_vals, name='Tensión Efectiva (σ\')', line=dict(color='red', width=3)))
    
    z_acum = 0
    colores = ['#f0f8ff', '#e6e6fa', '#fff0f5', '#f5fffa', '#ffebcd']
    for i, row in df_edit.iterrows():
        z_next = z_acum + row["Espesor (m)"]
        fig_tens.add_hrect(y0=z_acum, y1=z_next, fillcolor=colores[i%len(colores)], opacity=0.4, line_width=0, annotation_text=row["Estrato"])
        z_acum = z_next
    fig_tens.add_hline(y=zw, line_dash="dot", line_color="blue", annotation_text="Nivel Freático")
    fig_tens.update_yaxes(autorange="reversed", title="Profundidad (m)")
    fig_tens.update_xaxes(title="Presión (kPa)", side="top")
    
    with tab_tensiones:
        st.plotly_chart(fig_tens, use_container_width=True)

    # --- CÁLCULO DE LA MATRIZ ---
    resultados = []
    D_arr = np.arange(D_min, D_max + D_step/2, D_step)
    L_arr = np.arange(L_min, L_max + L_step/2, L_step)

    for D in D_arr:
        for L in L_arr:
            res = calcular_pilote(D, L, df_edit, zw, FS, sigma_tope_mpa)
            if res is not None: resultados.append(res)

    df_res = pd.DataFrame(resultados)
    df_pivot_geo_global = None
    df_pivot_final_global = None

    if not df_res.empty:
        columnas_con_fd = [f"Ø {d_val:.2f} m\n(fD={df_res[df_res['D'] == d_val]['fD'].iloc[0]:.2f})" for d_val in df_res['D'].unique()]
        
        df_pivot_punta = df_res.pivot(index="L", columns="D", values="Q_punta (kN)") / FS
        df_pivot_punta.index = [f"L = {idx:.1f} m" for idx in df_pivot_punta.index]
        df_pivot_punta.columns = columnas_con_fd
        with tab_matriz_punta:
            st.subheader(f"🔻 Carga Admisible SOLO por Punta (kN) - {sit_str} (FS = {FS:.2f})")
            st.dataframe(df_pivot_punta.style.background_gradient(cmap='Reds', axis=None).format("{:.0f}"), use_container_width=True)

        df_pivot_fuste = df_res.pivot(index="L", columns="D", values="Q_fuste (kN)") / FS
        df_pivot_fuste.index = [f"L = {idx:.1f} m" for idx in df_pivot_fuste.index]
        df_pivot_fuste.columns = columnas_con_fd
        with tab_matriz_fuste:
            st.subheader(f"🟫 Carga Admisible SOLO por Fuste (kN) - {sit_str} (FS = {FS:.2f})")
            st.dataframe(df_pivot_fuste.style.background_gradient(cmap='Oranges', axis=None).format("{:.0f}"), use_container_width=True)

        df_pivot_geo = df_res.pivot(index="L", columns="D", values="Q_adm_geo (kN)")
        df_pivot_geo.index = [f"L = {idx:.1f} m" for idx in df_pivot_geo.index]
        df_pivot_geo.columns = columnas_con_fd
        df_pivot_geo_global = df_pivot_geo
        with tab_matriz_geo:
            st.subheader(f"🌍 Carga Admisible GEOTÉCNICA Total (kN) - {sit_str}")
            st.markdown("*Esta es la capacidad exclusiva del terreno (Punta + Fuste) dividida por el FS, sin importar de qué material esté hecho el pilote.*")
            st.dataframe(df_pivot_geo.style.background_gradient(cmap='Greens', axis=None).format("{:.0f}"), use_container_width=True)

        df_pivot_final = df_res.pivot(index="L", columns="D", values="Q_final (kN)")
        df_pivot_control = df_res.pivot(index="L", columns="D", values="Control")
        
        df_final_formateada = pd.DataFrame(index=df_pivot_final.index, columns=df_pivot_final.columns)
        for c in df_pivot_final.columns:
            for r in df_pivot_final.index:
                val = df_pivot_final.loc[r, c]
                ctrl = df_pivot_control.loc[r, c]
                if ctrl == "ESTRUCTURAL":
                    df_final_formateada.loc[r, c] = f"{val:.0f} [EST]"
                else:
                    df_final_formateada.loc[r, c] = f"{val:.0f}"

        df_final_formateada.index = [f"L = {idx:.1f} m" for idx in df_final_formateada.index]
        df_final_formateada.columns = columnas_con_fd
        df_pivot_final_global = df_final_formateada
        
        with tab_matriz_tope:
            st.subheader(f"🛑 Matriz de Carga de Diseño FINAL (kN) - Limitada a {sigma_tope_mpa} MPa")
            st.markdown("*Muestra el valor mínimo entre la Resistencia Geotécnica y el Tope Estructural del Hormigón. Los valores marcados con **[EST]** indican que el pilote fallará antes por el hormigón que por el terreno.*")
            
            def color_tope(val):
                color = '#ffcccc' if '[EST]' in str(val) else '#ccffcc'
                return f'background-color: {color}'
            
            st.dataframe(df_final_formateada.style.applymap(color_tope), use_container_width=True)

        with tab_graficas:
            st.subheader("Desglose de Resistencia Geotécnica vs Tope Estructural")
            
            D_target = st.selectbox("Selecciona un Diámetro para evaluar el límite estructural:", D_arr)
            df_fil = df_res[df_res["D"] == D_target]
            tope_val = df_fil["Q_tope_est (kN)"].iloc[0]
            
            fig_break_UI = go.Figure()
            fig_break_UI.add_trace(go.Bar(x=df_fil["L"], y=df_fil["Q_punta (kN)"]/FS, name="Por Punta", marker_color='indianred'))
            fig_break_UI.add_trace(go.Bar(x=df_fil["L"], y=df_fil["Q_fuste (kN)"]/FS, name="Por Fuste", marker_color='lightsalmon'))
            fig_break_UI.add_trace(go.Scatter(x=df_fil["L"], y=[tope_val]*len(df_fil["L"]), mode="lines", name=f"Tope Estruct. ({tope_val:.0f} kN)", line=dict(color='red', width=3, dash='dash')))
            
            fig_break_UI.update_layout(barmode='stack', title=f"Capacidad (Ø {D_target:.2f} m | Tope = {sigma_tope_mpa} MPa)")
            st.plotly_chart(fig_break_UI, use_container_width=True)

        with tab_auditoria:
            st.subheader("🔍 Inspector de Resistencias Unitarias y Tope Estructural")
            col_aud1, col_aud2 = st.columns(2)
            d_aud = col_aud1.selectbox("Selecciona Diámetro Ø (m):", D_arr, key="d_aud")
            l_aud = col_aud2.selectbox("Selecciona Longitud L (m):", L_arr, key="l_aud")
            
            res_auditoria = df_res[(df_res['D'] == d_aud) & (df_res['L'] == l_aud)]
            
            if not res_auditoria.empty:
                fila_aud = res_auditoria.iloc[0]
                
                st.markdown(f"### ➡️ Carga de Diseño Final: **{fila_aud['Q_final (kN)']:.0f} kN** (Controlado por: **{fila_aud['Control']}**)")
                st.markdown(f"- Admisible del Terreno: {fila_aud['Q_adm_geo (kN)']:.0f} kN")
                st.markdown(f"- Tope Estructural del Hormigón: {fila_aud['Q_tope_est (kN)']:.0f} kN")
                st.markdown("---")
                
                st.markdown(f"#### 🟫 Desglose por Fuste")
                st.dataframe(pd.DataFrame(fila_aud['auditoria_fuste']).style.format({"Long. Roce (m)": "{:.2f}", "σ'_v media (kPa)": "{:.1f}", "Resist. Unitaria τ_f (kPa)": "{:.2f}", "Fuerza Tramo (kN)": "{:.0f}"}), use_container_width=True)
                
                st.markdown(f"#### 🔻 Desglose por Punta (Método q_p Promediado)")
                st.dataframe(pd.DataFrame([fila_aud['auditoria_punta']]).style.format({
                    "Profundidad Punta (m)": "{:.2f}", 
                    "σ'_v efectiva base (kPa)": "{:.1f}", 
                    "Factor de escala (fD)": "{:.3f}", 
                    "Resist. Unitaria q_p (kPa)": "{:.2f}", 
                    "Fuerza Total Punta (kN)": "{:.0f}"
                }), use_container_width=True)

        with tab_formulacion:
            st.subheader("📖 Ecuaciones y Formulación Geotécnica")
            st.markdown("Las fórmulas empleadas en este software se rigen por la **Guía de Cimentaciones en Obras de Carretera (GCOC)**.")
            
            st.markdown("### 1. Resistencia por Punta ($q_p$)")
            st.markdown("**1.1. Factor de Efecto de Escala ($f_D$)**")
            st.markdown("Aplica una reducción a la resistencia en punta para pilotes de gran tamaño:")
            st.latex(r"f_D = 1 - \frac{D}{3} \ge \frac{2}{3}")
            
            st.markdown("**1.2. Punta a Corto Plazo (Condición No Drenada)**")
            st.latex(r"q_p = 9 \cdot c_u \cdot f_D")
            
            st.markdown("**1.3. Punta a Largo Plazo (Condición Drenada)**")
            st.latex(r"q_p = \sigma'_{v,p} \cdot N_q^* + c' \cdot N_c^*")
            st.markdown("Donde los factores de capacidad de carga, afectados por la escala, son:")
            st.latex(r"N_q^* = 1.5 \cdot \frac{1 + \sin(\phi)}{1 - \sin(\phi)} \cdot e^{\pi \cdot \tan(\phi)} \cdot f_D")
            st.latex(r"N_c^* = \frac{N_q^* - 1}{\tan(\phi)}")
            st.markdown("*Nota: La tensión vertical efectiva en punta $\sigma'_{v,p}$ está limitada normativamente a una profundidad máxima de $20D$. Además, el cálculo emplea una resistencia $q_p$ promediada en un bulbo de rotura de $6D$ y $3D$.*")
            
            st.markdown("---")
            st.markdown("### 2. Resistencia por Fuste ($\tau_f$)")
            st.markdown("**2.1. Fuste a Corto Plazo (Ley Hiperbólica)**")
            st.latex(r"\tau_f = c_u \cdot \frac{100}{100 + c_u} \le 70 \text{ kPa}")
            
            st.markdown("**2.2. Fuste a Largo Plazo (Ley Mohr-Coulomb)**")
            st.latex(r"\tau_f = c' + \sigma'_v \cdot K_0 \cdot \tan(\delta) \le 90 \text{ kPa}")
            st.markdown("*Nota: Para pilotes perforados y hormigonados in situ, el programa adopta de forma automática el valor recomendado por la GCOC de $K_0 \tan(\delta) = 0.30$.*")
            
            st.markdown("---")
            st.markdown("### 3. Tope Estructural")
            st.markdown("La capacidad final del pilote se limita por la resistencia del material constitutivo (hormigón).")
            st.latex(r"Q_{tope} = A \cdot \sigma_{tope} \cdot 1000")
            st.markdown("Donde $A$ es el área transversal del pilote ($\text{m}^2$) y $\sigma_{tope}$ es la resistencia admisible del hormigón impuesta por el método de ejecución (en MPa).")

    with tab_informe:
        st.subheader("📄 Generación de Memoria de Cálculo")
        if st.button("🛠️ Generar Informe (.docx)", type="primary"):
            if not df_res.empty and df_pivot_geo_global is not None and df_pivot_final_global is not None:
                with st.spinner("Generando documento Word..."):
                    buffer = generar_word_pilotes(df_edit, fig_tens, df_pivot_geo_global, df_pivot_final_global, df_res, FS, sit_str, D_arr, zw, sigma_tope_mpa)
                    st.session_state.word_buffer = buffer
                st.success("✅ ¡Informe generado correctamente con formato profesional!")
            else:
                st.error("No hay resultados calculados.")

        if st.session_state.word_buffer is not None:
            st.markdown("---")
            st.download_button(label="⬇️ Descargar Informe Generado", data=st.session_state.word_buffer, file_name="Anejo_Pilotes_GCOC.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")